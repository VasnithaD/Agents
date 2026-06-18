import json
import os
import smtplib
import sys
import threading
import time
import traceback
import uuid
import re
from datetime import datetime
from io import BytesIO
from email.message import EmailMessage
from pathlib import Path
 
from flask import Flask, jsonify, redirect, render_template, request, send_file, url_for
from werkzeug.serving import make_server
 
# ── Ensure project root is importable ─────────────────────────────
SCRIPT_DIR = Path(__file__).parent.resolve()
sys.path.insert(0, str(SCRIPT_DIR))
 
from requirement_extractor import (
    approve_card,
    get_project_catalog,
    refresh_brd_artifacts_from_card,
    run_requirements_pipeline,
    build_pending_approval,
    SCRIPT_DIR as REQ_DIR,
)
 
app = Flask(__name__)
app.config["SECRET_KEY"] = (os.getenv("FLASK_SECRET_KEY") or "change-me-local-secret").strip()

DEFAULT_USER = (os.getenv("APP_DEFAULT_USER") or "local-user").strip()
DEFAULT_ROLE = (os.getenv("APP_DEFAULT_ROLE") or "admin").strip().lower()

# ── Register chatbot routes (RAG-based AI chatbot) ─────────────────
try:
    from chatbot_routes import register_chatbot_routes
    register_chatbot_routes(app)
except Exception as _chatbot_routes_err:
    print(f"[app] Warning: Could not register chatbot routes: {_chatbot_routes_err}")

# ── Register GitHub integration routes (separate testable module) ───
if (os.getenv("GITHUB_INTEGRATION_ENABLED") or "true").strip().lower() in {"1", "true", "yes", "on"}:
    try:
        from github_integration import register_github_routes
        register_github_routes(app)
        print("[app] GitHub integration routes enabled at /api/github/*")
    except Exception as _github_routes_err:
        print(f"[app] Warning: Could not register github integration routes: {_github_routes_err}")

# ── Auto-indexing: initialize + watch for new documents ────────────
def _get_projects_snapshot(projects_dir: Path) -> dict:
    """Return a dict of {filepath: mtime} for all files under projects_dir."""
    snapshot = {}
    if not projects_dir.exists():
        return snapshot
    for root, _, files in os.walk(projects_dir):
        for fname in files:
            fp = Path(root) / fname
            try:
                snapshot[str(fp)] = fp.stat().st_mtime
            except OSError:
                pass
    return snapshot


def _embeddings_auto_index_worker():
    """
    Background daemon: initializes embeddings then polls the projects folder
    every 60 seconds and rebuilds the FAISS index whenever files are added
    or modified.
    """
    projects_dir = SCRIPT_DIR / "projects"
    poll_interval = 60  # seconds

    try:
        from vector_embeddings import initialize_embeddings, get_embeddings_manager
    except Exception as e:
        print(f"[embeddings] Cannot import vector_embeddings: {e}")
        return

    # Initial load (uses cached index if it exists)
    try:
        initialize_embeddings()
        print("[embeddings] Initial index ready.")
    except Exception as e:
        print(f"[embeddings] Warning: initial indexing failed: {e}")

    last_snapshot = _get_projects_snapshot(projects_dir)

    while True:
        time.sleep(poll_interval)
        try:
            current_snapshot = _get_projects_snapshot(projects_dir)
            if current_snapshot != last_snapshot:
                print("[embeddings] Project files changed — rebuilding index...")
                manager = get_embeddings_manager()
                if manager:
                    docs = manager.load_documents_from_projects()
                    chunks = manager.chunk_documents(docs)
                    if chunks:
                        manager.build_faiss_index(chunks)
                        print(f"[embeddings] Index rebuilt: {len(chunks)} chunks from {len(docs)} docs.")
                    else:
                        print("[embeddings] No chunks generated; index not updated.")
                last_snapshot = current_snapshot
        except Exception as e:
            print(f"[embeddings] Warning: auto-index error: {e}")


threading.Thread(target=_embeddings_auto_index_worker, daemon=True).start()
 
# ══════════════════════════════════════════════════════════════════
# In-memory job store
# ══════════════════════════════════════════════════════════════════
_jobs: dict = {}
_jobs_lock = threading.Lock()
 
 
def _set_job(job_id: str, status: str, result=None, error=None):
    with _jobs_lock:
        _jobs[job_id] = {"status": status, "result": result, "error": error}
 
 
def _get_job(job_id: str):
    with _jobs_lock:
        return _jobs.get(job_id)
 
 
# ══════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════
 
def _safe_path(path_str: str) -> Path:
    """Resolve path and assert it stays within SCRIPT_DIR (security check)."""
    p = Path(path_str)
    if not p.is_absolute():
        p = SCRIPT_DIR / p
    resolved = p.resolve()
    try:
        resolved.relative_to(SCRIPT_DIR.resolve())
    except ValueError:
        raise PermissionError(f"Path outside workspace: {path_str}")
    return resolved
 
 
def _read_transcript(request_form, request_files) -> tuple[str, Path]:
    """
    Extract transcript text from the request.
    Returns (text, source_path).
    Priority: file upload > file path > pasted text.
    """
    uploaded = request_files.get("transcript_file")
    if uploaded and uploaded.filename:
        original_name = Path(uploaded.filename).name or "uploaded_transcript.txt"
        suffix = Path(original_name).suffix.lower() or ".txt"
        uploaded_bytes = uploaded.read()
        src = Path(original_name)
        if suffix == ".docx":
            from docx import Document as _Doc
            doc = _Doc(BytesIO(uploaded_bytes))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = uploaded_bytes.decode("utf-8", errors="replace")
        return _validate_transcript_text(text), src
 
    path_str = (request_form.get("transcript_path") or "").strip()
    if path_str:
        p = Path(path_str)
        if not p.is_absolute():
            p = SCRIPT_DIR / p
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path_str}")
        return _validate_transcript_text(p.read_text(encoding="utf-8", errors="replace")), p
 
    pasted = (request_form.get("transcript_text") or "").strip()
    if pasted:
        return _validate_transcript_text(pasted), Path("pasted_input.txt")
 
    raise ValueError("No transcript provided (file, path, or pasted text required).")


def _validate_transcript_text(text: str) -> str:
    """
    Reject free-form content that does not resemble a speaker-based transcript.
    This keeps the Requirement Agent from generating output for arbitrary input.
    """
    cleaned = (text or "").strip()
    if not cleaned:
        raise ValueError("Invalid transcript format: transcript is empty.")

    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    if len(lines) < 4 or len(cleaned) < 120:
        raise ValueError(
            "Invalid transcript format: expected a multi-line transcript with speaker turns."
        )

    speaker_line_pattern = re.compile(
        r"^(?:Speaker\s*\d+|Interviewer|Moderator|Agent|User|"
        r"[A-Za-z][A-Za-z0-9 ._/-]{0,40})\s*:\s+\S+",
        re.IGNORECASE,
    )
    speaker_lines = [line for line in lines if speaker_line_pattern.match(line)]
    if len(speaker_lines) < 2:
        raise ValueError(
            "Invalid transcript format: expected at least two speaker-labeled lines such as 'Speaker 1: ...'."
        )

    # Reject common non-transcript payloads that can slip through as plain text.
    if re.fullmatch(r"[\s\S]*\{[\s\S]*\}[\s\S]*", cleaned) and cleaned.count(":") < 2:
        raise ValueError("Invalid transcript format: JSON-like content is not accepted as a transcript.")

    return cleaned


def _approve_card_file(card_path: Path, approved_by: str, approver_role: str, comments: str) -> dict:
    from datetime import datetime as _dt

    card = json.loads(card_path.read_text(encoding="utf-8"))
    approval = {
        "status": "approved",
        "approved_by": approved_by,
        "approver_role": approver_role,
        "approved_at": _dt.now().isoformat(),
        "comments": comments,
    }
    card["approval"] = approval
    card["status"] = "approved_for_estimation"
    card["updated_at"] = _dt.now().isoformat()
    card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
    refresh_brd_artifacts_from_card(card_path)
    return approval


def _approve_stage_card_file(
    card_path: Path,
    approval_key: str,
    approved_by: str,
    approver_role: str,
    comments: str,
    next_status: str | None = None,
    artifact_updates: dict | None = None,
) -> dict:
    from datetime import datetime as _dt

    card = json.loads(card_path.read_text(encoding="utf-8"))
    approval = {
        "status": "approved",
        "approved_by": approved_by,
        "approver_role": approver_role,
        "approved_at": _dt.now().isoformat(),
        "comments": comments,
    }
    card[approval_key] = approval
    if artifact_updates:
        artifacts = card.get("artifacts") or {}
        artifacts.update({k: v for k, v in artifact_updates.items() if v})
        card["artifacts"] = artifacts
    if next_status:
        card["status"] = next_status
    card["updated_at"] = _dt.now().isoformat()
    card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
    return approval
 
 
def _as_bool(raw: str, default: bool = False) -> bool:
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}


def _parse_ports(raw_ports: str, default: int = 5000) -> list[int]:
    """
    Parse comma-separated ports from env APP_PORTS (e.g., "5000,5001").
    Returns unique ports in input order.
    """
    text = (raw_ports or "").strip()
    if not text:
        return [default]

    ports: list[int] = []
    seen = set()
    for chunk in text.split(","):
        part = chunk.strip()
        if not part:
            continue
        try:
            port = int(part)
        except ValueError as exc:
            raise ValueError(f"Invalid port '{part}' in APP_PORTS") from exc
        if port < 1 or port > 65535:
            raise ValueError(f"Port out of range: {port}")
        if port not in seen:
            seen.add(port)
            ports.append(port)

    return ports or [default]


def _run_multi_port(app_obj, host: str, ports: list[int]):
    """
    Run the same Flask app on multiple ports using WSGI servers.
    """
    servers = []
    for port in ports:
        server = make_server(host, port, app_obj)
        t = threading.Thread(target=server.serve_forever, daemon=True)
        t.start()
        servers.append(server)
        print(f"[app] Serving on http://localhost:{port}")

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n[app] Shutting down servers...")
        for server in servers:
            server.shutdown()


def _latest_card_for_project(project_slug: str) -> dict:
    cards_dir = SCRIPT_DIR / "projects" / project_slug / "cards"
    if not cards_dir.exists():
        return {}

    candidates = sorted(
        cards_dir.glob("a2a_card_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        return {}

    try:
        return json.loads(candidates[0].read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}


def _resolve_card_artifact_path(path_str: str) -> Path | None:
    value = (path_str or "").strip()
    if not value:
        return None
    p = Path(value)
    if not p.is_absolute():
        p = SCRIPT_DIR / p
    try:
        return p.resolve()
    except OSError:
        return None


def _list_project_brd_options(project_slug: str) -> list[dict]:
    project_dir = SCRIPT_DIR / "projects" / project_slug
    cards_dir = project_dir / "cards"
    brd_dir = project_dir / "brd"

    card_by_brd: dict[str, dict] = {}
    if cards_dir.exists():
        for card_file in sorted(cards_dir.glob("a2a_card_*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
            try:
                card = json.loads(card_file.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                continue

            artifacts = card.get("artifacts") or {}
            brd_txt_path = _resolve_card_artifact_path(artifacts.get("brd_txt", ""))
            if not brd_txt_path:
                continue

            key = str(brd_txt_path)
            if key not in card_by_brd:
                approval = card.get("approval") or {}
                card_by_brd[key] = {
                    "card_path": str(card_file.resolve()),
                    "card_status": card.get("status", ""),
                    "approval_status": approval.get("status", ""),
                }

    options: list[dict] = []
    if not brd_dir.exists():
        return options

    for brd_txt in sorted(brd_dir.glob("BRD_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True):
        brd_txt_resolved = str(brd_txt.resolve())
        linked = card_by_brd.get(brd_txt_resolved, {})
        options.append({
            "brd_name": brd_txt.name,
            "brd_txt_path": brd_txt_resolved,
            "card_path": linked.get("card_path", ""),
            "card_status": linked.get("card_status", ""),
            "approval_status": linked.get("approval_status", ""),
            "has_card": bool(linked.get("card_path")),
        })

    return options


def _list_project_fds_options(project_slug: str) -> list[dict]:
    project_dir = SCRIPT_DIR / "projects" / project_slug
    cards_dir = project_dir / "cards"
    fds_dir = project_dir / "fds"

    card_by_fds: dict[str, dict] = {}
    if cards_dir.exists():
        for card_file in sorted(cards_dir.glob("a2a_card_*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
            try:
                card = json.loads(card_file.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                continue

            fds_artifacts = card.get("fds_artifacts") or {}
            fds_txt_path = _resolve_card_artifact_path(fds_artifacts.get("fds_txt", ""))
            if not fds_txt_path:
                continue

            key = str(fds_txt_path)
            if key not in card_by_fds:
                tds_card_path = _resolve_card_artifact_path(fds_artifacts.get("tds_card", ""))
                card_by_fds[key] = {
                    "card_path": str(card_file.resolve()),
                    "tds_card_path": str(tds_card_path) if tds_card_path else "",
                    "has_tds_card": bool(tds_card_path and tds_card_path.exists()),
                }

    options: list[dict] = []
    if not fds_dir.exists():
        return options

    for fds_txt in sorted(fds_dir.glob("FDS_*.txt"), key=lambda p: p.stat().st_mtime, reverse=True):
        fds_txt_resolved = str(fds_txt.resolve())
        linked = card_by_fds.get(fds_txt_resolved, {})
        options.append({
            "fds_name": fds_txt.name,
            "fds_txt_path": fds_txt_resolved,
            "card_path": linked.get("card_path", ""),
            "tds_card_path": linked.get("tds_card_path", ""),
            "has_tds_card": bool(linked.get("has_tds_card", False)),
        })

    return options


def _list_project_tds_options(project_slug: str) -> list[dict]:
    project_dir = SCRIPT_DIR / "projects" / project_slug
    tds_dir = project_dir / "tds"

    options: list[dict] = []
    if not tds_dir.exists():
        return options

    for tds_json in sorted(tds_dir.glob("TDS_*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        options.append({
            "tds_name": tds_json.name,
            "tds_json_path": str(tds_json.resolve()),
        })

    return options


def _find_card_for_brd_txt(brd_txt_path: Path) -> Path | None:
    target = str(brd_txt_path.resolve())
    projects_dir = SCRIPT_DIR / "projects"
    if not projects_dir.exists():
        return None

    candidates = sorted(
        projects_dir.glob("*/cards/a2a_card_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    for card_file in candidates:
        try:
            card = json.loads(card_file.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue

        artifacts = card.get("artifacts") or {}
        linked_brd = _resolve_card_artifact_path(artifacts.get("brd_txt", ""))
        if not linked_brd:
            continue
        if str(linked_brd) == target:
            return card_file.resolve()

    return None


def _store_uploaded_estimation_card(uploaded_file) -> Path:
    filename = Path((uploaded_file.filename or "").strip()).name
    if not filename:
        raise ValueError("Uploaded card file name is missing.")
    if Path(filename).suffix.lower() != ".json":
        raise ValueError("Uploaded card must be a .json file.")

    upload_dir = SCRIPT_DIR / "temp" / "estimation_cards"
    upload_dir.mkdir(parents=True, exist_ok=True)
    saved_path = upload_dir / f"uploaded_card_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.json"
    uploaded_file.save(saved_path)

    # Validate the uploaded JSON before running the estimation pipeline.
    json.loads(saved_path.read_text(encoding="utf-8"))
    return saved_path.resolve()


def _store_uploaded_stage_card(uploaded_file, folder_name: str) -> Path:
    filename = Path((uploaded_file.filename or "").strip()).name
    if not filename:
        raise ValueError("Uploaded card file name is missing.")
    if Path(filename).suffix.lower() != ".json":
        raise ValueError("Uploaded card must be a .json file.")

    upload_dir = SCRIPT_DIR / "temp" / folder_name
    upload_dir.mkdir(parents=True, exist_ok=True)
    saved_path = upload_dir / f"uploaded_card_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.json"
    uploaded_file.save(saved_path)

    # Validate JSON payload to fail fast on malformed uploads.
    json.loads(saved_path.read_text(encoding="utf-8"))
    return saved_path.resolve()


def _find_tds_card_for_brd_txt(brd_txt_path: Path) -> Path | None:
    base_card = _find_card_for_brd_txt(brd_txt_path)
    if not base_card or not base_card.exists():
        return None

    try:
        card = json.loads(base_card.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None

    fds_artifacts = card.get("fds_artifacts") or {}
    tds_card_path = _resolve_card_artifact_path(fds_artifacts.get("tds_card", ""))
    if not tds_card_path or not tds_card_path.exists():
        return None
    return tds_card_path


def _find_tds_card_for_fds_txt(fds_txt_path: Path) -> Path | None:
    target = str(fds_txt_path.resolve())
    projects_dir = SCRIPT_DIR / "projects"
    if not projects_dir.exists():
        return None

    candidates = sorted(
        projects_dir.glob("*/cards/a2a_card_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    for card_file in candidates:
        try:
            card = json.loads(card_file.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue

        fds_artifacts = card.get("fds_artifacts") or {}
        linked_fds = _resolve_card_artifact_path(fds_artifacts.get("fds_txt", ""))
        if not linked_fds or str(linked_fds) != target:
            continue

        tds_card = _resolve_card_artifact_path(fds_artifacts.get("tds_card", ""))
        if tds_card and tds_card.exists():
            return tds_card

    return None


def _find_tds_card_for_tds_json(tds_json_path: Path) -> Path | None:
    target = str(tds_json_path.resolve())
    projects_dir = SCRIPT_DIR / "projects"
    if not projects_dir.exists():
        return None

    candidates = sorted(
        projects_dir.glob("*/cards/*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    for card_file in candidates:
        try:
            card = json.loads(card_file.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue

        tds_artifacts = card.get("tds_artifacts") or {}
        card_artifacts = card.get("artifacts") or {}

        linked_tds_json = _resolve_card_artifact_path(
            tds_artifacts.get("tds_json") or card_artifacts.get("tds_json", "")
        )
        if not linked_tds_json:
            continue
        if str(linked_tds_json) == target:
            return card_file.resolve()

    return None


def _auto_create_estimation_card(brd_txt_path: Path) -> Path:
    """
    Auto-generate estimation A2A card for BRD if it doesn't exist.
    Returns the card path (existing or newly created).
    """
    # Check if card already exists
    existing_card = _find_card_for_brd_txt(brd_txt_path)
    if existing_card and existing_card.exists():
        return existing_card

    # Auto-create card
    project_dir = brd_txt_path.parent.parent  # Up from 'brd' to 'project'
    cards_dir = project_dir / "cards"
    cards_dir.mkdir(parents=True, exist_ok=True)

    card_name = brd_txt_path.stem + "_a2a_card.json"
    card_path = cards_dir / card_name

    # If file already exists (edge case), return it
    if card_path.exists():
        return card_path

    # Create minimal A2A card
    new_card = {
        "status": "estimation_ready",
        "project": {
            "name": project_dir.name,
            "brd_path": str(brd_txt_path.resolve())
        },
        "payload": {
            "brd_content": f"[Linked to {brd_txt_path.name}]"
        },
        "artifacts": {
            "brd_txt": str(brd_txt_path.resolve()),
            "output_dir": str(project_dir.resolve())
        },
        "created_at": datetime.now().isoformat(),
        "created_by": "auto_generation"
    }

    try:
        card_path.write_text(json.dumps(new_card, indent=2), encoding="utf-8")
        print(f"[auto-card] ✅ Created estimation card: {card_path}")
    except OSError as exc:
        print(f"[auto-card] ❌ Failed to create card: {exc}")
        raise

    return card_path


def _auto_create_fds_tds_card(brd_txt_path: Path, estimation_card_path: Path) -> Path:
    """
    Auto-generate FDS→TDS A2A card if it doesn't exist.
    Links it to the estimation card.
    Returns the card path (existing or newly created).
    """
    project_dir = brd_txt_path.parent.parent
    cards_dir = project_dir / "cards"

    card_name = brd_txt_path.stem + "_fds_tds_a2a_card.json"
    card_path = cards_dir / card_name

    # If card already exists, return it
    if card_path.exists():
        return card_path

    # Create FDS→TDS card
    new_card = {
        "status": "ready_for_tds",
        "project": {
            "name": project_dir.name
        },
        "payload": {
            "fds_sections": {
                "functional_requirements": "[Auto-linked to FDS output]",
                "integration_points": "[Pending FDS generation]",
                "data_requirements": "[Pending FDS generation]"
            },
            "summary": {}
        },
        "artifacts": {
            "output_dir": str(project_dir.resolve()),
            "estimation_card": str(estimation_card_path.resolve())
        },
        "created_at": datetime.now().isoformat(),
        "created_by": "auto_generation"
    }

    try:
        card_path.write_text(json.dumps(new_card, indent=2), encoding="utf-8")
        print(f"[auto-card] ✅ Created FDS→TDS card: {card_path}")
    except OSError as exc:
        print(f"[auto-card] ❌ Failed to create FDS→TDS card: {exc}")
        raise

    return card_path


def _extract_business_from_text(text: str) -> str:
    if not text:
        return ""

    patterns = [
        r"\bfor\s+([A-Z][A-Za-z0-9&' -]{2,70})",
        r"\b([A-Z][A-Za-z0-9&' -]{2,70})\s+aims\b",
        r"\b([A-Z][A-Za-z0-9&' -]{2,70})\s+plans\b",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            return m.group(1).strip(" .,")
    return ""


_CHAT_STOPWORDS = {
    "the", "is", "are", "a", "an", "and", "or", "for", "to", "of", "in", "on", "at",
    "with", "about", "project", "projects", "this", "that", "it", "its", "from", "by", "as",
    "me", "show", "tell", "give", "what", "who", "how", "when", "where", "why", "which",
    "can", "could", "would", "should", "please", "any", "all", "our", "we", "us",
}


def _tokenize_query(text: str) -> list[str]:
    raw = re.findall(r"[a-zA-Z0-9]{3,}", (text or "").lower())
    return [tok for tok in raw if tok not in _CHAT_STOPWORDS]


def _canonicalize_common_question(text: str) -> str:
    q = (text or "").strip().lower()
    q = re.sub(r"[^a-z0-9\s?]", " ", q)
    q = re.sub(r"\s+", " ", q).strip()

    canonical_map = [
        (r"^(hi|hello|hey|yo|hola)$", "hi"),
        (r"^who (is|s) (the )?business.*$", "who is the business for this project"),
        (r"^who (are|is) (the )?(stakeholders|team leads|team|owners?).*$", "who are the stakeholders or team leads"),
        (r"^how (is|do you do) estimation.*$", "how is estimation defined"),
        (r"^what does fds mean[s]?$", "what does fds mean"),
        (r"^what does tds mean[s]?$", "what does tds mean"),
        (r"^what does brd mean[s]?$", "what does brd mean"),
        (r"^what does it do$", "what does it do"),
        (r"^list project status summary$", "list project status summary"),
        (r"^summari[sz]e (this )?project( end to end)?$", "summarize this project end-to-end"),
    ]
    for pattern, canonical in canonical_map:
        if re.match(pattern, q):
            return canonical
    return (text or "").strip()


def _suggested_questions_from_history(history: list, defaults: list[str], limit: int = 4) -> list[str]:
    if not isinstance(defaults, list):
        defaults = []

    learned: list[str] = []
    seen = set()
    for turn in reversed(history if isinstance(history, list) else []):
        if not isinstance(turn, dict):
            continue
        if str(turn.get("role") or "").lower() != "user":
            continue

        message = str(turn.get("message") or "").strip()
        if not message:
            continue

        normalized = _canonicalize_common_question(message)
        nl = normalized.lower()
        if nl in {"hi", "hello", "hey", "yo", "hola"}:
            continue
        if len(normalized) < 8 or len(normalized) > 90:
            continue
        if nl in seen:
            continue
        seen.add(nl)
        learned.append(normalized)
        if len(learned) >= 2:
            break

    result = learned[:]
    for d in defaults:
        d = str(d or "").strip()
        if not d:
            continue
        if d.lower() in seen:
            continue
        seen.add(d.lower())
        result.append(d)
        if len(result) >= limit:
            break
    return result[:limit]


def _best_matches(items: list[str], query_tokens: list[str], limit: int = 5) -> list[str]:
    if not items:
        return []
    if not query_tokens:
        return items[:limit]

    scored: list[tuple[int, str]] = []
    for item in items:
        item_l = (item or "").lower()
        score = sum(1 for tok in query_tokens if tok in item_l)
        if score > 0:
            scored.append((score, item))

    if not scored:
        return items[: min(limit, 3)]

    scored.sort(key=lambda x: x[0], reverse=True)
    return [text for _, text in scored[:limit]]


def _project_line(p: dict) -> str:
    return (
        f"- {p.get('name', p.get('slug', ''))} ({p.get('slug', '')}): "
        f"status={p.get('status', 'unknown')}, req={p.get('requirements_count', 0)}, "
        f"risks={p.get('risks_count', 0)}, assumptions={p.get('assumptions_count', 0)}, "
        f"dependencies={p.get('dependencies_count', 0)}"
    )


def _project_snapshot(selected: dict) -> str:
    return (
        f"Project: {selected.get('name', selected.get('slug', ''))}\n"
        f"Status: {selected.get('status', 'unknown')}\n"
        f"Business: {selected.get('business') or 'Not explicitly identified'}\n"
        f"Stakeholders: {selected.get('stakeholders') or 'Not listed'}\n"
        f"Requirements: {selected.get('requirements_count', 0)} | Risks: {selected.get('risks_count', 0)} | "
        f"Assumptions: {selected.get('assumptions_count', 0)} | Constraints: {selected.get('constraints_count', 0)} | "
        f"Dependencies: {selected.get('dependencies_count', 0)}"
    )


def _format_bullets(items: list[str], empty_text: str, limit: int = 6) -> str:
    if not items:
        return empty_text
    return "\n".join(f"- {x}" for x in items[:limit])


def _build_project_knowledge() -> dict:
    catalog = get_project_catalog(SCRIPT_DIR)
    projects = []

    for item in catalog:
        card = _latest_card_for_project(item["slug"])
        payload = card.get("payload", {}) if isinstance(card, dict) else {}
        overview = payload.get("overview", {}) if isinstance(payload, dict) else {}
        stats = card.get("stats", {}) if isinstance(card, dict) else {}

        purpose = str(overview.get("PURPOSE", "")).strip()
        background = str(overview.get("BACKGROUND", "")).strip()
        stakeholders = str(overview.get("STAKEHOLDERS", "")).strip()
        business_entity = _extract_business_from_text(purpose) or _extract_business_from_text(background)

        requirements_list = payload.get("requirements", []) if isinstance(payload, dict) else []
        requirements_text = []
        for req in requirements_list if isinstance(requirements_list, list) else []:
            if not isinstance(req, dict):
                continue
            rid = str(req.get("id", "")).strip()
            category = str(req.get("category", "")).strip()
            priority = str(req.get("priority", "")).strip()
            desc = str(req.get("description", "")).strip()
            if desc:
                prefix = " ".join(part for part in [rid, category, priority] if part)
                requirements_text.append((prefix + " - " + desc) if prefix else desc)

        risks = payload.get("risks", []) if isinstance(payload.get("risks", []), list) else []
        assumptions = payload.get("assumptions", []) if isinstance(payload.get("assumptions", []), list) else []
        constraints = payload.get("constraints", []) if isinstance(payload.get("constraints", []), list) else []
        dependencies = payload.get("dependencies", []) if isinstance(payload.get("dependencies", []), list) else []

        projects.append({
            "slug": item.get("slug", ""),
            "name": item.get("display_name", ""),
            "status": item.get("status", "unknown"),
            "business": business_entity,
            "stakeholders": stakeholders,
            "purpose": purpose,
            "scope_in": str(overview.get("SCOPE_IN", "")).strip(),
            "scope_out": str(overview.get("SCOPE_OUT", "")).strip(),
            "requirements_count": int(stats.get("total_requirements", 0) or 0),
            "risks_count": int(stats.get("total_risks", 0) or 0),
            "assumptions_count": int(stats.get("total_assumptions", 0) or 0),
            "constraints_count": int(stats.get("total_constraints", 0) or 0),
            "dependencies_count": int(stats.get("total_dependencies", 0) or 0),
            "requirements": requirements_text,
            "risks": [str(x).strip() for x in risks if str(x).strip()],
            "assumptions": [str(x).strip() for x in assumptions if str(x).strip()],
            "constraints": [str(x).strip() for x in constraints if str(x).strip()],
            "dependencies": [str(x).strip() for x in dependencies if str(x).strip()],
            "approval_status": str((card.get("approval") or {}).get("status", "")).strip(),
            "estimation_approval_status": str((card.get("estimation_approval") or {}).get("status", "")).strip(),
            "fds_approval_status": str((card.get("fds_approval") or {}).get("status", "")).strip(),
            "tds_approval_status": str((card.get("tds_approval") or {}).get("status", "")).strip(),
            "has_fds": bool(card.get("fds_artifacts")),
            "has_tds": bool(card.get("tds_artifacts")),
        })

    return {
        "generated_at": datetime.now().isoformat(),
        "projects": projects,
    }


def _project_from_history(history: list) -> str:
    if not isinstance(history, list):
        return ""

    for turn in reversed(history):
        if not isinstance(turn, dict):
            continue
        meta = turn.get("meta") or {}
        if isinstance(meta, dict):
            project_slug = str(meta.get("project") or "").strip()
            if project_slug:
                return project_slug
    return ""


def _last_defined_term_from_history(history: list) -> str:
    if not isinstance(history, list):
        return ""

    for turn in reversed(history):
        if not isinstance(turn, dict):
            continue
        meta = turn.get("meta") or {}
        if not isinstance(meta, dict):
            continue
        if str(meta.get("intent") or "").strip() != "term_definition":
            continue
        term = str(meta.get("term") or "").strip().lower()
        if term:
            return term
    return ""


def _pick_project(knowledge: dict, question: str, preferred_slug: str = "", history: list | None = None) -> dict:
    projects = knowledge.get("projects", [])
    if not projects:
        return {}

    if preferred_slug:
        for p in projects:
            if p.get("slug") == preferred_slug:
                return p

    previous_slug = _project_from_history(history or [])
    if previous_slug:
        for p in projects:
            if p.get("slug") == previous_slug:
                return p

    q = (question or "").lower()
    for p in projects:
        slug = p.get("slug", "").lower()
        name = p.get("name", "").lower()
        if slug and slug in q:
            return p
        if name and name in q:
            return p

    return projects[0]


def _answer_chat_question(question: str, preferred_slug: str = "", history: list | None = None) -> tuple[str, dict]:
    q = _canonicalize_common_question((question or "").strip())
    if not q:
        return "Please ask a question about projects, business context, stakeholders, or estimation.", {"intent": "empty"}

    knowledge = _build_project_knowledge()
    projects = knowledge.get("projects", [])
    if not projects:
        return "No project data found yet. Generate at least one BRD/A2A card first so I can answer project questions.", {"intent": "no_data"}

    ql = q.lower()
    selected = _pick_project(knowledge, q, preferred_slug, history)
    default_suggested = [
        "Summarize this project end-to-end",
        "List top risks and dependencies",
        "Show high-priority requirements",
        "What is the current approval/status state?",
    ]
    suggested = _suggested_questions_from_history(history or [], default_suggested)

    greeting_words = {
        "hi", "hello", "hey", "yo", "hola", "good morning", "good afternoon", "good evening"
    }
    if ql in greeting_words:
        answer = (
            f"Hi! I can help with project questions for '{selected.get('name', selected.get('slug', 'this project'))}'. "
            "Ask me about summary, requirements, risks, dependencies, approvals, scope, or estimation."
        )
        return answer, {"intent": "greeting", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["estimation", "estimate", "sizing", "investment", "hours"]):
        answer = (
            "Estimation is derived per requirement using 5 complexity dimensions: "
            "Integration, Data Volume, Business Logic, UI Complexity, and NFR Burden. "
            "The total effort is then mapped to a size bucket (XXS to XXXL) and an investment range."
        )
        return answer, {"intent": "estimation_definition", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["team", "lead", "manager", "stakeholder", "owner", "pm"]):
        names = selected.get("stakeholders", "") or "No named stakeholders available in the latest BRD card."
        answer = (
            f"For project '{selected.get('name', selected.get('slug', ''))}', the listed stakeholders are: {names}. "
            "If you want clear role mapping (Team Lead vs Project Manager), add those roles explicitly in the transcript/BRD so the card can capture them."
        )
        return answer, {"intent": "team_info", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["business", "client", "customer", "who is the business", "for which project"]):
        business = selected.get("business", "") or "Not explicitly identified"
        answer = (
            f"Project '{selected.get('name', selected.get('slug', ''))}' appears to be for: {business}. "
            f"Purpose summary: {selected.get('purpose', 'Not available')}"
        )
        return answer, {"intent": "business_info", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["all projects", "across projects", "compare", "comparison"]):
        lines = [_project_line(p) for p in projects]
        return (
            "Cross-project view:\n" + "\n".join(lines) +
            "\n\nAsk follow-up like: 'compare risks only' or 'show details for <project name>'."
        ), {"intent": "cross_project_summary", "suggested_questions": suggested}

    if any(k in ql for k in ["list", "projects", "status", "summary", "information", "overview"]):
        lines = []
        for p in projects:
            lines.append(_project_line(p))
        return "Available project information:\n" + "\n".join(lines), {"intent": "project_summary", "suggested_questions": suggested}

    if any(k in ql for k in ["scope", "in scope", "out of scope"]):
        answer = (
            f"Project '{selected.get('name', selected.get('slug', ''))}' scope in: {selected.get('scope_in', 'Not available')}\n"
            f"Scope out: {selected.get('scope_out', 'Not available')}"
        )
        return answer, {"intent": "scope", "project": selected.get("slug", ""), "suggested_questions": suggested}

    # Handle acronym meaning/definition questions before stage/status routing.
    if any(k in ql for k in ["mean", "means", "full form", "stand for", "stands for", "what is", "what does"]):
        glossary = {
            "brd": "BRD means Business Requirements Document.",
            "fds": "FDS means Functional Design Specification.",
            "tds": "TDS means Technical Design Specification.",
            "a2a": "A2A card means Agent-to-Agent handoff card used between pipeline stages.",
            "nfr": "NFR means Non-Functional Requirement.",
            "otp": "OTP means One-Time Password.",
            "api": "API means Application Programming Interface.",
            "ui": "UI means User Interface.",
            "ux": "UX means User Experience.",
            "sdlc": "SDLC means Software Development Life Cycle.",
        }
        term_actions = {
            "brd": "BRD captures business goals, scope, stakeholders, and requirements so downstream agents can estimate and design accurately.",
            "fds": "FDS translates approved requirements into functional behavior: feature flows, module behavior, validations, business rules, and handoff details for technical design.",
            "tds": "TDS converts functional design into technical implementation detail: architecture, components, interfaces, data structures, and sequence/flow diagrams.",
            "a2a": "The A2A card is the handoff package between agents, carrying structured project context, artifacts, approvals, and status between stages.",
            "nfr": "NFR defines quality attributes such as performance, reliability, security, usability, and maintainability.",
            "otp": "OTP is used for secure one-time user verification during sign-in or sensitive actions.",
            "api": "An API is the contract by which systems exchange data and trigger operations.",
            "ui": "UI is the visual and interactive layer users directly work with.",
            "ux": "UX is the end-to-end user experience, including ease of use, flow clarity, and satisfaction.",
            "sdlc": "SDLC is the full lifecycle from requirements through estimation, design, build, test, and release.",
        }
        found_keys = [k for k in glossary if re.search(r"\b" + re.escape(k) + r"\b", ql)]
        found = [glossary[k] for k in found_keys]
        if found:
            answer = "\n".join(found)
            meta = {"intent": "term_definition", "project": selected.get("slug", ""), "suggested_questions": suggested}
            if found_keys:
                meta["term"] = found_keys[0]
                meta["term_action"] = term_actions.get(found_keys[0], "")
            return answer, meta

    follow_up_term = _last_defined_term_from_history(history or [])
    if follow_up_term and any(k in ql for k in ["it", "that", "this", "do", "does", "work", "purpose", "use"]):
        follow_up_actions = {
            "brd": "BRD is used to align business and delivery teams on what should be built, why it matters, and what is in or out of scope.",
            "fds": "FDS defines how the approved requirements should function in the solution. It is the bridge between requirements and technical design, and it prepares handoff data for TDS.",
            "tds": "TDS specifies how the system will be implemented technically, including architecture and integration details used by development teams.",
            "a2a": "The A2A card carries approved structured outputs from one agent to the next so each stage can run with full context.",
            "nfr": "NFR guides non-functional behavior targets like performance, security, and reliability during design and implementation.",
            "otp": "OTP secures user access by validating identity with a one-time code.",
            "api": "API enables service-to-service communication and data exchange across components.",
            "ui": "UI lets users perform tasks through screens, forms, and controls.",
            "ux": "UX ensures the product flow is intuitive and effective for users.",
            "sdlc": "SDLC organizes end-to-end software delivery from planning to release.",
        }
        meaning_map = {
            "brd": "Business Requirements Document",
            "fds": "Functional Design Specification",
            "tds": "Technical Design Specification",
            "a2a": "Agent-to-Agent handoff card",
            "nfr": "Non-Functional Requirement",
            "otp": "One-Time Password",
            "api": "Application Programming Interface",
            "ui": "User Interface",
            "ux": "User Experience",
            "sdlc": "Software Development Life Cycle",
        }
        answer = (
            f"{follow_up_term.upper()} ({meaning_map.get(follow_up_term, follow_up_term.upper())}) is used in this pipeline as follows: "
            + follow_up_actions.get(follow_up_term, "It provides structure and guidance for this stage of the project lifecycle.")
        )
        return answer, {"intent": "term_follow_up", "term": follow_up_term, "project": selected.get("slug", ""), "suggested_questions": suggested}

    query_tokens = _tokenize_query(q)

    if not query_tokens:
        answer = (
            "I can help with project-related details, but I need a specific question. "
            "Try: 'summarize this project', 'list risks', 'show dependencies', or 'approval status'."
        )
        return answer, {"intent": "clarify", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["requirement", "feature", "need", "must"]) or "req-" in ql:
        matches = _best_matches(selected.get("requirements", []), query_tokens)
        answer = (
            f"Requirements for '{selected.get('name', selected.get('slug', ''))}':\n" +
            _format_bullets(matches, "No requirement details found in the latest card.")
        )
        return answer, {"intent": "requirements", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["risk", "issue", "threat", "challenge"]):
        matches = _best_matches(selected.get("risks", []), query_tokens)
        answer = (
            f"Risks for '{selected.get('name', selected.get('slug', ''))}':\n" +
            _format_bullets(matches, "No risks listed in the latest card.")
        )
        return answer, {"intent": "risks", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["assumption", "assume"]):
        matches = _best_matches(selected.get("assumptions", []), query_tokens)
        answer = (
            f"Assumptions for '{selected.get('name', selected.get('slug', ''))}':\n" +
            _format_bullets(matches, "No assumptions listed in the latest card.")
        )
        return answer, {"intent": "assumptions", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["constraint", "limitation", "not include", "out of scope"]):
        matches = _best_matches(selected.get("constraints", []), query_tokens)
        answer = (
            f"Constraints for '{selected.get('name', selected.get('slug', ''))}':\n" +
            _format_bullets(matches, "No constraints listed in the latest card.")
        )
        return answer, {"intent": "constraints", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["dependency", "depends", "integration", "gateway"]):
        matches = _best_matches(selected.get("dependencies", []), query_tokens)
        answer = (
            f"Dependencies for '{selected.get('name', selected.get('slug', ''))}':\n" +
            _format_bullets(matches, "No dependencies listed in the latest card.")
        )
        return answer, {"intent": "dependencies", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["approval", "approved", "stage", "pipeline", "status of fds", "status of tds"]):
        answer = (
            f"Approval and stage state for '{selected.get('name', selected.get('slug', ''))}':\n"
            f"- Overall status: {selected.get('status', 'unknown')}\n"
            f"- BRD approval: {selected.get('approval_status') or 'not recorded'}\n"
            f"- Estimation approval: {selected.get('estimation_approval_status') or 'not recorded'}\n"
            f"- FDS approval: {selected.get('fds_approval_status') or 'not recorded'}\n"
            f"- TDS approval: {selected.get('tds_approval_status') or 'not recorded'}"
        )
        return answer, {"intent": "approval_status", "project": selected.get("slug", ""), "suggested_questions": suggested}

    corpus = []
    corpus.extend(selected.get("requirements", []))
    corpus.extend(selected.get("risks", []))
    corpus.extend(selected.get("assumptions", []))
    corpus.extend(selected.get("constraints", []))
    corpus.extend(selected.get("dependencies", []))
    if selected.get("purpose"):
        corpus.append("Purpose: " + selected.get("purpose", ""))
    if selected.get("background"):
        corpus.append("Background: " + selected.get("background", ""))
    if selected.get("scope_in"):
        corpus.append("Scope in: " + selected.get("scope_in", ""))
    if selected.get("scope_out"):
        corpus.append("Scope out: " + selected.get("scope_out", ""))

    matches = _best_matches(corpus, query_tokens, limit=4)
    if matches:
        answer = (
            f"Here is what I found for '{selected.get('name', selected.get('slug', ''))}' related to your question:\n" +
            _format_bullets(matches, "No relevant details found.") +
            "\n\nIf you want, ask a narrower follow-up (for example: risks, dependencies, approvals, or specific requirements)."
        )
        return answer, {"intent": "semantic_match", "project": selected.get("slug", ""), "suggested_questions": suggested}

    if any(k in ql for k in ["what about", "and", "also", "more", "next"]) and selected.get("name"):
        answer = (
            f"We are currently discussing project '{selected.get('name', selected.get('slug', ''))}'. "
            "I can continue with requirements, risks, assumptions, constraints, dependencies, approvals, scope, status, or estimation. "
            "Tell me what you want next."
        )
        return answer, {"intent": "follow_up", "project": selected.get("slug", ""), "suggested_questions": suggested}

    answer = (
        "I can answer most project-related questions from BRD/A2A data: overview, stakeholders, requirements, risks, assumptions, "
        "constraints, dependencies, approvals, and pipeline status.\n\n"
        + _project_snapshot(selected)
        + "\n\nTry: 'show high-priority requirements', 'list dependencies', 'what are the main risks', or 'compare all projects'."
    )
    return answer, {"intent": "fallback", "project": selected.get("slug", ""), "suggested_questions": suggested}
 
 
def _send_email_smtp(to_email: str, subject: str, body: str, attachment_path: str = ""):
    smtp_host = (os.getenv("SMTP_HOST") or "").strip()
    smtp_port = int((os.getenv("SMTP_PORT") or "587").strip())
    smtp_user = (os.getenv("SMTP_USER") or "").strip()
    smtp_pass = (os.getenv("SMTP_PASS") or "").strip()
    smtp_from = (os.getenv("SMTP_FROM") or smtp_user).strip()
    smtp_use_tls = _as_bool(os.getenv("SMTP_USE_TLS"), True)
    smtp_use_ssl = _as_bool(os.getenv("SMTP_USE_SSL"), False)
 
    if not smtp_host:
        raise RuntimeError("SMTP_HOST is missing in .env")
    if not smtp_from:
        raise RuntimeError("SMTP_FROM (or SMTP_USER) is missing in .env")
 
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp_from
    msg["To"] = to_email
    msg.set_content(body or "BRD report attached/generated from Requirement Agent.")
 
    if attachment_path:
        p = _safe_path(attachment_path)
        if not p.exists():
            raise FileNotFoundError(f"Attachment not found: {attachment_path}")
        data = p.read_bytes()
        # Send as text/plain for .txt and octet-stream for other files.
        subtype = "plain" if p.suffix.lower() == ".txt" else "octet-stream"
        msg.add_attachment(
            data,
            maintype="text" if subtype == "plain" else "application",
            subtype=subtype,
            filename=p.name,
        )
 
    if smtp_use_ssl:
        with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=30) as server:
            if smtp_user and smtp_pass:
                server.login(smtp_user, smtp_pass)
            server.send_message(msg)
    else:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as server:
            server.ehlo()
            if smtp_use_tls:
                server.starttls()
                server.ehlo()
            if smtp_user and smtp_pass:
                server.login(smtp_user, smtp_pass)
            server.send_message(msg)


def _generic_mail_body(subject: str, attachment_path: str = "") -> str:
    subject_l = (subject or "").lower()
    attachment_name = Path(attachment_path).name.lower() if attachment_path else ""

    doc_type = "document"
    if "estimation" in subject_l or "estimation" in attachment_name:
        doc_type = "estimation"
    elif "fds" in subject_l or "fds" in attachment_name:
        doc_type = "FDS"
    elif "tds" in subject_l or "tds" in attachment_name:
        doc_type = "TDS"
    elif any(k in subject_l for k in ["brd", "requirement"]) or any(
        k in attachment_name for k in ["brd", "requirement"]
    ):
        doc_type = "requirements"

    return (
        "Hi,\n\n"
        f"Your {doc_type} is ready. Please find the attached file.\n\n"
        "Thanks."
    )


def _read_env_file_value(env_path: Path, key: str) -> str:
    """Read a single key from an env file without loading the whole file into process env."""
    if not env_path.exists():
        return ""

    prefix = f"{key}="
    try:
        for raw_line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or not line.startswith(prefix):
                continue
            value = line[len(prefix):].strip()
            if len(value) >= 2 and value[0] == value[-1] and value[0] in {'"', "'"}:
                value = value[1:-1]
            return value.strip()
    except Exception:
        return ""

    return ""


def _normalize_host(host: str) -> str:
    host_clean = (host or "").strip()
    if not host_clean or host_clean in {"0.0.0.0", "::"}:
        return "localhost"
    return host_clean


def _resolve_codex_portal_url() -> str:
    """
    Resolve the URL opened from the main Agents page Codex button.

    Priority:
    1) CODEX_PORTAL_URL env override
    2) MCP_SERVER_URL env
    3) mcp-client/.env MCP_SERVER_URL
    4) mcp-client/.env MCP_SERVER_HOST + MCP_SERVER_PORT
    5) fallback localhost:3000
    """
    direct_url = (os.getenv("CODEX_PORTAL_URL") or "").strip()
    if direct_url:
        return direct_url

    env_server_url = (os.getenv("MCP_SERVER_URL") or "").strip()
    if env_server_url:
        return env_server_url

    mcp_env_path = SCRIPT_DIR / "mcp-client" / ".env"

    mcp_server_url = _read_env_file_value(mcp_env_path, "MCP_SERVER_URL")
    if mcp_server_url:
        return mcp_server_url

    mcp_host = _normalize_host(_read_env_file_value(mcp_env_path, "MCP_SERVER_HOST"))
    mcp_port = (_read_env_file_value(mcp_env_path, "MCP_SERVER_PORT") or "3000").strip()

    if not mcp_port.isdigit():
        mcp_port = "3000"

    return f"http://{mcp_host}:{mcp_port}"
 
 
# ══════════════════════════════════════════════════════════════════
# Routes
# ══════════════════════════════════════════════════════════════════


def _render_main_page(active_page: str):
    return render_template(
        "index.html",
        active_page=active_page,
        current_user=DEFAULT_USER,
        current_role=DEFAULT_ROLE,
        codex_portal_url=_resolve_codex_portal_url(),
    )


@app.route("/login", methods=["GET", "POST"])
def login_page():
    return redirect(url_for("home_page"))


@app.route("/logout", methods=["GET"])
def logout_page():
    return redirect(url_for("home_page"))
 
@app.route("/")
def index():
    return redirect(url_for("home_page"))


@app.route("/home")
def home_page():
    return _render_main_page("home")


@app.route("/requirement")
def requirement_page():
    return _render_main_page("requirement")


@app.route("/estimation")
def estimation_page():
    return _render_main_page("estimation")


@app.route("/fds")
def fds_page():
    return _render_main_page("fds")


@app.route("/tds")
def tds_page():
    return _render_main_page("tds")
 
 
@app.route("/api/projects", methods=["GET"])
def list_projects():
    try:
        catalog = get_project_catalog(SCRIPT_DIR)
        return jsonify(catalog)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/project-brd-options", methods=["GET"])
def project_brd_options():
    project_slug = (request.args.get("project_slug") or "").strip()
    if not project_slug:
        return jsonify({"error": "project_slug is required"}), 400

    try:
        catalog = get_project_catalog(SCRIPT_DIR)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    project_map = {p.get("slug", ""): p for p in catalog}
    selected = project_map.get(project_slug)
    if not selected:
        return jsonify({"error": f"Project slug not found: {project_slug}"}), 404

    options = _list_project_brd_options(project_slug)
    return jsonify({
        "project_slug": project_slug,
        "project_name": selected.get("display_name", project_slug),
        "options": options,
    })


@app.route("/api/project-fds-options", methods=["GET"])
def project_fds_options():
    project_slug = (request.args.get("project_slug") or "").strip()
    if not project_slug:
        return jsonify({"error": "project_slug is required"}), 400

    try:
        catalog = get_project_catalog(SCRIPT_DIR)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    project_map = {p.get("slug", ""): p for p in catalog}
    selected = project_map.get(project_slug)
    if not selected:
        return jsonify({"error": f"Project slug not found: {project_slug}"}), 404

    options = _list_project_fds_options(project_slug)
    return jsonify({
        "project_slug": project_slug,
        "project_name": selected.get("display_name", project_slug),
        "options": options,
    })


@app.route("/api/project-tds-options", methods=["GET"])
def project_tds_options():
    project_slug = (request.args.get("project_slug") or "").strip()
    if not project_slug:
        return jsonify({"error": "project_slug is required"}), 400

    try:
        catalog = get_project_catalog(SCRIPT_DIR)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    project_map = {p.get("slug", ""): p for p in catalog}
    selected = project_map.get(project_slug)
    if not selected:
        return jsonify({"error": f"Project slug not found: {project_slug}"}), 404

    options = _list_project_tds_options(project_slug)
    return jsonify({
        "project_slug": project_slug,
        "project_name": selected.get("display_name", project_slug),
        "options": options,
    })


@app.route("/api/brd-journey", methods=["GET"])
def brd_journey():
    """Return the SDLC journey status for a single BRD's A2A card."""
    card_path_str = (request.args.get("card_path") or "").strip()
    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": "Card file not found"}), 404

    try:
        card = json.loads(card_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        return jsonify({"error": f"Cannot read card: {exc}"}), 500

    artifacts   = card.get("artifacts") or {}
    est_arts    = card.get("estimation_artifacts") or {}
    fds_arts    = card.get("fds_artifacts") or {}
    tds_arts    = card.get("tds_artifacts") or {}
    approval    = card.get("approval") or {}
    est_app     = card.get("estimation_approval") or {}
    fds_app     = card.get("fds_approval") or {}
    tds_app     = card.get("tds_approval") or {}

    def _exists(path_str):
        if not path_str:
            return False
        try:
            return Path(path_str).exists()
        except Exception:
            return False

    def _any_exists(*paths):
        return any(_exists(p) for p in paths)

    linked_tds_card = {}
    linked_tds_arts = {}
    linked_tds_app = {}
    linked_fds_app = {}
    linked_tds_artifacts = {}
    linked_tds_card_path = _resolve_card_artifact_path(fds_arts.get("tds_card", ""))
    if linked_tds_card_path and linked_tds_card_path.exists():
        try:
            linked_tds_card = json.loads(linked_tds_card_path.read_text(encoding="utf-8"))
            linked_tds_arts = linked_tds_card.get("tds_artifacts") or {}
            linked_tds_app = linked_tds_card.get("tds_approval") or {}
            linked_fds_app = linked_tds_card.get("fds_approval") or {}
            linked_tds_artifacts = linked_tds_card.get("artifacts") or {}
        except (json.JSONDecodeError, OSError):
            linked_tds_card = {}

    brd_txt = artifacts.get("brd_txt", "")
    fds_approved = (fds_app.get("status") == "approved") or (linked_fds_app.get("status") == "approved")
    fds_approver = fds_app.get("approved_by", "") or linked_fds_app.get("approved_by", "")
    return jsonify({
        "brd_name": Path(brd_txt).name if brd_txt else card_path.stem,
        "req_generated":    _exists(brd_txt),
        "req_approved":     approval.get("status") == "approved",
        "req_approver":     approval.get("approved_by", ""),
        "est_generated":    _exists(est_arts.get("txt", "")),
        "est_approved":     est_app.get("status") == "approved",
        "est_approver":     est_app.get("approved_by", ""),
        "fds_generated":    _any_exists(
            fds_arts.get("txt", ""),
            fds_arts.get("fds_txt", ""),
            fds_arts.get("docx", ""),
            fds_arts.get("fds_docx", ""),
        ),
        "fds_approved":     fds_approved,
        "fds_approver":     fds_approver,
        "tds_generated":    _any_exists(
            tds_arts.get("docx", ""),
            tds_arts.get("tds_docx", ""),
            artifacts.get("tds_docx", ""),
            linked_tds_arts.get("docx", ""),
            linked_tds_arts.get("tds_docx", ""),
            linked_tds_artifacts.get("tds_docx", ""),
        ),
        "tds_approved":     (tds_app.get("status") == "approved") or (linked_tds_app.get("status") == "approved"),
        "tds_approver":     tds_app.get("approved_by", "") or linked_tds_app.get("approved_by", ""),
    })


@app.route("/api/chatbot", methods=["POST"])
def chatbot_api():
    data = request.get_json(silent=True) or {}
    question = (data.get("question") or "").strip()
    project_slug = (data.get("project_slug") or "").strip()
    history = data.get("history") or []

    try:
        answer, meta = _answer_chat_question(question, project_slug, history)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "answer": answer,
        "meta": meta,
    })


@app.route("/api/codex/chat", methods=["POST"])
def codex_chat_api():
    data = request.get_json(silent=True) or {}
    question = (data.get("question") or data.get("message") or "").strip()
    project_slug = (data.get("project_slug") or "").strip()
    history = data.get("history") or []

    try:
        answer, meta = _answer_chat_question(question, project_slug, history)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    meta = dict(meta or {})
    meta["mode"] = "codex"
    meta["retrieval"] = "rag"

    return jsonify({
        "answer": answer,
        "meta": meta,
    })
 
 
@app.route("/api/job/<job_id>", methods=["GET"])
def get_job(job_id):
    job = _get_job(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)
 
 
# ── Generate BRD ──────────────────────────────────────────────────
 
@app.route("/api/generate-brd", methods=["POST"])
def generate_brd():
    project_name = (request.form.get("project_name") or "").strip()
    project_slug = (request.form.get("project_slug") or "").strip()
 
    if project_slug:
        catalog = get_project_catalog(SCRIPT_DIR)
        project_map = {p["slug"]: p for p in catalog}
        selected = project_map.get(project_slug)
        if not selected:
            return jsonify({"error": f"Project slug not found: {project_slug}"}), 400
        # Keep updates inside the selected existing project folder.
        project_name = selected.get("display_name", project_name).strip() or project_name
 
    try:
        transcript_text, src_path = _read_transcript(request.form, request.files)
    except (ValueError, FileNotFoundError) as exc:
        return jsonify({"error": str(exc)}), 400
 
    if not transcript_text.strip():
        return jsonify({"error": "Transcript is empty."}), 400
 
    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")
 
    def worker():
        try:
            result = run_requirements_pipeline(
                transcript_text, src_path, SCRIPT_DIR, project_name
            )
            brd_text = ""
            if result.get("txt_path"):
                p = Path(result["txt_path"])
                if p.exists():
                    brd_text = p.read_text(encoding="utf-8", errors="replace")
 
            _set_job(job_id, "done", result={
                "brd_text":     brd_text,
                "card_path":    result["card_path"],
                "docx_path":    result["docx_path"],
                "txt_path":     result["txt_path"],
                "related_folder": result.get("brd_dir", ""),
                "related_file": result.get("brd_file", ""),
                "project_name": result["overview"].get("PROJECT_NAME", ""),
                "project_slug": Path(result["project_dir"]).name,
                "stats":        result["stats"],
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())
 
    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/save-brd-edit", methods=["POST"])
def save_brd_edit():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    brd_text = (data.get("brd_text") or "").strip()
    txt_path_str = (data.get("txt_path") or "").strip()
    docx_path_str = (data.get("docx_path") or "").strip()

    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400
    if not brd_text:
        return jsonify({"error": "brd_text is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404

    try:
        card = json.loads(card_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        return jsonify({"error": f"Unable to read card: {exc}"}), 500

    artifacts = card.get("artifacts") or {}
    txt_path_raw = txt_path_str or (artifacts.get("brd_txt") or "")
    docx_path_raw = docx_path_str or (artifacts.get("brd_docx") or "")

    if not txt_path_raw:
        return jsonify({"error": "No BRD TXT path found in card artifacts."}), 400

    try:
        txt_path = _safe_path(txt_path_raw)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    txt_path.parent.mkdir(parents=True, exist_ok=True)
    normalized_text = brd_text.replace("\r\n", "\n")
    txt_path.write_text(normalized_text + "\n", encoding="utf-8")
    artifacts["brd_txt"] = str(txt_path)

    saved_docx_path = ""
    if docx_path_raw:
        try:
            docx_path = _safe_path(docx_path_raw)
        except PermissionError as exc:
            return jsonify({"error": str(exc)}), 403

        docx_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            from docx import Document

            doc = Document()
            for line in normalized_text.split("\n"):
                doc.add_paragraph(line)
            doc.save(str(docx_path))
            saved_docx_path = str(docx_path)
        except PermissionError:
            from datetime import datetime as _dt

            timestamp = _dt.now().strftime("%Y%m%d_%H%M%S")
            fallback_docx = docx_path.with_name(f"{docx_path.stem}_edited_{timestamp}{docx_path.suffix}")
            from docx import Document

            doc = Document()
            for line in normalized_text.split("\n"):
                doc.add_paragraph(line)
            doc.save(str(fallback_docx))
            saved_docx_path = str(fallback_docx)
        except Exception as exc:
            return jsonify({"error": f"Failed to write BRD DOCX: {exc}"}), 500

    if saved_docx_path:
        artifacts["brd_docx"] = saved_docx_path

    card["artifacts"] = artifacts
    card["approval"] = build_pending_approval()
    card["status"] = "pending_approval"
    card["updated_at"] = datetime.now().isoformat()

    try:
        card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
    except OSError as exc:
        return jsonify({"error": f"Failed to save card: {exc}"}), 500

    return jsonify({
        "message": "BRD edits saved. Approval reset to pending.",
        "card_path": str(card_path),
        "txt_path": str(txt_path),
        "docx_path": saved_docx_path or (artifacts.get("brd_docx") or ""),
        "approval": card["approval"],
    })


@app.route("/api/approve-card", methods=["POST"])
def approve_card_api():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    approved_by = (data.get("approved_by") or "Web User").strip()
    approver_role = (data.get("approver_role") or "Business Owner").strip()
    comments = (data.get("comments") or "Approved via Web UI.").strip()

    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404

    try:
        approval = _approve_card_file(card_path, approved_by, approver_role, comments)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "message": "Card approved.",
        "approval": approval,
        "card_path": str(card_path),
    })


@app.route("/api/approve-fds", methods=["POST"])
def approve_fds_api():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    approved_by = (data.get("approved_by") or "Web User").strip()
    approver_role = (data.get("approver_role") or "Solution Architect").strip()
    comments = (data.get("comments") or "Approved via Web UI.").strip()

    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404

    try:
        approval = _approve_stage_card_file(
            card_path,
            approval_key="fds_approval",
            approved_by=approved_by,
            approver_role=approver_role,
            comments=comments,
            next_status="approved_for_tds",
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "message": "FDS approved.",
        "fds_approval": approval,
        "card_path": str(card_path),
    })


@app.route("/api/approve-estimation", methods=["POST"])
def approve_estimation_api():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    approved_by = (data.get("approved_by") or "Web User").strip()
    approver_role = (data.get("approver_role") or "Delivery Manager").strip()
    comments = (data.get("comments") or "Approved via Web UI.").strip()

    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404

    try:
        approval = _approve_stage_card_file(
            card_path,
            approval_key="estimation_approval",
            approved_by=approved_by,
            approver_role=approver_role,
            comments=comments,
            next_status="approved_for_fds",
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "message": "Estimation approved.",
        "estimation_approval": approval,
        "card_path": str(card_path),
    })


@app.route("/api/approve-tds", methods=["POST"])
def approve_tds_api():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    approved_by = (data.get("approved_by") or "Web User").strip()
    approver_role = (data.get("approver_role") or "Technical Lead").strip()
    comments = (data.get("comments") or "Approved via Web UI.").strip()
    tds_docx_path = (data.get("tds_docx_path") or "").strip()
    tds_json_path = (data.get("tds_json_path") or "").strip()

    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400

    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403

    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404

    try:
        approval = _approve_stage_card_file(
            card_path,
            approval_key="tds_approval",
            approved_by=approved_by,
            approver_role=approver_role,
            comments=comments,
            next_status="tds_approved",
            artifact_updates={
                "tds_docx": tds_docx_path,
                "tds_json": tds_json_path,
            },
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "message": "TDS approved.",
        "tds_approval": approval,
        "card_path": str(card_path),
    })
 
 
# ── Approve BRD + Run Estimation ─────────────────────────────────
 
@app.route("/api/approve-and-estimate", methods=["POST"])
def approve_and_estimate():
    data = request.get_json(silent=True) or {}
    card_path_str = (data.get("card_path") or "").strip()
    approved_by   = (data.get("approved_by")   or "Web User").strip()
    approver_role = (data.get("approver_role") or "Business Owner").strip()
    comments      = (data.get("comments")      or "Approved via Web UI.").strip()
 
    if not card_path_str:
        return jsonify({"error": "card_path is required"}), 400
 
    try:
        card_path = _safe_path(card_path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403
 
    if not card_path.exists():
        return jsonify({"error": f"Card not found: {card_path_str}"}), 404
 
    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")
 
    def worker():
        try:
            _approve_card_file(card_path, approved_by, approver_role, comments)

            from estimation_agents import run_estimation_pipeline
            est = run_estimation_pipeline(str(card_path))
 
            _set_job(job_id, "done", result={
                "estimation_text":      est.get("txt_content", ""),
                "estimation_txt_path":  est.get("txt_path", ""),
                "estimation_docx_path": est.get("docx_path", ""),
                "size":       est.get("size", ""),
                "investment": est.get("investment", ""),
                "total_low":  est.get("total_low", 0),
                "total_mid":  est.get("total_mid", 0),
                "total_high": est.get("total_high", 0),
                "total_with_buffer": est.get("total_with_buffer", 0),
                "breakdown":  est.get("breakdown"),
                "req_count":  est.get("req_count", 0),
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


# ── Run Estimation only (card already approved) ───────────────────
 
@app.route("/api/run-estimation", methods=["POST"])
def run_estimation():
    card_path = None
    brd_txt_path_str = ""

    uploaded_card = request.files.get("card_file")
    if uploaded_card and uploaded_card.filename:
        try:
            card_path = _store_uploaded_estimation_card(uploaded_card)
        except json.JSONDecodeError:
            return jsonify({"error": "Uploaded card file is not valid JSON."}), 400
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except OSError as exc:
            return jsonify({"error": f"Failed to store uploaded card: {exc}"}), 500

    if request.is_json:
        data = request.get_json(silent=True) or {}
        if (data.get("card_path") or "").strip():
            return jsonify({"error": "card_path is no longer supported. Use brd_txt_path or upload a card file."}), 400
        brd_txt_path_str = (data.get("brd_txt_path") or "").strip()
    else:
        brd_txt_path_str = (request.form.get("brd_txt_path") or "").strip()

    if not card_path and brd_txt_path_str:
        try:
            brd_txt_path = _safe_path(brd_txt_path_str)
        except PermissionError as exc:
            return jsonify({"error": str(exc)}), 403

        if not brd_txt_path.exists():
            return jsonify({"error": f"BRD file not found: {brd_txt_path_str}"}), 404

        # Auto-create estimation card if it doesn't exist
        try:
            card_path = _auto_create_estimation_card(brd_txt_path)
        except Exception as exc:
            return jsonify({"error": f"Failed to auto-create estimation card: {exc}"}), 500

    if not card_path:
        return jsonify({"error": "Provide brd_txt_path or upload a card_file."}), 400
 
    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")
 
    def worker():
        try:
            from estimation_agents import run_estimation_pipeline
            est = run_estimation_pipeline(str(card_path))
            _set_job(job_id, "done", result={
                "card_path": str(card_path),
                "estimation_text":      est.get("txt_content", ""),
                "estimation_txt_path":  est.get("txt_path", ""),
                "estimation_docx_path": est.get("docx_path", ""),
                "size":       est.get("size", ""),
                "investment": est.get("investment", ""),
                "total_low":  est.get("total_low", 0),
                "total_mid":  est.get("total_mid", 0),
                "total_high": est.get("total_high", 0),
                "total_with_buffer": est.get("total_with_buffer", 0),
                "breakdown":  est.get("breakdown"),
                "req_count":  est.get("req_count", 0),
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())
 
    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


# ── Run FDS from approved Requirement card ───────────────────────

@app.route("/api/run-fds", methods=["POST"])
def run_fds():
    card_path = None
    brd_txt_path_str = ""

    uploaded_card = request.files.get("card_file")
    if uploaded_card and uploaded_card.filename:
        try:
            card_path = _store_uploaded_stage_card(uploaded_card, "fds_cards")
        except json.JSONDecodeError:
            return jsonify({"error": "Uploaded card file is not valid JSON."}), 400
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except OSError as exc:
            return jsonify({"error": f"Failed to store uploaded card: {exc}"}), 500

    if request.is_json:
        data = request.get_json(silent=True) or {}
        if (data.get("card_path") or "").strip():
            return jsonify({"error": "card_path is no longer supported. Use brd_txt_path or upload a card file."}), 400
        brd_txt_path_str = (data.get("brd_txt_path") or "").strip()
    else:
        brd_txt_path_str = (request.form.get("brd_txt_path") or "").strip()

    if not card_path and brd_txt_path_str:
        try:
            brd_txt_path = _safe_path(brd_txt_path_str)
        except PermissionError as exc:
            return jsonify({"error": str(exc)}), 403

        if not brd_txt_path.exists():
            return jsonify({"error": f"BRD file not found: {brd_txt_path_str}"}), 404

        # Auto-create estimation card if it doesn't exist
        try:
            estimation_card = _auto_create_estimation_card(brd_txt_path)
            # Auto-create FDS→TDS card before running FDS
            _auto_create_fds_tds_card(brd_txt_path, estimation_card)
            card_path = estimation_card
        except Exception as exc:
            return jsonify({"error": f"Failed to auto-create cards: {exc}"}), 500

    if not card_path:
        return jsonify({"error": "Provide brd_txt_path or upload a card_file."}), 400

    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")

    def worker():
        try:
            from fds import run_fds_pipeline

            result = run_fds_pipeline(str(card_path))
            sections = result.get("sections") or {}

            _set_job(job_id, "done", result={
                "card_path": str(card_path),
                "project_name": result.get("project_name", ""),
                "fds_docx_path": result.get("docx_path", ""),
                "fds_txt_path": result.get("txt_path", ""),
                "related_folder": result.get("fds_dir", ""),
                "related_file": result.get("fds_file", ""),
                "tds_card_path": result.get("tds_card_path", ""),
                "fds_text": result.get("txt_content", ""),
                "sections_count": len(sections),
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


# ── Run TDS from ready_for_tds card ──────────────────────────────

@app.route("/api/run-tds", methods=["POST"])
def run_tds():
    card_path = None
    fds_txt_path_str = ""

    uploaded_card = request.files.get("card_file")
    if uploaded_card and uploaded_card.filename:
        try:
            card_path = _store_uploaded_stage_card(uploaded_card, "tds_cards")
        except json.JSONDecodeError:
            return jsonify({"error": "Uploaded card file is not valid JSON."}), 400
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except OSError as exc:
            return jsonify({"error": f"Failed to store uploaded card: {exc}"}), 500

    if request.is_json:
        data = request.get_json(silent=True) or {}
        if (data.get("card_path") or "").strip():
            return jsonify({"error": "card_path is no longer supported. Use fds_txt_path or upload a card file."}), 400
        fds_txt_path_str = (data.get("fds_txt_path") or "").strip()
    else:
        fds_txt_path_str = (request.form.get("fds_txt_path") or "").strip()

    if not card_path and fds_txt_path_str:
        try:
            fds_txt_path = _safe_path(fds_txt_path_str)
        except PermissionError as exc:
            return jsonify({"error": str(exc)}), 403

        if not fds_txt_path.exists():
            return jsonify({"error": f"FDS file not found: {fds_txt_path_str}"}), 404

        resolved_card = _find_tds_card_for_fds_txt(fds_txt_path)
        if not resolved_card:
            return jsonify({"error": "No TDS card is linked with the selected FDS. Generate FDS first."}), 400
        card_path = resolved_card

    if not card_path:
        return jsonify({"error": "Provide fds_txt_path or upload a card_file."}), 400

    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")

    def worker():
        try:
            from tds import (
                build_tds_docx,
                load_tds_a2a_card,
                resolve_output_dir,
                run_tds_from_card,
            )

            card = load_tds_a2a_card(card_path)
            result = run_tds_from_card(card)
            output_dir = resolve_output_dir(card_path, card)
            output_dir.mkdir(parents=True, exist_ok=True)

            timestamp = time.strftime("%Y%m%d_%H%M%S")
            docx_path = output_dir / f"TDS_{timestamp}.docx"
            json_path = output_dir / f"TDS_{timestamp}.json"
            source_fds_docx = card.get("artifacts", {}).get("fds_docx", "Not specified")
            display_folder = result.get("related_folder") or str(output_dir)
            display_file = result.get("related_file") or ""

            doc = build_tds_docx(
                result["project_name"],
                result.get("requirements", []),
                result["sections"],
                result.get("section_order", list(result["sections"].keys())),
                card_path.name,
                source_fds_docx,
                related_folder=display_folder,
                related_file=display_file,
            )
            doc.save(str(docx_path))

            json_payload = {
                "project": result["project_name"],
                "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
                "source_card": str(card_path),
                "related_file": display_file,
                "related_folder": display_folder,
                "requirements": result.get("requirements", []),
                "section_plan": result.get("section_plan", []),
                "sections": result["sections"],
            }
            json_path.write_text(json.dumps(json_payload, indent=2), encoding="utf-8")

            # Stamp generated TDS artifact paths to the active TDS card (approval remains separate).
            try:
                card["tds_artifacts"] = {
                    "tds_docx": str(docx_path),
                    "tds_json": str(json_path),
                }
                card_artifacts = card.get("artifacts") or {}
                card_artifacts["tds_docx"] = str(docx_path)
                card_artifacts["tds_json"] = str(json_path)
                card["artifacts"] = card_artifacts
                card["status"] = "tds_generated"
                card["updated_at"] = datetime.now().isoformat()
                card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
            except OSError:
                pass

            ordered_sections = result.get("section_order", list(result["sections"].keys()))
            preview_parts = []
            for section_name in ordered_sections:
                preview_parts.append(section_name)
                preview_parts.append("-" * len(section_name))
                preview_parts.append(result["sections"].get(section_name, ""))
                preview_parts.append("")

            _set_job(job_id, "done", result={
                "card_path": str(card_path),
                "project_name": result.get("project_name", ""),
                "tds_docx_path": str(docx_path),
                "tds_json_path": str(json_path),
                "related_file": display_file,
                "related_folder": display_folder,
                "tds_preview": "\n".join(preview_parts).strip(),
                "sections_count": len(result.get("sections", {})),
                "section_order": ordered_sections,
                "sections": result.get("sections", {}),
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


# ── Run Test Cases (pseudocode + test cases from TDS card) ────────

@app.route("/api/run-test-cases", methods=["POST"])
def run_test_cases():
    card_path = None
    tds_json_path_str = ""

    uploaded_card = request.files.get("card_file")
    if uploaded_card and uploaded_card.filename:
        try:
            card_path = _store_uploaded_stage_card(uploaded_card, "tds_cards")
        except json.JSONDecodeError:
            return jsonify({"error": "Uploaded card file is not valid JSON."}), 400
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        except OSError as exc:
            return jsonify({"error": f"Failed to store uploaded card: {exc}"}), 500

    if request.is_json:
        data = request.get_json(silent=True) or {}
        tds_json_path_str = (data.get("tds_json_path") or "").strip()
    else:
        tds_json_path_str = (request.form.get("tds_json_path") or "").strip()

    if not card_path and tds_json_path_str:
        try:
            tds_json_path = _safe_path(tds_json_path_str)
        except PermissionError as exc:
            return jsonify({"error": str(exc)}), 403

        if not tds_json_path.exists():
            return jsonify({"error": f"TDS file not found: {tds_json_path_str}"}), 404

        resolved_card = _find_tds_card_for_tds_json(tds_json_path)
        if not resolved_card:
            return jsonify({"error": "No TDS card linked to selected TDS file. Generate TDS first."}), 400
        card_path = resolved_card

    if not card_path:
        return jsonify({"error": "Provide tds_json_path or upload a card_file."}), 400

    job_id = str(uuid.uuid4())
    _set_job(job_id, "running")

    def worker():
        try:
            from tds import load_tds_a2a_card, resolve_output_dir
            from test_case_generator import build_test_cases_docx, run_test_cases_from_card

            card = load_tds_a2a_card(card_path)
            result = run_test_cases_from_card(card)

            output_dir = resolve_output_dir(card_path, card)
            output_dir.mkdir(parents=True, exist_ok=True)

            timestamp = time.strftime("%Y%m%d_%H%M%S")
            docx_path = output_dir / f"TestCases_{timestamp}.docx"
            json_path = output_dir / f"TestCases_{timestamp}.json"

            doc = build_test_cases_docx(
                result["project_name"],
                result["sections"],
                result.get("section_order", list(result["sections"].keys())),
                card_path.name,
            )
            doc.save(str(docx_path))

            json_payload = {
                "project": result["project_name"],
                "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
                "source_card": str(card_path),
                "sections": result["sections"],
            }
            json_path.write_text(json.dumps(json_payload, indent=2), encoding="utf-8")

            # Build plain-text preview
            ordered = result.get("section_order", list(result["sections"].keys()))
            preview_parts = []
            for name in ordered:
                preview_parts.append(f"{'='*60}\n{name}\n{'='*60}")
                preview_parts.append(result["sections"].get(name, ""))
                preview_parts.append("")

            _set_job(job_id, "done", result={
                "card_path": str(card_path),
                "project_name": result.get("project_name", ""),
                "tc_docx_path": str(docx_path),
                "tc_json_path": str(json_path),
                "tc_preview": "\n".join(preview_parts).strip(),
                "sections": result.get("sections", {}),
                "section_order": ordered,
            })
        except Exception as exc:
            _set_job(job_id, "error", error=str(exc) + "\n" + traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id})


# ── Download file ─────────────────────────────────────────────────
 
@app.route("/api/download", methods=["GET"])
def download_file():
    path_str = (request.args.get("path") or "").strip()
    if not path_str:
        return jsonify({"error": "path is required"}), 400
    try:
        p = _safe_path(path_str)
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 403
    if not p.exists():
        return jsonify({"error": "File not found"}), 404
    return send_file(str(p), as_attachment=True, download_name=p.name)
 
 
# ── Send Mail (stub — configure SMTP to enable) ───────────────────
 
@app.route("/api/send-mail", methods=["POST"])
def send_mail():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip()
    subject = (data.get("subject") or "BRD Report").strip()
    attachment_path = (data.get("attachment_path") or "").strip()
    body = _generic_mail_body(subject, attachment_path)
 
    if not email:
        return jsonify({"error": "email is required"}), 400
 
    try:
        _send_email_smtp(email, subject, body, attachment_path)
    except Exception as exc:
        return jsonify({"error": f"Mail send failed: {exc}"}), 500
 
    return jsonify({"message": f"Email sent to {email}."})
 
 
# ══════════════════════════════════════════════════════════════════
 
if __name__ == "__main__":
    host = os.getenv("APP_HOST", "0.0.0.0").strip() or "0.0.0.0"
    try:
        ports = _parse_ports(os.getenv("APP_PORTS", "5000"), default=5000)
    except ValueError as exc:
        print(f"[app] {exc}")
        sys.exit(1)

    print("[app] Starting Requirement Agent web server")
    if len(ports) == 1:
        print(f"[app] Serving on http://localhost:{ports[0]}")
        app.run(debug=True, host=host, port=ports[0], use_reloader=False)
    else:
        print(f"[app] Multi-port mode enabled: {', '.join(str(p) for p in ports)}")
        print("[app] Note: Debug reloader is disabled in multi-port mode.")
        _run_multi_port(app, host, ports)