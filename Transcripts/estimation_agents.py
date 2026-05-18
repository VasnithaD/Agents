"""
Estimation.py
─────────────
A2A Estimation Agent
Reads a2a_card_*.json produced by RequirementsAgent,
scores each requirement on 5 complexity dimensions using the LLM,
maps total hours to the sizing scale, and writes Estimation.docx + .txt
"""
 
# ── standard library ───────────────────────────────────────────────
import os
import re
import sys
import json
import uuid
import argparse
import time
import glob
from pathlib import Path
from datetime import datetime
 
# ── third-party ────────────────────────────────────────────────────
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
# ══════════════════════════════════════════════════════════════════
# 1. ENV + CONFIG
# ══════════════════════════════════════════════════════════════════
HERE = Path(__file__).parent
load_dotenv(HERE / ".env")
 
CLIENT_ID  = os.getenv("CLIENT_ID")
AUTH_TOKEN = os.getenv("AUTHENTICATION_TOKEN")
MODEL      = os.getenv("MODEL_NAME", "gpt-4o-mini")
API_URL    = "https://api.chathpe.it.hpe.com/v2.8/"
SCRIPT_DIR = Path(__file__).parent.resolve()
ENV_PATH   = SCRIPT_DIR / ".env"
PEM_PATH   = SCRIPT_DIR / "cacert 1 (1).pem"
SESSION = {
    "user_id": None,
    "username": None,
    "session_id": None,
}
 
 
def project_output_paths(card_path: Path, card: dict):
    """Resolve project-scoped estimation output directory and artifact paths."""
    artifacts = card.get("artifacts", {})
    project_root = artifacts.get("output_dir", "")
 
    if project_root:
        project_dir = Path(project_root)
    else:
        # Fallback for older cards: infer from .../projects/<slug>/cards/a2a_card_*.json
        card_parent = card_path.parent
        project_dir = card_parent.parent if card_parent.name == "cards" else HERE
 
    estimation_dir = project_dir / "estimation"
    estimation_dir.mkdir(parents=True, exist_ok=True)
 
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = estimation_dir / f"Estimation_{timestamp}.docx"
    txt_path = estimation_dir / f"Estimation_{timestamp}.txt"
    return estimation_dir, docx_path, txt_path
 
 
def clean(value: str) -> str:
    return value.strip().strip('"').strip("'") if value else ""
 
 
def auth_header() -> str:
    token = clean(AUTH_TOKEN)
    if not token:
        raise RuntimeError("AUTHENTICATION_TOKEN missing from .env")
    if token.lower().startswith("bearer "):
        return token
    return "Bearer " + token
 
 
def verify_tls_value():
    return str(PEM_PATH) if PEM_PATH.exists() else True
 
 
def api_post(endpoint: str, payload: dict, timeout_sec: int = 60) -> dict:
    url = API_URL + endpoint
    headers = {
        "Client-ID": clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Content-Type": "application/json",
        "Cache-Control": "no-cache",
    }
 
    last_exc = None
    max_attempts = 6
    for attempt in range(max_attempts):
        try:
            resp = requests.post(
                url,
                headers=headers,
                json=payload,
                verify=verify_tls_value(),
                timeout=timeout_sec,
            )
            resp.raise_for_status()
            return resp.json()
        except requests.RequestException as exc:
            last_exc = exc
            code = None
            if getattr(exc, "response", None) is not None:
                code = exc.response.status_code
            should_retry = (code in (429, 500, 502, 503, 504)) or (code is None)
            if attempt < (max_attempts - 1) and should_retry:
                # Longer backoff for rate limiting
                if code == 429:
                    wait_time = min(30, (3 ** attempt))  # 1, 3, 9, 27, 30, 30 seconds
                    print(f"  [rate-limit] 429 — backing off {wait_time}s (attempt {attempt + 1}/{max_attempts})")
                else:
                    wait_time = 2 ** attempt  # 1, 2, 4, 8, 16, 32 seconds
                    print(f"  [retry] HTTP {code} — backing off {wait_time}s (attempt {attempt + 1}/{max_attempts})")
                time.sleep(wait_time)
                continue
            raise
 
    if last_exc:
        raise last_exc
    raise RuntimeError("Unknown API error")
 
 
def api_get(endpoint: str, timeout_sec: int = 30) -> str:
    url = API_URL + endpoint
    headers = {
        "Client-ID": clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Cache-Control": "no-cache",
    }
    resp = requests.get(
        url,
        headers=headers,
        verify=verify_tls_value(),
        timeout=timeout_sec,
    )
    resp.raise_for_status()
    return resp.text
 
 
def parse_session_id(raw_session_text: str) -> str:
    # Typical response contains a UUID-like session id; keep the whole id.
    uuid_match = re.search(
        r"\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b",
        raw_session_text,
    )
    if uuid_match:
        return uuid_match.group(0)
 
    # Fallback for older numeric formats.
    num_match = re.search(r"\b\d+\b", raw_session_text)
    if num_match:
        return num_match.group(0)
 
    return "-1"
 
 
def setup_session() -> None:
    if SESSION["user_id"] and SESSION["username"] and SESSION["session_id"]:
        return
 
    login_resp = api_post("login", {"appId": "1"}, timeout_sec=60)
    bot_data = login_resp.get("chatHPE_bot_data", {})
 
    user_id = str(bot_data.get("userId", "")).strip()
    username = str(bot_data.get("username", "")).strip()
    if not user_id or not username:
        raise RuntimeError("Could not establish chat session: missing user data")
 
    api_post(
        "preferences",
        {
            "agreement": True,
            "dark_mode": True,
            "stream": False,
            "chatHPE_bot_data": {
                "appId": "1",
                "sessionId": "-1",
                "userId": user_id,
                "username": username,
            },
            "webScraping": False,
        },
        timeout_sec=60,
    )
 
    raw_session = api_get("sessionId_generator", timeout_sec=60)
    session_id = parse_session_id(raw_session)
 
    SESSION["user_id"] = user_id
    SESSION["username"] = username
    SESSION["session_id"] = session_id
 
 
def call_llm(system_prompt: str, user_prompt: str) -> str:
    setup_session()
 
    endpoint = (
        "call/chatlite"
        "?force_async=false"
        "&session_management_support=true"
        "&internal_call=false"
        "&proxy=false"
    )
 
    merged_prompt = (
        "SYSTEM INSTRUCTION:\n"
        + system_prompt
        + "\n\nUSER REQUEST:\n"
        + user_prompt
    )
 
    payload = {
        "chatHPE_bot_data": {
            "appId": "1",
            "sessionId": SESSION["session_id"],
            "userId": SESSION["user_id"],
            "username": SESSION["username"],
        },
        "model_name": MODEL,
        "stream": False,
        "webScraping": False,
        "user_query": merged_prompt,
    }
 
    try:
        resp = api_post(endpoint, payload, timeout_sec=90)
    except requests.HTTPError as exc:
        status = exc.response.status_code if exc.response is not None else None
        if status == 422 and MODEL != "gpt-4o-mini":
            payload["model_name"] = "gpt-4o-mini"
            resp = api_post(endpoint, payload, timeout_sec=90)
        else:
            raise
 
    text = resp.get("Response", "")
    if not text:
        text = json.dumps(resp)
    return text.strip()
 
 
# ══════════════════════════════════════════════════════════════════
# 2. SCORING CONSTANTS
# ══════════════════════════════════════════════════════════════════
 
DIMENSIONS = [
    "integration",
    "data_volume",
    "business_logic",
    "ui_complexity",
    "nfr_burden",
]
 
SCORE_MAP = {
    "integration":    {"none": 0,     "1-2 apis": 2,    "3+ systems": 4},
    "data_volume":    {"low": 0,      "medium": 2,      "high/rf": 4},
    "business_logic": {"simple": 0,   "moderate": 2,    "complex": 4},
    "ui_complexity":  {"forms": 0,    "dashboards": 2,  "custom ui": 4},
    "nfr_burden":     {"standard": 0, "performance": 2, "compliance": 4},
}
 
DIM_LABELS = {
    "integration":    "Integration",
    "data_volume":    "Data Volume",
    "business_logic": "Business Logic",
    "ui_complexity":  "UI Complexity",
    "nfr_burden":     "NFR Burden",
}
 
SIZING_SCALE = [
    {"size": "XXS",  "max_hrs": 100,    "investment": "$15,000",
     "scope": "Simple feature, 1-2 screens, no integration"},
    {"size": "XS",   "max_hrs": 500,    "investment": "$35,000",
     "scope": "Small module, 1 integration, basic CRUD"},
    {"size": "S",    "max_hrs": 1500,   "investment": "$100,000",
     "scope": "Multi-feature module, 2-3 integrations"},
    {"size": "M",    "max_hrs": 2500,   "investment": "$170,000",
     "scope": "Full sub-system, complex workflows"},
    {"size": "L",    "max_hrs": 3500,   "investment": "$230,000",
     "scope": "Large platform, multi-team"},
    {"size": "XL",   "max_hrs": 5000,   "investment": "$350,000",
     "scope": "Full product stream, multiple integrations"},
    {"size": "XXL",  "max_hrs": 10000,  "investment": "$650,000",
     "scope": "Programme-scale, org-wide impact"},
    {"size": "XXXL", "max_hrs": 999999, "investment": "Custom SOW",
     "scope": "Enterprise transformation, needs decomposition"},
]
 
SIZE_COLOURS = {
    "XXS":  RGBColor(0xD9, 0xEA, 0xD3),
    "XS":   RGBColor(0xD9, 0xEA, 0xD3),
    "S":    RGBColor(0xFF, 0xF2, 0xCC),
    "M":    RGBColor(0xFF, 0xF2, 0xCC),
    "L":    RGBColor(0xFF, 0xE6, 0xCC),
    "XL":   RGBColor(0xFF, 0xE6, 0xCC),
    "XXL":  RGBColor(0xF4, 0xCC, 0xCC),
    "XXXL": RGBColor(0xEA, 0x9A, 0x9A),
}
 
SYSTEM_SCORE = """You are a senior software estimation expert.
Given a software requirement, score it on exactly 5 dimensions.
Respond ONLY with valid JSON and no extra text, no markdown fences:
{
  "integration":    "<None|1-2 APIs|3+ systems>",
  "data_volume":    "<Low|Medium|High/RF>",
  "business_logic": "<Simple|Moderate|Complex>",
  "ui_complexity":  "<Forms|Dashboards|Custom UI>",
  "nfr_burden":     "<Standard|Performance|Compliance>",
  "reasoning":      "<one sentence explaining your scores>"
}"""
 
def score_requirement(req):
 
    req_id = req.get("id", "?")
    req_cat = req.get("category", "")
    req_pri = req.get("priority", "")
    req_desc = req.get("description", "")
 
    line1 = "Requirement ID   : " + req_id
    line2 = "Category         : " + req_cat
    line3 = "Priority         : " + req_pri
    line4 = "Description      : " + req_desc
    line5 = "Score this requirement on the 5 complexity dimensions."
    line6 = "Return ONLY valid JSON, no markdown."
 
    user_prompt = line1 + "\n" + line2 + "\n" + line3 + "\n" + line4 + "\n\n" + line5 + "\n" + line6
 
    raw = call_llm(SYSTEM_SCORE, user_prompt)
 
    scored = None
    fence = chr(96) + chr(96) + chr(96)
 
    try:
        clean_text = raw
        clean_text = clean_text.replace(fence + "json", "")
        clean_text = clean_text.replace(fence, "")
        clean_text = clean_text.strip()
        scored = json.loads(clean_text)
    except json.JSONDecodeError:
        print("  [warn] Could not parse score for " + req_id)
        scored = dict()
        scored["integration"] = "None"
        scored["data_volume"] = "Low"
        scored["business_logic"] = "Simple"
        scored["ui_complexity"] = "Forms"
        scored["nfr_burden"] = "Standard"
        scored["reasoning"] = "Could not parse LLM response"
 
    total = 0
    for dim in DIMENSIONS:
        val = scored.get(dim, "")
        val = val.lower().strip()
        dim_score = SCORE_MAP[dim].get(val, 0)
        total = total + dim_score
 
    hrs_low = 0
    hrs_high = 0
 
    if total <= 4:
        hrs_low = 10
        hrs_high = 50
    elif total <= 8:
        hrs_low = 50
        hrs_high = 150
    elif total <= 12:
        hrs_low = 150
        hrs_high = 300
    elif total <= 16:
        hrs_low = 300
        hrs_high = 500
    else:
        hrs_low = 500
        hrs_high = 800
 
    hrs_mid = (hrs_low + hrs_high) // 2
 
    result = dict()
    result["id"] = req_id
    result["description"] = req_desc
    result["category"] = req_cat
    result["priority"] = req_pri
    result["integration"] = scored.get("integration", "None")
    result["data_volume"] = scored.get("data_volume", "Low")
    result["business_logic"] = scored.get("business_logic", "Simple")
    result["ui_complexity"] = scored.get("ui_complexity", "Forms")
    result["nfr_burden"] = scored.get("nfr_burden", "Standard")
    result["reasoning"] = scored.get("reasoning", "")
    result["complexity_score"] = total
    result["hrs_low"] = hrs_low
    result["hrs_high"] = hrs_high
    result["hrs_mid"] = hrs_mid
 
    return result
 
 
def map_to_size(total_hrs):
    for band in SIZING_SCALE:
        if total_hrs <= band["max_hrs"]:
            return band
    return SIZING_SCALE[-1]
 
 
def breakdown_hours_by_role(total_mid_hours):
    """
    Break down total hours into roles/categories.
    Distribution: Frontend 40%, Backend 35%, Testing 15%, Documentation 10%
    Add 8 hours buffer to each role.
    Returns updated breakdown dict and new total.
    """
    frontend = int(total_mid_hours * 0.40) + 8
    backend = int(total_mid_hours * 0.35) + 8
    testing = int(total_mid_hours * 0.15) + 8
    documentation = int(total_mid_hours * 0.10) + 8
   
    new_total = frontend + backend + testing + documentation
   
    breakdown = {
        "frontend": frontend,
        "backend": backend,
        "testing": testing,
        "documentation": documentation,
        "total_original": total_mid_hours,
        "total_with_buffer": new_total,
        "buffer_added": 32,  # 8 hours per role * 4 roles
    }
   
    return breakdown, new_total
 
 
def set_cell_shading(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)
 
 
def style_header_row(row, hex_color="1F4E79"):
    for cell in row.cells:
        set_cell_shading(cell, hex_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
 
 
def add_styled_row(table, values, bold_first=False):
    row = table.add_row()
    for idx in range(len(values)):
        cell = row.cells[idx]
        cell.text = str(values[idx])
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if bold_first and idx == 0:
                    run.font.bold = True
    return row
 
def build_docx(scored_list, size_info, totals, project_name, out_path, breakdown=None):
    """Build simplified docx with size, investment summary, and hours breakdown by role."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
 
    # ── Title ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Project Estimation Report")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
 
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(project_name)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
 
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("Generated: " + datetime.now().strftime("%Y-%m-%d %H:%M"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
 
    doc.add_paragraph()
 
    # ── Summary ────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
 
    size_label = size_info.get("size", "?")
    investment = size_info.get("investment", "?")
 
    summary_table = doc.add_table(rows=2, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Light Grid Accent 1"
 
    s_hdr = summary_table.rows[0].cells
    s_hdr[0].text = "Project Size"
    s_hdr[1].text = "Investment"
    style_header_row(summary_table.rows[0])
 
    s_val = summary_table.rows[1].cells
    s_val[0].text = size_label
    s_val[1].text = investment
 
    # ── Hours Breakdown by Role ────────────────────────────────
    if breakdown:
        doc.add_paragraph()
        doc.add_heading("Effort Breakdown", level=1)
       
        doc.add_paragraph(f"Frontend Development    : {breakdown.get('frontend', 0)} hours")
        doc.add_paragraph(f"Backend Development     : {breakdown.get('backend', 0)} hours")
        doc.add_paragraph(f"Testing & QA            : {breakdown.get('testing', 0)} hours")
        doc.add_paragraph(f"Documentation           : {breakdown.get('documentation', 0)} hours")
        doc.add_paragraph()
        doc.add_paragraph(f"TOTAL EFFORT            : {breakdown.get('total_with_buffer', 0)} hours")
       
    # ── Save ───────────────────────────────────────────────────
    doc.save(str(out_path))
    print("  [docx] Saved: " + str(out_path))
 
 
def build_txt(scored_list, size_info, totals, project_name, out_path, breakdown=None):
    """Build simplified txt with size, investment summary, and hours breakdown by role."""
    lines = []
    lines.append("ESTIMATION REPORT")
    lines.append("=" * 60)
    lines.append("")
    lines.append("Project: " + project_name)
    lines.append("Date: " + datetime.now().strftime("%Y-%m-%d %H:%M"))
    lines.append("")
    lines.append("EXECUTIVE SUMMARY")
    lines.append("-" * 60)
    lines.append("PROJECT SIZE     : " + size_info.get("size", "?"))
    lines.append("INVESTMENT RANGE : " + size_info.get("investment", "?"))
    lines.append("")
    lines.append("SCOPE PROFILE")
    lines.append("-" * 60)
    lines.append(size_info.get("scope", "?"))
    lines.append("")
   
    # ── Hours Breakdown by Role ────────────────────────────────
    if breakdown:
        lines.append("EFFORT BREAKDOWN")
        lines.append("-" * 60)
        lines.append(f"Frontend Development    : {breakdown.get('frontend', 0):4d} hours")
        lines.append(f"Backend Development     : {breakdown.get('backend', 0):4d} hours")
        lines.append(f"Testing & QA            : {breakdown.get('testing', 0):4d} hours")
        lines.append(f"Documentation           : {breakdown.get('documentation', 0):4d} hours")
        lines.append("-" * 60)
        lines.append(f"TOTAL EFFORT            : {breakdown.get('total_with_buffer', 0):4d} hours")
        lines.append(f"  (Original: {breakdown.get('total_original', 0)} hours)")
        lines.append("")
 
    content = "\n".join(lines)
    out_path.write_text(content, encoding="utf-8")
    print("  [txt] Saved: " + str(out_path))
    return content
 
 
def _stamp_estimation_approval(card_path: Path, docx_path: Path, txt_path: Path,
                               approved_by: str, approver_role: str) -> None:
    """Mark the card as estimation-approved and record the artifact paths."""
    try:
        card = json.loads(card_path.read_text(encoding="utf-8"))
        card["estimation_approval"] = {
            "status": "approved",
            "approved_by": approved_by,
            "approver_role": approver_role,
            "approved_at": datetime.now().isoformat(),
            "comments": "Estimation approved in terminal.",
        }
        card["estimation_artifacts"] = {
            "docx": str(docx_path),
            "txt": str(txt_path),
        }
        card["status"] = "estimation_approved"
        card["updated_at"] = datetime.now().isoformat()
        card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
        print("[EstimationAgent] Estimation approved and card updated.")
    except Exception as exc:
        print(f"[EstimationAgent] Could not update card: {exc}")
 
 
def _get_updated_transcript_path(card_path: Path) -> Path:
    """Ask reviewer whether to supply updated transcript via text or file."""
    source_transcript = Path(
        json.loads(card_path.read_text(encoding="utf-8"))
        .get("project", {})
        .get("source_transcript", "")
    )
    confirm = input(
        "[EstimationAgent] Do you want to provide an updated transcript before re-running? (yes/no): "
    ).strip().lower()
 
    if confirm not in ("yes", "y"):
        return source_transcript
 
    mode = input(
        '[EstimationAgent] Type "text" to paste transcript, or "file" to provide a .txt path: '
    ).strip().lower()
 
    if mode == "text":
        print('[EstimationAgent] Paste updated transcript. Type "END" on a new line to finish.')
        lines = []
        while True:
            line = input()
            if line.strip() == "END":
                break
            lines.append(line)
        updated_content = "\n".join(lines).strip()
        if not updated_content:
            print("[EstimationAgent] No text provided. Using existing transcript.")
            return source_transcript
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        updated_path = source_transcript.parent / f"{source_transcript.stem}_updated_{timestamp}.txt"
        updated_path.write_text(updated_content, encoding="utf-8")
        print(f"[EstimationAgent] Updated transcript saved → {updated_path.name}")
        return updated_path
 
    if mode == "file":
        raw = input("[EstimationAgent] Enter updated .txt file path: ").strip().strip('"')
        file_path = Path(raw)
        if file_path.exists() and file_path.suffix.lower() == ".txt":
            return file_path
        print("[EstimationAgent] Invalid path. Using existing transcript.")
        return source_transcript
 
    print("[EstimationAgent] Unknown option. Using existing transcript.")
    return source_transcript
 
 
def _retrigger_requirements(card_path: Path) -> None:
    """Re-run Requirements agent with an optionally updated transcript."""
    transcript_path = _get_updated_transcript_path(card_path)
    requirements_script = HERE / "Requirements.py"
    if not requirements_script.exists():
        print(f"[EstimationAgent] Requirements.py not found at {requirements_script}")
        return
    print("[EstimationAgent] Re-triggering Requirements Agent...")
    import subprocess
    try:
        subprocess.run(
            [sys.executable, str(requirements_script), str(transcript_path)],
            check=True,
        )
    except subprocess.CalledProcessError as exc:
        print(f"[EstimationAgent] Requirements Agent failed: {exc}")
 
 
def run_estimation_pipeline(card_path_str: str) -> dict:
    """
    Run the estimation pipeline programmatically (called from app.py / web UI).
    The card must already be approved_for_estimation.
    Returns a dict with txt_content, txt_path, docx_path, size, investment, totals.
    """
    card_path = Path(card_path_str)
    if not card_path.exists():
        raise FileNotFoundError(f"Card not found: {card_path}")
 
    card = json.loads(card_path.read_text(encoding="utf-8"))
 
    approval = card.get("approval", {})
    card_status = card.get("status", "pending_approval")
    approval_status = approval.get("status", "pending")
    allowed_statuses = {"approved_for_estimation", "estimation_approved"}
    if card_status not in allowed_statuses or approval_status != "approved":
        raise RuntimeError(
            f"Card is not approved for estimation. "
            f"status={card_status}, approval.status={approval_status}"
        )
 
    project_info = card.get("project", {})
    project_name = project_info.get("name", "Unknown Project")
 
    payload = card.get("payload", {})
    requirements = payload.get("requirements", [])
 
    # Normalize requirement shape
    cleaned = []
    for req in requirements:
        cleaned.append({
            "id":          req.get("id", req.get("req_id", "REQ-?")),
            "category":    req.get("category", req.get("type", "Functional")),
            "priority":    req.get("priority", "Medium"),
            "description": req.get("description", req.get("title", "")),
        })
    requirements = cleaned
 
    if not requirements:
        raise ValueError("No requirements found in card.")
 
    setup_session()
 
    scored_list = []
    for req in requirements:
        print(f"  Scoring: {req.get('id', '?')}")
        scored_list.append(score_requirement(req))
        time.sleep(2)
 
    total_low  = sum(s["hrs_low"]  for s in scored_list)
    total_mid  = sum(s["hrs_mid"]  for s in scored_list)
    total_high = sum(s["hrs_high"] for s in scored_list)
 
    totals = {
        "total_low":  total_low,
        "total_mid":  total_mid,
        "total_high": total_high,
        "req_count":  len(scored_list),
    }
 
    # ── Break down hours by role and add 8-hour buffer per role ──
    breakdown, new_total_mid = breakdown_hours_by_role(total_mid)
   
    # Re-map to size using the new total with buffer
    size_info = map_to_size(new_total_mid)
    out_dir, docx_path, txt_path = project_output_paths(card_path, card)
 
    build_docx(scored_list, size_info, totals, project_name, docx_path, breakdown=breakdown)
    build_txt(scored_list, size_info, totals, project_name, txt_path, breakdown=breakdown)
 
    # Stamp card with estimation artifacts
    _stamp_estimation_approval(card_path, docx_path, txt_path, "Web UI", "Web Approval")
 
    txt_content = txt_path.read_text(encoding="utf-8") if txt_path.exists() else ""
 
    return {
        "txt_content":  txt_content,
        "txt_path":     str(txt_path),
        "docx_path":    str(docx_path),
        "size":         size_info["size"],
        "investment":   size_info["investment"],
        "total_low":    total_low,
        "total_mid":    total_mid,
        "total_high":   total_high,
        "total_with_buffer": new_total_mid,
        "breakdown":    breakdown,
        "req_count":    len(scored_list),
    }
 
 
if __name__ == "__main__":
 
    parser = argparse.ArgumentParser(description="Estimation Agent")
    parser.add_argument(
        "--card",
        type=str,
        default="",
        help="Path to a2a_card_*.json from Requirements Agent (auto-detected if not provided)",
    )
    args = parser.parse_args()
 
    card_path = None
   
    # If card path provided, use it
    if args.card:
        card_path = Path(args.card)
    else:
        # Auto-detect: prefer project-scoped cards, then fallback to legacy root cards.
        nested_cards = glob.glob(str(HERE / "projects" / "*" / "cards" / "a2a_card_*.json"))
        legacy_cards = glob.glob(str(HERE / "a2a_card_*.json"))
        cards = sorted(
            nested_cards + legacy_cards,
            key=lambda p: Path(p).stat().st_mtime,
            reverse=True,
        )
        if cards:
            card_path = Path(cards[0])
            print(f"[EstimationAgent] Auto-detected card: {card_path.name}")
   
    if not card_path or not card_path.exists():
        if args.card:
            print(f"[error] Card not found: {args.card}")
        else:
            print(f"[error] No a2a_card_*.json found in {HERE} or {HERE / 'projects'}")
        sys.exit(1)
 
    print(f"[EstimationAgent] Reading card: {card_path.name}")
    card_text = card_path.read_text(encoding="utf-8")
    card = json.loads(card_text)
 
    approval = card.get("approval", {})
    card_status = card.get("status", "pending_approval")
    approval_status = approval.get("status", "pending")
    allowed_statuses = {"approved_for_estimation", "estimation_approved"}
    if card_status not in allowed_statuses or approval_status != "approved":
        print("[error] Card is not approved for estimation.")
        print("[error] Generate the BRD, review it, then approve the card first.")
        print("[error] Expected status in {approved_for_estimation, estimation_approved} and approval.status=approved")
        print("[error] Current status=" + str(card_status) + ", approval.status=" + str(approval_status))
        sys.exit(1)
 
        # ── Read from nested A2A card structure ──────────────────
    project_info = card.get("project", {})
    project_name = project_info.get("name", "Unknown Project")
 
    payload = card.get("payload", {})
    requirements = payload.get("requirements", [])
 
    # ── Fallback: if requirements are dicts with "description" only ──
    # Ensure each requirement has the keys the scorer expects
    cleaned = []
    for req in requirements:
        r = dict()
        r["id"] = req.get("id", req.get("req_id", "REQ-?"))
        r["category"] = req.get("category", req.get("type", "Functional"))
        r["priority"] = req.get("priority", "Medium")
        r["description"] = req.get("description", req.get("title", ""))
        cleaned.append(r)
    requirements = cleaned
 
    print("[EstimationAgent] Project: " + project_name)
    print("[EstimationAgent] Requirements found: " + str(len(requirements)))
 
    if not requirements:
        print("[error] No requirements found in card.")
        sys.exit(1)
 
    print("[EstimationAgent] Scoring " + str(len(requirements)) + " requirements...")
 
    scored_list = []
    for req in requirements:
        print("  Scoring: " + req.get("id", "?"))
        scored = score_requirement(req)
        scored_list.append(scored)
        # Small delay to respect rate limits between API calls
        time.sleep(2)
 
    total_low = 0
    total_mid = 0
    total_high = 0
 
    for s in scored_list:
        total_low = total_low + s["hrs_low"]
        total_mid = total_mid + s["hrs_mid"]
        total_high = total_high + s["hrs_high"]
 
    req_count = len(scored_list)
 
    totals = dict()
    totals["total_low"] = total_low
    totals["total_mid"] = total_mid
    totals["total_high"] = total_high
    totals["req_count"] = req_count
 
    size_info = map_to_size(total_mid)
 
    print("[EstimationAgent] Size: " + size_info["size"])
    print("[EstimationAgent] Hours: " + str(total_low) + " - " + str(total_high))
 
    out_dir, docx_path, txt_path = project_output_paths(card_path, card)
    print("[EstimationAgent] Output directory: " + str(out_dir))
 
    build_docx(scored_list, size_info, totals, project_name, docx_path)
    build_txt(scored_list, size_info, totals, project_name, txt_path)
 
    print("[EstimationAgent] Done!")
 
    # ── Estimation Approval Block ─────────────────────────────────
    print("\n[EstimationAgent] Estimation report generated.")
    est_approval = input(
        '[EstimationAgent] Review the estimation, then type "approved" to finalise, '
        '"not approved" to re-run requirements, or press Enter to leave pending: '
    ).strip().lower()
 
    if est_approval == "approved":
        approver_name = input("[EstimationAgent] Enter approver name: ").strip() or "Terminal Approval"
        approver_position = input("[EstimationAgent] Enter approver position: ").strip() or "Interactive Review"
        _stamp_estimation_approval(card_path, docx_path, txt_path, approver_name, approver_position)
    elif est_approval == "not approved":
        _retrigger_requirements(card_path)
    else:
        print("[EstimationAgent] Estimation left as pending. No further action taken.")
 
 