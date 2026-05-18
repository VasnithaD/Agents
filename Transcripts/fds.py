# === Paths (.env, pem) ===
from dotenv import load_dotenv
import os
import json
import argparse
import glob
import urllib.request
import urllib.error
import ssl
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
 
# Load environment variables
SCRIPT_DIR = Path(__file__).parent.resolve()
ENV_PATH   = SCRIPT_DIR / ".env"
PEM_PATH   = SCRIPT_DIR / "cacert 1 (1).pem"
 
load_dotenv(dotenv_path=ENV_PATH, override=True)
 
CLIENT_ID  = os.getenv("CLIENT_ID")
USER_NAME = os.getenv("USER_NAME")
AUTH_TOKEN = os.getenv("AUTHENTICATION_TOKEN")
API_URL    = "https://api.chathpe.it.hpe.com/v2.8/"
 
#  SSL                                                                 #
# ------------------------------------------------------------------ #
 
def build_ssl_context():
    if not PEM_PATH.exists():
        print(f"[ssl] ❌ PEM not found at {PEM_PATH}")
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.load_verify_locations(cafile=str(PEM_PATH))
        print(f"[ssl] ✅ PEM loaded")
        return ctx
    except Exception as e:
        print(f"[ssl] ❌ Failed to load PEM: {e}")
        return None
 
SSL_CONTEXT = build_ssl_context()
 
#  Auth / API                                                          #
# ------------------------------------------------------------------ #
 
def clean(val):
    return val.strip().strip('"').strip("'") if val else val
 
def auth_header():
    if not AUTH_TOKEN:
        raise RuntimeError("AUTHENTICATION_TOKEN missing from .env")
    token = clean(AUTH_TOKEN)
    if token.lower().startswith("bearer "):
        return token
    return f"Bearer {token}"
 
def post(endpoint, payload):
    url     = API_URL + endpoint
    headers = {
        "Client-ID":     clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Content-Type":  "application/json",
        "Cache-Control": "no-cache",
    }
    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, headers=headers, data=data, method="POST")
    try:
        with urllib.request.urlopen(req, context=SSL_CONTEXT) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        print(f"[error] HTTP {e.code} — {body}")
        raise
 
def get(endpoint):
    url     = API_URL + endpoint
    headers = {
        "Client-ID":     clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Cache-Control": "no-cache",
    }
    req = urllib.request.Request(url, headers=headers, method="GET")
    try:
        with urllib.request.urlopen(req, context=SSL_CONTEXT) as resp:
            return resp.read()
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        print(f"[error] HTTP {e.code} — {body}")
        raise
 
# ------------------------------------------------------------------ #
#  Session                                                             #
# ------------------------------------------------------------------ #
 
def setup_session():
    print("[session] Logging in ...")
    resp     = post("login", {"appId": "1"})
    bot_data = resp["chatHPE_bot_data"]
    user_id  = bot_data["userId"]
    username = bot_data["username"]
    print(f"[session] Logged in as {username}")
 
    post("preferences", {
        "agreement": True,
        "dark_mode": True,
        "stream":    False,
        "chatHPE_bot_data": {
            "appId":     "1",
            "sessionId": "-1",
            "userId":    user_id,
            "username":  username,
        },
        "webScraping": False,
    })
 
    raw        = get("sessionId_generator")
    session_id = str(raw.split()[2])[2:-2]
    print(f"[session] session_id={session_id}")
    return user_id, username, session_id
 
# ------------------------------------------------------------------ #
#  Ask                                                                 #
# ------------------------------------------------------------------ #
 
def ask(prompt, user_id, username, session_id):
    endpoint = (
        "call/chatlite"
        "?force_async=false"
        "&session_management_support=true"
        "&internal_call=false"
        "&proxy=false"
    )
    resp = post(endpoint, {
        "chatHPE_bot_data": {
            "appId":     "1",
            "sessionId": session_id,
            "userId":    user_id,
            "username":  username,
        },
        "model_name":  "gpt-4o-mini",
        "stream":      False,
        "webScraping": False,
        "user_query":  prompt,
    })
    return resp.get("Response", "")
 
# ------------------------------------------------------------------ #
#  Prompts                                                           #
# ------------------------------------------------------------------ #
FDS_SECTION_ORDER = [
    "Document Control",
    "Executive Summary",
    "Introduction",
    "Business Context",
    "Stakeholder Analysis",
    "System Overview Diagram",
    "Functional Requirements",
    "User Workflows",
    "Feature Specifications",
    "Data Requirements",
    "Business Rules",
    "Non-Functional Requirements",
    "User Interface Requirements",
    "Integration Points",
    "Assumptions & Constraints",
    "Risks & Mitigation",
    "Acceptance Criteria",
    "Traceability Matrix",
    "Appendices",
]
 
PROMPT_FDS = """You are the FDS Agent. Based on the following BRD, produce a detailed Functional Design Specification Document (FDS).
 
Output MUST contain exactly these 19 sections in this exact order and title format:
1. Document Control
2. Executive Summary
3. Introduction
4. Business Context
5. Stakeholder Analysis
6. System Overview Diagram
7. Functional Requirements
8. User Workflows
9. Feature Specifications
10. Data Requirements
11. Business Rules
12. Non-Functional Requirements
13. User Interface Requirements
14. Integration Points
15. Assumptions & Constraints
16. Risks & Mitigation
17. Acceptance Criteria
18. Traceability Matrix
19. Appendices
 
Instructions:
- Use clear, professional language suitable for a formal specification document.
- Document Control: Include Version, Date, Author, Approval Status
- Executive Summary: High-level overview of the application and its objectives
- Introduction: Include Purpose, Project Scope, System Overview subsections
- Business Context: Include Project Background, Business Needs, Objectives, Expected Benefits
- Stakeholder Analysis: List stakeholders with their roles, responsibilities, and interests
- System Overview Diagram: Describe the system architecture and component interactions (can be mermaid format)
- Functional Requirements: Organize by feature/module with FR-### codes and descriptions
- User Workflows: Include sequence diagrams or step-by-step workflows for key processes
- Feature Specifications: Include Inputs, Outputs, Business Rules, Validation for each major feature
- Data Requirements: Include Field Name, Data Type, Validation Rules, Format as a table
- Business Rules: List business constraints and operational rules
- Non-Functional Requirements: Cover Performance, Security, Scalability, Usability
- User Interface Requirements: Include Screen Layouts, Navigation Flow, UI Behavior
- Integration Points: List external systems, interfaces, and data exchanges
- Assumptions & Constraints: List assumptions and regulatory/technical constraints
- Risks & Mitigation: Include Risk, Likelihood, Impact, Mitigation Strategy in table format
- Acceptance Criteria: List criteria that must be met for acceptance
- Traceability Matrix: Map User Stories to Functional Requirements
- Appendices: Include Glossary, Acronyms, and References
 
BRD Content:
{requirements}
"""
 
PROMPT_SUMMARY = """Provide a concise summary of the FDS output in this exact format.
Output exactly these keys in plain text (no markdown, no bold, no quotes):
KEY_ACTORS: <comma-separated actors/stakeholders>
KEY_SYSTEMS: <comma-separated systems/services/integrations>
KEY_ASSUMPTIONS: <comma-separated assumptions>
 
If information is not available, use: Not specified
 
FDS Content:
{fds}
"""
 
PROMPT_WORKFLOW_DIAGRAM = """You are the FDS Agent.
Create one end-to-end workflow diagram from the FDS content below.
 
Rules:
- Output ONLY Mermaid flowchart syntax, no explanations.
- First line must be exactly: flowchart TD
- Use rectangle nodes for process steps, e.g., A[Submit Request]
- Use one decision node, e.g., D{{Manager Approves?}}
- Use arrows between nodes, e.g., A --> B
- Include 8 to 15 nodes.
- Keep node labels short and business-readable.
 
Reference pattern:
flowchart TD
    A[Employee Submits Leave Request] --> B[Validate Request Data]
    B --> C[Send to Reporting Manager]
    C --> D{{Manager Approves?}}
    D -->|Yes| E[Update Leave Balance]
    D -->|No| F[Notify Rejection]
    E --> G[Send Approval Notification]
    F --> G
 
FDS Content:
{fds}
"""
 
 
def sanitize_line(text: str) -> str:
    """Normalize noisy model output while keeping business content intact."""
    if not text:
        return ""
 
    cleaned = text.strip()
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^[-*]\s+", "", cleaned)
 
    # Remove ALL bold/italic markdown markers anywhere in the line.
    cleaned = re.sub(r"\*{1,3}", "", cleaned)
    cleaned = re.sub(r"_{1,2}(.*?)_{1,2}", r"\1", cleaned)
 
    if (
        (cleaned.startswith('"') and cleaned.endswith('"'))
        or (cleaned.startswith("'") and cleaned.endswith("'"))
    ) and len(cleaned) > 1:
        cleaned = cleaned[1:-1].strip()
 
    cleaned = cleaned.strip("`")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()
 
 
def sanitize_multiline_text(text: str) -> str:
    lines = [sanitize_line(line) for line in text.splitlines()]
    compact = []
    blank_pending = False
 
    for line in lines:
        if not line:
            if compact and not blank_pending:
                compact.append("")
            blank_pending = True
            continue
        compact.append(line)
        blank_pending = False
 
    return "\n".join(compact).strip()
 
 
def sanitize_mermaid(raw_text: str) -> str:
    text = raw_text.strip()
    fence = re.search(r"```(?:mermaid)?\s*(.*?)```", text, flags=re.S | re.I)
    if fence:
        text = fence.group(1).strip()
 
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "flowchart TD\n    A[Start] --> B[Process] --> C[End]"
 
    if not lines[0].lower().startswith("flowchart"):
        lines.insert(0, "flowchart TD")
 
    # Remove quote wrappers around diagram lines.
    lines = [sanitize_line(line) for line in lines if sanitize_line(line)]
 
    return "\n".join(lines)
 
 
def build_fallback_workflow_diagram(sections: dict) -> str:
    process_text = sections.get("Process Flows", "")
    candidates = []
    for line in process_text.splitlines():
        cleaned = sanitize_line(line)
        cleaned = re.sub(r"^\d+[\.)\-\s]+", "", cleaned)
        if cleaned:
            candidates.append(cleaned)
 
    steps = candidates[:6] if candidates else [
        "Employee submits leave request",
        "System validates leave data",
        "Manager reviews request",
        "HR receives outcome",
    ]
 
    node_ids = ["A", "B", "C", "E", "F", "G"]
    diagram_lines = ["flowchart TD"]
 
    # Build a linear pre-decision path with rectangle nodes.
    prev_id = None
    for idx, step in enumerate(steps[:3]):
        nid = node_ids[idx]
        label = step[:60]
        diagram_lines.append(f"    {nid}[{label}]")
        if prev_id:
            diagram_lines.append(f"    {prev_id} --> {nid}")
        prev_id = nid
 
    if prev_id is None:
        diagram_lines.append("    A[Start Leave Request]")
        prev_id = "A"
 
    diagram_lines.append("    D{Manager Approves?}")
    diagram_lines.append(f"    {prev_id} --> D")
 
    approve_step = steps[3] if len(steps) > 3 else "Update leave balance"
    reject_step = steps[4] if len(steps) > 4 else "Notify rejection to employee"
    close_step = steps[5] if len(steps) > 5 else "Close request and audit log"
 
    diagram_lines.append(f"    E[{approve_step[:60]}]")
    diagram_lines.append(f"    F[{reject_step[:60]}]")
    diagram_lines.append(f"    G[{close_step[:60]}]")
    diagram_lines.append("    D -->|Yes| E")
    diagram_lines.append("    D -->|No| F")
    diagram_lines.append("    E --> G")
    diagram_lines.append("    F --> G")
 
    return "\n".join(diagram_lines)
 
 
def ensure_workflow_diagram_format(diagram_text: str, sections: dict) -> str:
    diagram = sanitize_mermaid(diagram_text)
    has_arrow = "-->" in diagram
    has_rect = "[" in diagram and "]" in diagram
    has_decision = "{" in diagram and "}" in diagram
 
    if has_arrow and has_rect and has_decision:
        return diagram
 
    return build_fallback_workflow_diagram(sections)
 
 
def _shade_table_header(table):
    """Apply shading to the first row (header) of a table."""
    for cell in table.rows[0].cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="1F2937"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
        cell._element.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
 
 
def add_table_to_doc(doc, headers: list, rows: list, title: str = ""):
    """Add a formatted table to the document."""
    if title:
        h = doc.add_heading(title, level=2)
        _set_para_spacing(h, space_before_pt=10, space_after_pt=6, line_spacing_pt=16)
   
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
   
    # Add header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
   
    # Add data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = str(cell_value)
            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
   
    _shade_table_header(table)
    return table
 
 
def detect_and_add_table(doc, section: str, content: str):
    """Detect if content should be a table and add it appropriately."""
    # Sections that should have tables
    table_sections = [
        "Stakeholder Analysis",
        "Data Requirements",
        "Traceability Matrix",
        "Risks & Mitigation",
        "Feature Specifications"
    ]
   
    if section not in table_sections:
        return False
   
    # Try to parse table-like content (pipe-separated or multi-line structured data)
    lines = content.strip().split('\n')
    if len(lines) < 2:
        return False
   
    # Check for pipe-separated format (markdown table style)
    if '|' in lines[0]:
        try:
            # Parse markdown-style table
            rows = []
            headers = None
            for line in lines:
                if not line.strip() or '---' in line:
                    continue
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if not cells:
                    continue
                if headers is None:
                    headers = cells
                else:
                    rows.append(cells)
           
            if headers and rows:
                add_table_to_doc(doc, headers, rows)
                return True
        except:
            pass
   
    return False
 
 
# === Parsers (adapted from Estimation style) ===
def parse_fds_sections(raw_text):
    """
    Parse the raw FDS output into structured sections.
    Expected sections are defined by FDS_SECTION_ORDER (19 sections).
    """
    sections = {}
    current_key = None
    buffer = []
 
    def normalize_header(line: str) -> str:
        header = line.strip().lstrip("#").strip()
        header = re.sub(r"^\d+[\.)\-\s]+", "", header)
        header = re.sub(r"\s+", " ", header)
        return header.strip().lower()
 
    def detect_section_key(line: str):
        header = normalize_header(line)
 
        if "document control" in header:
            return "Document Control"
        if "executive summary" in header:
            return "Executive Summary"
        if "introduction" in header and "scope" not in header:
            return "Introduction"
        if "business context" in header:
            return "Business Context"
        if "stakeholder analysis" in header or "stakeholders" in header:
            return "Stakeholder Analysis"
        if "system overview diagram" in header or "system diagram" in header:
            return "System Overview Diagram"
        if "functional requirements" in header:
            return "Functional Requirements"
        if "user workflows" in header or "workflows" in header:
            return "User Workflows"
        if "feature specifications" in header or "feature specification" in header:
            return "Feature Specifications"
        if "data requirements" in header:
            return "Data Requirements"
        if "business rules" in header:
            return "Business Rules"
        if "non-functional requirements" in header or "non functional requirements" in header:
            return "Non-Functional Requirements"
        if "user interface" in header and "requirements" in header:
            return "User Interface Requirements"
        if "integration points" in header:
            return "Integration Points"
        if "assumptions" in header and "constraints" in header:
            return "Assumptions & Constraints"
        if "risks" in header and "mitigation" in header:
            return "Risks & Mitigation"
        if "acceptance criteria" in header:
            return "Acceptance Criteria"
        if "traceability" in header and "matrix" in header:
            return "Traceability Matrix"
        if "appendices" in header or "appendix" in header:
            return "Appendices"
 
        return None
 
    for line in raw_text.strip().splitlines():
        stripped = line.strip()
        if not stripped:
            continue
 
        detected_key = detect_section_key(stripped)
        if detected_key:
            if current_key and buffer:
                sections[current_key] = "\n".join(buffer).strip()
                buffer = []
            current_key = detected_key
        else:
            if current_key:
                buffer.append(sanitize_line(stripped))
 
    # Save last section
    if current_key and buffer:
        sections[current_key] = "\n".join(buffer).strip()
 
    # Ensure the result always contains all required sections in order.
    return {
        key: sanitize_multiline_text(
            sections.get(key, "Not enough detail provided in requirements.")
        )
        for key in FDS_SECTION_ORDER
    }
 
def parse_fds_summary(raw_text):
    """
    Parse summary-style key:value lines from FDS output.
    """
    result = {}
    for line in raw_text.strip().splitlines():
        if ":" in line:
            key, _, val = line.partition(":")
            normalized_key = sanitize_line(key).upper().replace(" ", "_")
            cleaned_val = sanitize_line(val)
 
            if "ACTOR" in normalized_key:
                result["KEY_ACTORS"] = cleaned_val or "Not specified"
            elif "SYSTEM" in normalized_key:
                result["KEY_SYSTEMS"] = cleaned_val or "Not specified"
            elif "ASSUMPTION" in normalized_key:
                result["KEY_ASSUMPTIONS"] = cleaned_val or "Not specified"
 
    for required_key in ["KEY_ACTORS", "KEY_SYSTEMS", "KEY_ASSUMPTIONS"]:
        result.setdefault(required_key, "Not specified")
 
    return result
 
# ------------------------------------------------------------------ #
#  DOCX Builder                                                      #
# ------------------------------------------------------------------ #
def _set_para_spacing(para, space_before_pt=0, space_after_pt=6, line_spacing_pt=14):
    """Apply paragraph spacing and line height via pPr XML."""
    from docx.shared import Pt as _Pt
    pf = para.paragraph_format
    pf.space_before = _Pt(space_before_pt)
    pf.space_after  = _Pt(space_after_pt)
    pf.line_spacing = _Pt(line_spacing_pt)
 
 
def _add_body_para(doc, text, indent_left=0):
    """Add a justified body paragraph with consistent spacing and font."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(para, space_before_pt=2, space_after_pt=6, line_spacing_pt=14)
    if indent_left:
        para.paragraph_format.left_indent = Inches(indent_left)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return para
 
 
def build_fds_docx(sections: dict, summary: dict, user_name: str, workflow_diagram: str = ""):
    doc = Document()
 
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
 
    # Cover / title block
    title = doc.add_heading("Functional Specification Document (FDS)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(title, space_before_pt=0, space_after_pt=12, line_spacing_pt=18)
 
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(meta, space_before_pt=0, space_after_pt=4, line_spacing_pt=13)
    run = meta.add_run(f"Generated by: {user_name}     |     Date: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
 
    doc.add_paragraph()  # spacer
 
    # Main FDS sections in order
    for section, content in sections.items():
        h = doc.add_heading(section, level=1)
        _set_para_spacing(h, space_before_pt=14, space_after_pt=4, line_spacing_pt=16)
 
        # Try to detect and add as table first
        if not detect_and_add_table(doc, section, content):
            # Otherwise add as regular paragraphs
            for line in content.splitlines():
                line = line.strip()
                if not line:
                    continue
                # Detect list items (lines starting with -, •, or digit+dot)
                is_bullet = re.match(r'^[-•]\s+', line) or re.match(r'^\d+[.)\s]\s+', line)
                text = re.sub(r'^[-•]\s+', '', line)
                text = re.sub(r'^\d+[.)\s]\s+', '', text)
                p = _add_body_para(doc, ("• " if is_bullet else "") + text,
                                   indent_left=0.3 if is_bullet else 0)
 
    return doc
 
# ------------------------------------------------------------------ #
#  Main entry point (standalone)                                     #
# ------------------------------------------------------------------ #
def run_fds(requirements_text: str) -> dict:
    """
    Run the FDS agent on a block of requirements text.
    Returns dict with 'sections', 'summary'.
    Can be called from app.py / orchestrator.
    """
    user_id, username, session_id = setup_session()
 
    print("[fds] Step 1/2 — Generating FDS ...")
    raw_fds = ask(
        PROMPT_FDS.format(requirements=requirements_text),
        user_id, username, session_id
    )
 
    print("[fds] Step 2/2 — Building executive summary ...")
    raw_summary = ask(
        PROMPT_SUMMARY.format(fds=raw_fds),
        user_id, username, session_id
    )
 
    raw_fds_clean = sanitize_multiline_text(raw_fds)
    raw_summary_clean = sanitize_multiline_text(raw_summary)
 
    sections = parse_fds_sections(raw_fds_clean)
    summary  = parse_fds_summary(raw_summary_clean)
 
    print(f"[fds] Parsed {len(sections)} sections")
    return {
        "sections":   sections,
        "summary":    summary,
        "raw_fds":    raw_fds_clean,
        "raw_summary": raw_summary_clean,
    }
 
 
def load_a2a_card(card_path: Path) -> dict:
    """Load and validate an a2a JSON card. Returns the card dict."""
    if card_path.suffix.lower() != ".json":
        raise ValueError(
            f"❌ Unsupported file type: {card_path.suffix}. Only .json (a2a format) is accepted."
        )
 
    try:
        with open(card_path, "r", encoding="utf-8") as f:
            card = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"❌ Invalid JSON file: {e}")
    except Exception as e:
        raise ValueError(f"❌ Failed to read JSON file: {e}")
 
    if not isinstance(card, dict):
        raise ValueError("❌ Invalid a2a card: root must be a JSON object")
 
    payload = card.get("payload", {})
    if not isinstance(payload, dict):
        raise ValueError("❌ Invalid a2a card: 'payload' must be an object")
 
    requirements = payload.get("requirements", [])
    if not isinstance(requirements, list) or not requirements:
        raise ValueError("❌ Invalid a2a card: 'requirements' must be a non-empty array")
 
    return card
 
 
def extract_requirements_text(card: dict) -> str:
    """Convert a2a card requirements into estimator-friendly text lines."""
    requirements = card.get("payload", {}).get("requirements", [])
    lines = []
    for req in requirements:
        if not isinstance(req, dict):
            continue
        req_id = str(req.get("id", "")).strip()
        category = str(req.get("category", "Functional")).strip() or "Functional"
        priority = str(req.get("priority", "Medium")).strip() or "Medium"
        description = str(req.get("description", "")).strip()
        if req_id and description:
            lines.append(f"{req_id} | {category} | {priority} | {description}")
 
    if not lines:
        raise ValueError("❌ Invalid a2a card: no valid requirements found")
 
    return "\n".join(lines)
 
 
def get_project_name_from_card(card: dict, default: str = "Project") -> str:
    """Extract project name from an a2a card."""
    payload = card.get("payload", {})
    overview = payload.get("overview", {})
    project = card.get("project", {})
    return (
        overview.get("PROJECT_NAME")
        or project.get("name")
        or default
    )
 
 
def resolve_project_paths(card_path: Path, card: dict):
    """Resolve project-scoped directories for FDS artifacts."""
    artifacts = card.get("artifacts", {})
    project_root = artifacts.get("output_dir", "")

    if project_root:
        project_dir = Path(project_root)
    else:
        parent = card_path.parent
        project_dir = parent.parent if parent.name == "cards" else Path(".").resolve()

    fds_dir = project_dir / "fds"
    cards_dir = project_dir / "cards"
    fds_dir.mkdir(parents=True, exist_ok=True)
    cards_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = fds_dir / f"FDS_{timestamp}.docx"
    txt_path = fds_dir / f"FDS_{timestamp}.txt"
    return project_dir, fds_dir, cards_dir, docx_path, txt_path


def build_fds_txt(project_name: str, sections: dict, summary: dict, workflow_diagram: str = "") -> str:
    lines = []
    lines.append("FUNCTIONAL DESIGN SPECIFICATION")
    lines.append(f"Project: {project_name}")
    lines.append(f"Generated At: {datetime.now().isoformat()}")
    lines.append("")
    lines.append("SUMMARY")
    lines.append(f"KEY_ACTORS: {summary.get('KEY_ACTORS', 'Not specified')}")
    lines.append(f"KEY_SYSTEMS: {summary.get('KEY_SYSTEMS', 'Not specified')}")
    lines.append(f"KEY_ASSUMPTIONS: {summary.get('KEY_ASSUMPTIONS', 'Not specified')}")
    lines.append("")

    for name in FDS_SECTION_ORDER:
        lines.append(name.upper())
        lines.append(sections.get(name, "Not enough detail provided in requirements."))
        lines.append("")

    if workflow_diagram:
        lines.append("WORKFLOW_DIAGRAM_MERMAID")
        lines.append(workflow_diagram)
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def _stamp_fds_on_input_card(input_card_path: Path, fds_docx_path: Path, fds_txt_path: Path, tds_card_path: Path):
    """Update the Requirement A2A card with FDS artifacts for orchestration."""
    try:
        card = json.loads(input_card_path.read_text(encoding="utf-8"))
        card["fds_artifacts"] = {
            "fds_docx": str(fds_docx_path),
            "fds_txt": str(fds_txt_path),
            "tds_card": str(tds_card_path),
        }
        card["status"] = "fds_generated"
        card["updated_at"] = datetime.now().isoformat()
        input_card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
    except Exception as exc:
        print(f"[fds] Warning: could not stamp input card with FDS artifacts: {exc}")


def save_tds_a2a_card(
    sections: dict,
    summary: dict,
    raw_fds: str,
    project_name: str,
    input_card: dict,
    fds_docx_path: Path,
    fds_txt_path: Path,
    output_dir: Path,
) -> Path:
    """Save a2a card for the TDS agent to consume."""
    import uuid

    slug = re.sub(r"[^a-z0-9]+", "_", project_name.lower()).strip("_") or "project"

    card = {
        "a2a_version": "1.0",
        "card_id": str(uuid.uuid4()),
        "created_at": datetime.now().isoformat(),

        "source_agent": {
            "name": "FDSAgent",
            "version": "1.0",
            "script": "FDS.py",
        },

        "target_agent": {
            "name": "TDSAgent",
            "version": "1.0",
            "script": "TDS.py",
            "action": "run_tds_from_card",
        },

        "project": {
            "name": project_name,
            "source_transcript": input_card.get("project", {}).get("source_transcript", ""),
            "processed_at": datetime.now().isoformat(),
        },

        "artifacts": {
            "input_card": str(input_card.get("artifacts", {}).get("brd_txt", "")),
            "fds_docx": str(fds_docx_path),
            "fds_txt": str(fds_txt_path),
            "output_dir": str(output_dir),
        },

        "payload": {
            "fds_sections": sections,
            "raw_fds": raw_fds,
            "summary": summary,
        },

        "stats": {
            "total_fds_sections": len(sections),
        },

        "status": "ready_for_tds",
    }

    card_path = output_dir / f"tds_a2a_card_{slug}.json"
    with open(card_path, "w", encoding="utf-8") as f:
        json.dump(card, f, indent=2)

    print(f"[fds] A2A card for TDS saved -> {card_path}")
    return card_path


def run_fds_pipeline(card_path_str: str) -> dict:
    """Run FDS pipeline from Requirement A2A card and write project-scoped artifacts."""
    card_path = Path(card_path_str).expanduser().resolve()
    if not card_path.exists():
        raise FileNotFoundError(f"A2A card not found: {card_path}")

    input_card = load_a2a_card(card_path)
    requirements_text = extract_requirements_text(input_card)
    project_name = get_project_name_from_card(input_card)

    result = run_fds(requirements_text)

    user_id, username, session_id = setup_session()
    raw_workflow = ask(
        PROMPT_WORKFLOW_DIAGRAM.format(fds=result["raw_fds"]),
        user_id,
        username,
        session_id,
    )
    workflow_diagram = ensure_workflow_diagram_format(raw_workflow, result["sections"])

    project_dir, fds_dir, cards_dir, docx_path, txt_path = resolve_project_paths(card_path, input_card)

    doc = build_fds_docx(
        result["sections"],
        result["summary"],
        USER_NAME or "FDS Agent",
        workflow_diagram,
    )
    doc.save(str(docx_path))

    txt_content = build_fds_txt(project_name, result["sections"], result["summary"], workflow_diagram)
    txt_path.write_text(txt_content, encoding="utf-8")

    tds_card_path = save_tds_a2a_card(
        sections=result["sections"],
        summary=result["summary"],
        raw_fds=result["raw_fds"],
        project_name=project_name,
        input_card=input_card,
        fds_docx_path=docx_path,
        fds_txt_path=txt_path,
        output_dir=cards_dir,
    )

    _stamp_fds_on_input_card(card_path, docx_path, txt_path, tds_card_path)

    return {
        "project_name": project_name,
        "docx_path": str(docx_path),
        "txt_path": str(txt_path),
        "tds_card_path": str(tds_card_path),
        "txt_content": txt_content,
        "sections": result["sections"],
        "summary": result["summary"],
        "project_dir": str(project_dir),
        "fds_dir": str(fds_dir),
        "cards_dir": str(cards_dir),
    }


def auto_detect_a2a_card(base_dir: Path):
    """Auto-detect latest requirement A2A card in current project/workspace."""
    patterns = [
        str(base_dir / "cards" / "a2a_card_*.json"),
        str(base_dir / "a2a_card_*.json"),
        str(base_dir / "projects" / "*" / "cards" / "a2a_card_*.json"),
        str(SCRIPT_DIR / "projects" / "*" / "cards" / "a2a_card_*.json"),
    ]

    candidates = []
    for pattern in patterns:
        candidates.extend(glob.glob(pattern))

    if not candidates:
        return None

    candidates = sorted(candidates, key=lambda p: Path(p).stat().st_mtime, reverse=True)
    return Path(candidates[0]).resolve()


if __name__ == "__main__":
    if not CLIENT_ID or not AUTH_TOKEN:
        print("[fds] CLIENT_ID or AUTHENTICATION_TOKEN missing from .env")
        raise SystemExit(1)

    parser = argparse.ArgumentParser(description="FDS Agent")
    parser.add_argument(
        "--card",
        type=str,
        help="Path to Requirement a2a_card_*.json (auto-detected if omitted)",
    )
    args = parser.parse_args()

    if args.card:
        card_path = Path(args.card).expanduser().resolve()
    else:
        detected = auto_detect_a2a_card(Path.cwd())
        if not detected:
            print("[fds] No a2a_card_*.json found in current project/workspace.")
            raise SystemExit(1)
        card_path = detected
        print(f"[fds] Auto-detected card: {card_path}")

    try:
        out = run_fds_pipeline(str(card_path))
    except Exception as exc:
        print(f"[fds] Failed: {exc}")
        raise SystemExit(1)

    print("[fds] Done")
    print(f"[fds] FDS DOCX: {out['docx_path']}")
    print(f"[fds] FDS TXT: {out['txt_path']}")
    print(f"[fds] TDS A2A card: {out['tds_card_path']}")
    print(f"[fds] Project dir: {out['project_dir']}")