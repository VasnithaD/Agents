from dotenv import load_dotenv
import os
import json
import argparse
import sys
import urllib.request
import urllib.error
import ssl
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor ,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import subprocess
from rag_agent_context import ground_prompt_with_rag

# ------------------------------------------------------------------ #
#  Paths                                                               #
# ------------------------------------------------------------------ #

SCRIPT_DIR = Path(__file__).parent.resolve()
ENV_PATH   = SCRIPT_DIR / ".env"
PEM_PATH   = SCRIPT_DIR / "cacert 1 (1).pem"

load_dotenv(dotenv_path=ENV_PATH, override=True)

CLIENT_ID  = os.getenv("CLIENT_ID")
AUTH_TOKEN = os.getenv("AUTHENTICATION_TOKEN")
API_URL    = "https://api.chathpe.it.hpe.com/v2.8/"

# ------------------------------------------------------------------ #
#  SSL                                                                 #
# ------------------------------------------------------------------ #

def build_ssl_context():
    if not PEM_PATH.exists():
        print(f"[ssl] ❌ PEM not found at {PEM_PATH}")
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.load_verify_locations(cafile=str(PEM_PATH))
        print(f"[ssl] ✅ PEM loaded")
        return ctx
    except Exception as e:
        print(f"[ssl] ❌ Failed to load PEM: {e}")
        return None

SSL_CONTEXT = build_ssl_context()

# ------------------------------------------------------------------ #
#  Auth                                                                #
# ------------------------------------------------------------------ #

def clean(val):
    return val.strip().strip('"').strip("'") if val else val

def auth_header():
    if not AUTH_TOKEN:
        raise RuntimeError("AUTHENTICATION_TOKEN missing from .env")
    token = clean(AUTH_TOKEN)
    if token.lower().startswith("bearer "):
        return token
    return f"Bearer {token}"

# ------------------------------------------------------------------ #
#  API calls                                                           #
# ------------------------------------------------------------------ #

def post(endpoint, payload):
    url     = API_URL + endpoint
    headers = {
        "Client-ID":     clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Content-Type":  "application/json",
        "Cache-Control": "no-cache",
    }
    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, headers=headers, data=data, method="POST")
    try:
        with urllib.request.urlopen(req, context=SSL_CONTEXT) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        print(f"[error] HTTP {e.code} — {body}")
        raise

def get(endpoint):
    url     = API_URL + endpoint
    headers = {
        "Client-ID":     clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Cache-Control": "no-cache",
    }
    req = urllib.request.Request(url, headers=headers, method="GET")
    try:
        with urllib.request.urlopen(req, context=SSL_CONTEXT) as resp:
            return resp.read()
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        print(f"[error] HTTP {e.code} — {body}")
        raise

# ------------------------------------------------------------------ #
#  Session                                                             #
# ------------------------------------------------------------------ #

def setup_session():
    print("[session] Logging in ...")
    resp     = post("login", {"appId": "1"})
    bot_data = resp["chatHPE_bot_data"]
    user_id  = bot_data["userId"]
    username = bot_data["username"]
    print(f"[session] Logged in as {username}")

    post("preferences", {
        "agreement": True,
        "dark_mode": True,
        "stream":    False,
        "chatHPE_bot_data": {
            "appId":     "1",
            "sessionId": "-1",
            "userId":    user_id,
            "username":  username,
        },
        "webScraping": False,
    })

    raw        = get("sessionId_generator")
    session_id = str(raw.split()[2])[2:-2]
    print(f"[session] session_id={session_id}")
    return user_id, username, session_id

# ------------------------------------------------------------------ #
#  Ask                                                                 #
# ------------------------------------------------------------------ #

def ask(prompt, user_id, username, session_id):
    grounded_prompt = ground_prompt_with_rag(
        prompt,
        task_label="BRD extraction from transcript",
    )
    endpoint = (
        "call/chatlite"
        "?force_async=false"
        "&session_management_support=true"
        "&internal_call=false"
        "&proxy=false"
    )
    resp = post(endpoint, {
        "chatHPE_bot_data": {
            "appId":     "1",
            "sessionId": session_id,
            "userId":    user_id,
            "username":  username,
        },
        "model_name":  "gpt-4o-mini",
        "stream":      False,
        "webScraping": False,
        "user_query":  grounded_prompt,
    })
    return resp.get("Response", "")

# ------------------------------------------------------------------ #
#  Prompts                                                             #
# ------------------------------------------------------------------ #

PROMPT_OVERVIEW = """
You are a Business Analyst writing a Business Requirements Document (BRD).
Read the transcript below and extract:
1. Project Name (guess from context)
2. Project Purpose (2-3 sentences)
3. Project Scope (what is included and excluded)
4. Key Stakeholders mentioned
5. Project Background / Problem Statement

Format your response exactly like this with no extra text:

PROJECT_NAME: <name>
PURPOSE: <2-3 sentence purpose>
SCOPE_IN: <what is in scope>
SCOPE_OUT: <what is out of scope>
STAKEHOLDERS: <comma separated list>
BACKGROUND: <background paragraph>

TRANSCRIPT:
{transcript}
"""

PROMPT_REQUIREMENTS = """
You are a Business Analyst writing a Business Requirements Document (BRD).
Read the transcript below and extract ALL requirements — functional, non-functional,
technical, security, performance, and any constraints or assumptions.

For EACH requirement output exactly this format (one per line):
REQ-001 | Functional | High | <clear one sentence requirement>
REQ-002 | Non-Functional | Medium | <clear one sentence requirement>

Categories must be one of:
Functional | Non-Functional | Technical | Security | Performance | Constraint | Assumption

Priority must be one of:
High | Medium | Low

Extract at least 10-15 requirements if the transcript supports it.
Only output the requirement lines, nothing else, no extra commentary.

TRANSCRIPT:
{transcript}
"""

PROMPT_RISKS = """
You are a Business Analyst writing a Business Requirements Document (BRD).
Read the transcript below and identify risks, assumptions, constraints and dependencies.

Format exactly like this with no extra text:

RISK-001: <risk description>
RISK-002: <risk description>
ASSUMPTION-001: <assumption>
ASSUMPTION-002: <assumption>
CONSTRAINT-001: <constraint>
CONSTRAINT-002: <constraint>
DEPENDENCY-001: <dependency>

TRANSCRIPT:
{transcript}
"""

# ------------------------------------------------------------------ #
#  Parsers                                                             #
# ------------------------------------------------------------------ #

def parse_overview(raw_text):
    result = {}
    for line in raw_text.strip().splitlines():
        if ":" in line:
            key, _, val = line.partition(":")
            result[key.strip()] = val.strip()
    return result

def parse_requirements(raw_text):
    requirements = []
    for line in raw_text.strip().splitlines():
        line = line.strip()
        if not line or "|" not in line:
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) == 4:
            requirements.append({
                "id":          parts[0],
                "category":    parts[1],
                "priority":    parts[2],
                "description": parts[3],
            })
    return requirements

def parse_risks(raw_text):
    risks        = []
    assumptions  = []
    constraints  = []
    dependencies = []
    for line in raw_text.strip().splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        key, _, val = line.partition(":")
        key = key.strip().upper()
        val = val.strip()
        if key.startswith("RISK"):
            risks.append(val)
        elif key.startswith("ASSUMPTION"):
            assumptions.append(val)
        elif key.startswith("CONSTRAINT"):
            constraints.append(val)
        elif key.startswith("DEPENDENCY"):
            dependencies.append(val)
    return risks, assumptions, constraints, dependencies

# ------------------------------------------------------------------ #
#  DOCX Builder                                                        #
# ------------------------------------------------------------------ #

HPE_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
HPE_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

PRIORITY_COLORS = {
    "High":   "FFD7D7",
    "Medium": "FFF3CD",
    "Low":    "D4EDDA",
}

def set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml   import parse_xml
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shading_elm)

def add_heading(doc, text, level=1):
    p   = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = HPE_BLUE
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row styling
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run  = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = HPE_WHITE
        set_cell_bg(cell, "1F497D")

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            # Color priority column
            if i < len(headers) and headers[i] == "Priority":
                color = PRIORITY_COLORS.get(val)
                if color:
                    set_cell_bg(row_cells[i], color)
    return table

def slugify_project_name(project_name):
    slug = "".join(c if c.isalnum() else "_" for c in project_name).lower().strip("_")
    return slug or "unknown"

def build_pending_approval():
    return {
        "status": "pending",
        "approved_by": "",
        "approver_role": "",
        "approved_at": "",
        "comments": "Awaiting business approval before estimation."
    }

def build_approval_rows(approval):
    status = approval.get("status", "pending").replace("_", " ").title()
    return [
        ["Approval Status", status],
        ["Approved By", approval.get("approved_by", "") or "Pending"],
        ["Approver Role", approval.get("approver_role", "") or "Pending"],
        ["Approved At", approval.get("approved_at", "") or "Pending"],
        ["Comments", approval.get("comments", "") or "Pending approval"],
    ]

def write_brd_txt(output_path, overview, requirements, risks, assumptions,
                  constraints, dependencies, source_filename, approval, related_folder="", related_file=""):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=" * 70 + "\n")
        f.write("BUSINESS REQUIREMENTS DOCUMENT\n")
        f.write(f"Project : {overview.get('PROJECT_NAME', 'N/A')}\n")
        f.write(f"Related File : {related_file or 'N/A'}\n")
        f.write(f"Related Folder : {related_folder or 'N/A'}\n")
        f.write(f"Date    : {datetime.now().strftime('%Y-%m-%d')}\n")
        f.write(f"Source  : {source_filename}\n")
        f.write("=" * 70 + "\n\n")

        f.write("1. EXECUTIVE SUMMARY\n")
        f.write("-" * 40 + "\n")
        f.write(overview.get("PURPOSE", "N/A") + "\n\n")

        f.write("2. PROJECT BACKGROUND\n")
        f.write("-" * 40 + "\n")
        f.write(overview.get("BACKGROUND", "N/A") + "\n\n")

        f.write("3. SCOPE\n")
        f.write("-" * 40 + "\n")
        f.write(f"In Scope  : {overview.get('SCOPE_IN', 'N/A')}\n")
        f.write(f"Out Scope : {overview.get('SCOPE_OUT', 'N/A')}\n\n")

        f.write("4. STAKEHOLDERS\n")
        f.write("-" * 40 + "\n")
        for stakeholder in overview.get("STAKEHOLDERS", "N/A").split(","):
            f.write(f"  - {stakeholder.strip()}\n")
        f.write("\n")

        f.write("5. REQUIREMENTS\n")
        f.write("-" * 40 + "\n")
        f.write(f"{'ID':<12} {'Category':<18} {'Priority':<10} Description\n")
        f.write(f"{'-'*12} {'-'*18} {'-'*10} {'-'*40}\n")
        for requirement in requirements:
            f.write(
                f"{requirement['id']:<12} {requirement['category']:<18} "
                f"{requirement['priority']:<10} {requirement['description']}\n"
            )
        f.write("\n")

        f.write("6. RISKS\n")
        f.write("-" * 40 + "\n")
        for idx, risk in enumerate(risks, 1):
            f.write(f"  RISK-{idx:03d}: {risk}\n")
        f.write("\n")

        f.write("7. ASSUMPTIONS\n")
        f.write("-" * 40 + "\n")
        for idx, assumption in enumerate(assumptions, 1):
            f.write(f"  ASSUMPTION-{idx:03d}: {assumption}\n")
        f.write("\n")

        f.write("8. CONSTRAINTS\n")
        f.write("-" * 40 + "\n")
        for idx, constraint in enumerate(constraints, 1):
            f.write(f"  CONSTRAINT-{idx:03d}: {constraint}\n")
        f.write("\n")

        f.write("9. DEPENDENCIES\n")
        f.write("-" * 40 + "\n")
        for idx, dependency in enumerate(dependencies, 1):
            f.write(f"  DEPENDENCY-{idx:03d}: {dependency}\n")
        f.write("\n")

        f.write("10. APPROVAL\n")
        f.write("-" * 40 + "\n")
        for label, value in build_approval_rows(approval):
            f.write(f"{label:<16}: {value}\n")
        f.write("\n")

def refresh_brd_artifacts_from_card(card_path):
    card = json.loads(Path(card_path).read_text(encoding="utf-8"))
    payload = card.get("payload", {})
    artifacts = card.get("artifacts", {})
    project = card.get("project", {})
    approval = card.get("approval", build_pending_approval())

    docx_path = Path(artifacts.get("brd_docx", ""))
    txt_path = Path(artifacts.get("brd_txt", ""))
    source_filename = Path(project.get("source_transcript", "")).name or "Unknown"

    doc = build_docx(
        payload.get("overview", {}),
        payload.get("requirements", []),
        payload.get("risks", []),
        payload.get("assumptions", []),
        payload.get("constraints", []),
        payload.get("dependencies", []),
        source_filename,
        approval,
    )
    artifacts_updated = False

    try:
        doc.save(str(docx_path))
    except PermissionError:
        # Common on Windows when the DOCX is open in Word during approval.
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback_docx = docx_path.with_name(f"{docx_path.stem}_approved_{timestamp}{docx_path.suffix}")
        doc.save(str(fallback_docx))
        docx_path = fallback_docx
        artifacts["brd_docx"] = str(docx_path)
        artifacts_updated = True
        print(f"[agent] BRD DOCX was locked. Saved approved copy instead → {docx_path}")

    try:
        write_brd_txt(
            txt_path,
            payload.get("overview", {}),
            payload.get("requirements", []),
            payload.get("risks", []),
            payload.get("assumptions", []),
            payload.get("constraints", []),
            payload.get("dependencies", []),
            source_filename,
            approval,
            related_folder=str(Path(txt_path).parent),
        )
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback_txt = txt_path.with_name(f"{txt_path.stem}_approved_{timestamp}{txt_path.suffix}")
        write_brd_txt(
            fallback_txt,
            payload.get("overview", {}),
            payload.get("requirements", []),
            payload.get("risks", []),
            payload.get("assumptions", []),
            payload.get("constraints", []),
            payload.get("dependencies", []),
            source_filename,
            approval,
            related_folder=str(Path(fallback_txt).parent),
        )
        txt_path = fallback_txt
        artifacts["brd_txt"] = str(txt_path)
        artifacts_updated = True
        print(f"[agent] BRD TXT was locked. Saved approved copy instead → {txt_path}")

    if artifacts_updated:
        card["artifacts"] = artifacts
        card["updated_at"] = datetime.now().isoformat()
        Path(card_path).write_text(json.dumps(card, indent=2), encoding="utf-8")

def trigger_estimation(card_path):
    estimation_script = SCRIPT_DIR / "Estimation.py"
    if not estimation_script.exists():
        print(f"\n[agent] ℹ️  Estimation.py not found at {estimation_script}")
        print(f"[agent] Run manually: python Estimation.py --card \"{card_path}\"")
        return

    print(f"\n[agent] 🚀 Triggering Estimation Agent ...")
    try:
        subprocess.run(
            [
                os.sys.executable,
                str(estimation_script),
                "--card",
                str(card_path),
            ],
            check=True,
        )
        print("[agent] ✅ Estimation Agent completed!")
    except subprocess.CalledProcessError as e:
        print(f"[agent] ❌ Estimation Agent failed: {e}")

def approve_card(card_path, approved_by, approver_role, approval_comments):
    card_path = Path(card_path)
    if not card_path.exists():
        raise FileNotFoundError(f"Card not found: {card_path}")

    card = json.loads(card_path.read_text(encoding="utf-8"))
    card["approval"] = {
        "status": "approved",
        "approved_by": approved_by,
        "approver_role": approver_role,
        "approved_at": datetime.now().isoformat(),
        "comments": approval_comments or "Approved for estimation."
    }
    card["status"] = "approved_for_estimation"
    card["updated_at"] = datetime.now().isoformat()

    card_path.write_text(json.dumps(card, indent=2), encoding="utf-8")
    refresh_brd_artifacts_from_card(card_path)
    print(f"[agent] Card approved → {card_path.name}")
    trigger_estimation(card_path)

def read_terminal_textbox(prompt_message):
    print(prompt_message)
    print('[agent] Type "END" on a new line to finish.')
    lines = []
    while True:
        line = input()
        if line.strip() == "END":
            break
        lines.append(line)
    return "\n".join(lines).strip()

def save_text_variant(base_path, content, suffix):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    variant_path = base_path.parent / f"{base_path.stem}_{suffix}_{timestamp}.txt"
    variant_path.write_text(content, encoding="utf-8")
    print(f"[agent] Saved updated transcript → {variant_path}")
    return variant_path

def get_projects_root(base_dir):
    return Path(base_dir) / "projects"

def get_project_catalog(base_dir):
    projects_root = get_projects_root(base_dir)
    if not projects_root.exists():
        return []

    catalog = []
    for project_dir in sorted([p for p in projects_root.iterdir() if p.is_dir()]):
        cards_dir = project_dir / "cards"
        brd_dir = project_dir / "brd"

        card_files = sorted(
            cards_dir.glob("a2a_card_*.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        ) if cards_dir.exists() else []

        brd_files = sorted(
            brd_dir.glob("BRD_*.docx"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        ) if brd_dir.exists() else []

        display_name = project_dir.name.replace("_", " ").title()
        source_transcript = ""
        latest_status = "unknown"

        if card_files:
            try:
                latest_card_data = json.loads(card_files[0].read_text(encoding="utf-8"))
                display_name = latest_card_data.get("project", {}).get("name", display_name) or display_name
                source_transcript = latest_card_data.get("project", {}).get("source_transcript", "")
                latest_status = latest_card_data.get("status", "unknown")
            except (json.JSONDecodeError, OSError):
                pass

        catalog.append({
            "slug": project_dir.name,
            "display_name": display_name,
            "project_dir": str(project_dir),
            "cards_dir": str(cards_dir),
            "brd_dir": str(brd_dir),
            "card_count": len(card_files),
            "brd_count": len(brd_files),
            "latest_card": str(card_files[0]) if card_files else "",
            "latest_brd": str(brd_files[0]) if brd_files else "",
            "source_transcript": source_transcript,
            "status": latest_status,
        })

    return catalog

def print_project_catalog(catalog):
    print("\n[agent] Available projects:")
    if not catalog:
        print("[agent] No existing projects found.")
        return

    for idx, project in enumerate(catalog, 1):
        print(
            f"  {idx}. {project['display_name']} "
            f"(slug: {project['slug']}, cards: {project['card_count']}, "
            f"brds: {project['brd_count']}, status: {project['status']})"
        )

def prompt_transcript_path(default_path=""):
    if default_path:
        answer = input(
            f"Enter path to your transcript .txt file (press Enter to use {default_path}): "
        ).strip().strip('"')
        if not answer:
            return Path(default_path)
        return Path(answer)

    raw = input("Enter path to your transcript .txt file: ").strip().strip('"')
    return Path(raw)

def choose_project_workflow(base_dir):
    catalog = get_project_catalog(base_dir)
    print_project_catalog(catalog)

    if catalog:
        choice = input(
            '\n[agent] Type a project number to continue with that project, or type "new" to create a new project: '
        ).strip().lower()
    else:
        choice = "new"

    if choice == "new":
        project_name = input(
            "[agent] Enter a new project name (or press Enter to infer it from transcript): "
        ).strip()
        transcript_path = prompt_transcript_path()
        return {
            "selected_slug": "",
            "project_name_override": project_name,
            "transcript_path": transcript_path,
        }

    if not choice.isdigit():
        raise ValueError("Invalid selection. Please enter a valid project number or 'new'.")

    selection_idx = int(choice)
    if selection_idx < 1 or selection_idx > len(catalog):
        raise ValueError("Project selection out of range.")

    selected = catalog[selection_idx - 1]
    transcript_path = prompt_transcript_path(selected.get("source_transcript", ""))
    return {
        "selected_slug": selected["slug"],
        "project_name_override": selected["display_name"],
        "transcript_path": transcript_path,
    }

def get_updated_transcript_path(current_txt_path):
    confirm = input(
        '[agent] Do you want to update the BRD and card using terminal text input? (yes/no): '
    ).strip().lower()

    if confirm not in ("yes", "y"):
        return current_txt_path

    mode = input(
        '[agent] Type "transcript" to replace the transcript, or "suggestion" to apply change requests to the current transcript: '
    ).strip().lower()

    if mode == "transcript":
        updated_content = read_terminal_textbox('[agent] Paste the full updated transcript below.')
        if not updated_content:
            print("[agent] No updated text provided. Re-running with the existing transcript.")
            return current_txt_path
        return save_text_variant(current_txt_path, updated_content, "updated")

    if mode == "suggestion":
        suggestion_text = read_terminal_textbox('[agent] Paste the requested BRD/card changes below.')
        if not suggestion_text:
            print("[agent] No suggestions provided. Re-running with the existing transcript.")
            return current_txt_path

        current_transcript = current_txt_path.read_text(encoding="utf-8", errors="replace")
        combined_transcript = combine_transcript_and_suggestions(current_transcript, suggestion_text)
        return save_text_variant(current_txt_path, combined_transcript, "suggested")

    print("[agent] Unknown option. Re-running with the existing transcript.")
    return current_txt_path

def build_docx(overview, requirements, risks, assumptions,
               constraints, dependencies, source_filename, approval=None, related_folder="", related_file=""):

    doc  = Document()
    now  = datetime.now().strftime("%Y-%m-%d")
    proj = overview.get("PROJECT_NAME", "Project")

    # ── Page margins ───────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover Page ─────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    run.bold           = True
    run.font.size      = Pt(26)
    run.font.color.rgb = HPE_BLUE

    doc.add_paragraph()

    proj_para = doc.add_paragraph()
    proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = proj_para.add_run(proj)
    pr.font.size      = Pt(18)
    pr.font.color.rgb = HPE_BLUE

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Date: {now}\n"
        f"Source: {source_filename}\n"
        f"Related File: {related_file or 'N/A'}\n"
        f"Related Folder: {related_folder or 'N/A'}\n"
        f"Version: 1.0  |  Status: Draft"
    )

    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(overview.get("PURPOSE", "N/A"))
    doc.add_paragraph()

    # ── 2. Project Background ──────────────────────────────────────
    add_heading(doc, "2. Project Background", 1)
    doc.add_paragraph(overview.get("BACKGROUND", "N/A"))
    doc.add_paragraph()

    # ── 3. Scope ───────────────────────────────────────────────────
    add_heading(doc, "3. Scope", 1)
    add_heading(doc, "3.1 In Scope", 2)
    doc.add_paragraph(overview.get("SCOPE_IN", "N/A"))
    add_heading(doc, "3.2 Out of Scope", 2)
    doc.add_paragraph(overview.get("SCOPE_OUT", "N/A"))
    doc.add_paragraph()

    # ── 4. Stakeholders ────────────────────────────────────────────
    add_heading(doc, "4. Stakeholders", 1)
    stakeholders = overview.get("STAKEHOLDERS", "N/A").split(",")
    for s in stakeholders:
        s = s.strip()
        if s:
            doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph()

    # ── 5. Requirements ────────────────────────────────────────────
    add_heading(doc, "5. Business Requirements", 1)
    categories = {}
    for req in requirements:
        categories.setdefault(req["category"], []).append(req)

    for idx, (cat, reqs) in enumerate(categories.items(), 1):
        add_heading(doc, f"5.{idx} {cat} Requirements", 2)
        add_table(doc,
            ["ID", "Priority", "Description"],
            [[r["id"], r["priority"], r["description"]] for r in reqs]
        )
        doc.add_paragraph()

    # ── 6. Risks ───────────────────────────────────────────────────
    add_heading(doc, "6. Risks", 1)
    if risks:
        for i, r in enumerate(risks, 1):
            doc.add_paragraph(f"RISK-{i:03d}: {r}", style="List Bullet")
    else:
        doc.add_paragraph("No risks identified.")
    doc.add_paragraph()

    # ── 7. Assumptions ─────────────────────────────────────────────
    add_heading(doc, "7. Assumptions", 1)
    if assumptions:
        for i, a in enumerate(assumptions, 1):
            doc.add_paragraph(f"ASSUMPTION-{i:03d}: {a}", style="List Bullet")
    else:
        doc.add_paragraph("No assumptions identified.")
    doc.add_paragraph()

    # ── 8. Constraints ─────────────────────────────────────────────
    add_heading(doc, "8. Constraints", 1)
    if constraints:
        for i, c in enumerate(constraints, 1):
            doc.add_paragraph(f"CONSTRAINT-{i:03d}: {c}", style="List Bullet")
    else:
        doc.add_paragraph("No constraints identified.")
    doc.add_paragraph()

    # ── 9. Dependencies ───────────────────────────────────────────
    add_heading(doc, "9. Dependencies", 1)
    if dependencies:
        for i, d in enumerate(dependencies, 1):
            doc.add_paragraph(f"DEPENDENCY-{i:03d}: {d}", style="List Bullet")
    else:
        doc.add_paragraph("No dependencies identified.")
    doc.add_paragraph()

    # ── 10. Approval ───────────────────────────────────────────────
    approval = approval or build_pending_approval()
    add_heading(doc, "10. Approval", 1)
    add_table(doc, ["Field", "Details"], build_approval_rows(approval))
    doc.add_paragraph()

    return doc

def save_a2a_card(overview, requirements, risks, assumptions,
                  constraints, dependencies, stats,
                  transcript_path, docx_path, txt_path, cards_dir, project_dir):
    """Save A2A card as JSON for the Estimation Agent to consume.

    A new card is created for each BRD generation so Estimation can target
    a specific BRD/card pair from the project history.
    """

    project_name = overview.get("PROJECT_NAME", "Unknown").strip()
    project_slug = slugify_project_name(project_name)

    txt_stem = Path(txt_path).stem
    prefix = f"BRD_{project_slug}_"
    run_token = txt_stem[len(prefix):] if txt_stem.startswith(prefix) else datetime.now().strftime("%Y%m%d_%H%M%S")
    card_path = Path(cards_dir) / f"a2a_card_{project_slug}_{run_token}.json"

    new_payload = {
        "overview": overview,
        "requirements": requirements,
        "risks": risks,
        "assumptions": assumptions,
        "constraints": constraints,
        "dependencies": dependencies
    }

    approval = build_pending_approval()
    status = "pending_approval"

    # Build the updated card
    card = {
        "a2a_version": "1.0",
        "card_id": str(uuid.uuid4()),
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),

        "source_agent": {
            "name": "RequirementsAgent",
            "version": "1.0",
            "script": "Requirements.py"
        },

        "target_agent": {
            "name": "EstimationAgent",
            "version": "1.0",
            "script": "Estimation.py",
            "action": "estimate_from_approved_brd"
        },

        "project": {
            "name": project_name,
            "source_transcript": str(transcript_path),
            "processed_at": datetime.now().isoformat()
        },

        "artifacts": {
            "brd_docx": str(docx_path),
            "brd_txt":  str(txt_path),
            "output_dir": str(project_dir)
        },

        "payload": new_payload,

        "stats": {
            "total_requirements": stats["requirements"],
            "total_risks":        stats["risks"],
            "total_assumptions":  stats["assumptions"],
            "total_constraints":  stats["constraints"],
            "total_dependencies": stats["dependencies"]
        },

        "approval": approval,
        "status": status
    }

    with open(card_path, "w", encoding="utf-8") as f:
        json.dump(card, f, indent=2)

    print(f"[agent] A2A card created → {card_path.name}")
    return card_path

def combine_transcript_and_suggestions(transcript_text, suggestion_text):
    suggestion_text = (suggestion_text or "").strip()
    if not suggestion_text:
        return transcript_text
    return (
        f"{transcript_text.rstrip()}\n\n"
        "--- USER SUGGESTIONS / CHANGE REQUESTS ---\n"
        f"{suggestion_text}\n"
    )

def extract_related_artifacts_from_rag(rag_context: str) -> dict:
    """
    Extract most relevant file/folder from RAG context.
    Priority:
    1) first "> Path:" file hit
    2) first "- ..." related folder line
    3) RAG project root
    """
    result = {"related_file": "", "related_folder": "", "project_root": ""}
    if not rag_context or not rag_context.strip():
        return result

    lines = rag_context.split("\n")
    related_folders = []

    for line in lines:
        trimmed = line.strip()
        if not trimmed:
            continue

        if trimmed.startswith("> Path:"):
            full_path = trimmed.replace("> Path:", "", 1).strip()
            if full_path and len(full_path) > 5 and not result["related_file"]:
                result["related_file"] = full_path
                idx = max(full_path.rfind("\\"), full_path.rfind("/"))
                if idx > 0:
                    result["related_folder"] = full_path[:idx]

        if trimmed.startswith("- "):
            path = trimmed[2:].strip()
            if path and ("\\" in path or "/" in path) and path not in related_folders:
                related_folders.append(path)

        if "rag project root:" in trimmed.lower() and not result["project_root"]:
            key_idx = trimmed.lower().find("rag project root:")
            root = trimmed[key_idx + len("rag project root:"):].replace("*", "").strip()
            if root:
                result["project_root"] = root

    if not result["related_folder"] and related_folders:
        result["related_folder"] = related_folders[0]

    if not result["related_folder"] and result["project_root"]:
        result["related_folder"] = result["project_root"]

    return result

def run_requirements_pipeline(transcript_text, transcript_source_path, output_dir, project_name_override=""):
    if not transcript_text or not transcript_text.strip():
        raise ValueError("Transcript text is empty.")

    if not CLIENT_ID or not AUTH_TOKEN:
        raise RuntimeError("CLIENT_ID or AUTHENTICATION_TOKEN missing from .env")

    print("\n[agent] Setting up session ...")
    user_id, username, session_id = setup_session()

    # Retrieve RAG context and extract related file/folder from external project
    print("\n[agent] Retrieving RAG context for related files ...")
    rag_context = ground_prompt_with_rag(transcript_text, "BRD requirements extraction")
    rag_artifacts = extract_related_artifacts_from_rag(rag_context)
    related_file = rag_artifacts.get("related_file", "")
    related_folder = rag_artifacts.get("related_folder", "")
    project_root = rag_artifacts.get("project_root", "")
    if related_file:
        print(f"[agent] Related file from RAG: {related_file} ✅")
    if related_folder:
        print(f"[agent] Related folder from RAG: {related_folder} ✅")
    elif project_root:
        print(f"[agent] Using RAG project root: {project_root}")
    else:
        print("[agent] No related file/folder found in RAG context (using local folder)")
        related_folder = ""

    print("\n[agent] Step 1/3 — Extracting project overview ...")
    raw_overview = ask(
        PROMPT_OVERVIEW.format(transcript=transcript_text),
        user_id, username, session_id,
    )
    print("[agent] Overview response received ✅")

    print("\n[agent] Step 2/3 — Extracting requirements ...")
    raw_requirements = ask(
        PROMPT_REQUIREMENTS.format(transcript=transcript_text),
        user_id, username, session_id,
    )
    print("[agent] Requirements response received ✅")

    print("\n[agent] Step 3/3 — Extracting risks, assumptions, constraints ...")
    raw_risks = ask(
        PROMPT_RISKS.format(transcript=transcript_text),
        user_id, username, session_id,
    )
    print("[agent] Risks response received ✅")

    overview = parse_overview(raw_overview)
    requirements = parse_requirements(raw_requirements)
    risks, assumptions, constraints, deps = parse_risks(raw_risks)

    if project_name_override.strip():
        # Keeps updates for an existing/selected project in one folder.
        overview["PROJECT_NAME"] = project_name_override.strip()

    project_name = overview.get("PROJECT_NAME", "Unknown").strip()
    project_slug = slugify_project_name(project_name)

    # Master folder -> per-project folder -> cards/brd folders
    projects_root_dir = Path(output_dir) / "projects"
    project_dir = projects_root_dir / project_slug
    cards_dir = project_dir / "cards"
    brd_dir = project_dir / "brd"
    cards_dir.mkdir(parents=True, exist_ok=True)
    brd_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = project_slug

    docx_path = brd_dir / f"BRD_{base_name}_{timestamp}.docx"
    txt_out_path = brd_dir / f"BRD_{base_name}_{timestamp}.txt"

    # Use RAG-extracted folder/file if available, otherwise fallback to local brd_dir
    display_folder = related_folder or str(brd_dir)
    display_file = related_file or ""

    doc = build_docx(
        overview,
        requirements,
        risks,
        assumptions,
        constraints,
        deps,
        transcript_source_path.name,
        build_pending_approval(),
        related_folder=display_folder,
        related_file=display_file,
    )
    doc.save(str(docx_path))

    write_brd_txt(
        txt_out_path,
        overview,
        requirements,
        risks,
        assumptions,
        constraints,
        deps,
        transcript_source_path.name,
        build_pending_approval(),
        related_folder=display_folder,
        related_file=display_file,
    )

    stats = {
        "requirements": len(requirements),
        "risks": len(risks),
        "assumptions": len(assumptions),
        "constraints": len(constraints),
        "dependencies": len(deps),
    }

    card_path = save_a2a_card(
        overview,
        requirements,
        risks,
        assumptions,
        constraints,
        deps,
        stats,
        transcript_source_path,
        docx_path,
        txt_out_path,
        cards_dir,
        project_dir,
    )

    return {
        "overview": overview,
        "requirements": requirements,
        "risks": risks,
        "assumptions": assumptions,
        "constraints": constraints,
        "dependencies": deps,
        "stats": stats,
        "docx_path": str(docx_path),
        "txt_path": str(txt_out_path),
        "card_path": str(card_path),
        "project_dir": str(project_dir),
        "cards_dir": str(cards_dir),
        "brd_dir": display_folder,  # Use RAG-extracted folder if available
        "brd_file": display_file,
    }

def maybe_apply_terminal_suggestions(txt_path, transcript_text):
    use_suggestions = input(
        '[agent] Do you want to add change requests in a terminal text box before generating the BRD and card? (yes/no): '
    ).strip().lower()

    if use_suggestions not in ("yes", "y"):
        return txt_path, transcript_text

    suggestion_text = read_terminal_textbox('[agent] Paste the BRD/card change requests below.')
    if not suggestion_text:
        print("[agent] No suggestions provided. Proceeding with the original transcript.")
        return txt_path, transcript_text

    combined_transcript = combine_transcript_and_suggestions(transcript_text, suggestion_text)
    updated_txt_path = save_text_variant(txt_path, combined_transcript, "suggested")
    return updated_txt_path, combined_transcript
# ------------------------------------------------------------------ #
#  Main                                                                #
# ------------------------------------------------------------------ #

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Requirements Agent")
    parser.add_argument("transcript", nargs="?", help="Path to transcript .txt file")
    parser.add_argument("--approve-card", type=str, help="Approve an existing A2A card and trigger estimation")
    parser.add_argument("--approved-by", type=str, default="Business Owner", help="Approver name")
    parser.add_argument("--approver-role", type=str, default="Business Owner", help="Approver role")
    parser.add_argument("--approval-comments", type=str, default="", help="Optional approval comments")
    parser.add_argument("--project", type=str, help="Existing project slug from the projects folder")
    parser.add_argument("--project-name", type=str, default="", help="Override project name used for foldering")
    parser.add_argument("--list-projects", action="store_true", help="Print available projects as JSON and exit")
    args = parser.parse_args()

    if args.list_projects:
        print(json.dumps(get_project_catalog(SCRIPT_DIR), indent=2))
        raise SystemExit(0)

    if args.approve_card:
        approve_card(args.approve_card, args.approved_by, args.approver_role, args.approval_comments)
        raise SystemExit(0)

    selected_project_slug = ""
    project_name_override = args.project_name.strip()

    if args.project:
        catalog = get_project_catalog(SCRIPT_DIR)
        project_map = {p["slug"]: p for p in catalog}
        selected = project_map.get(args.project)
        if not selected:
            print(f"❌ Project slug not found: {args.project}")
            print("Use --list-projects to view available slugs.")
            raise SystemExit(1)
        selected_project_slug = selected["slug"]
        if not project_name_override:
            project_name_override = selected["display_name"]
        default_transcript = selected.get("source_transcript", "")
        txt_path = Path(args.transcript) if args.transcript else prompt_transcript_path(default_transcript)
    elif args.transcript:
        txt_path = Path(args.transcript)
    else:
        try:
            workflow_selection = choose_project_workflow(SCRIPT_DIR)
        except ValueError as e:
            print(f"❌ {e}")
            raise SystemExit(1)
        selected_project_slug = workflow_selection["selected_slug"]
        if workflow_selection["project_name_override"]:
            project_name_override = workflow_selection["project_name_override"]
        txt_path = workflow_selection["transcript_path"]

    # ── Validate file ─────────────────────────────────────────────
    if not txt_path.exists():
        print(f"❌ File not found: {txt_path}")
        sys.exit(1)

    if txt_path.suffix.lower() != ".txt":
        print(f"❌ File must be a .txt file, got: {txt_path.suffix}")
        sys.exit(1)

    print(f"✅ Reading transcript from: {txt_path.name}")
    transcript = txt_path.read_text(encoding="utf-8", errors="replace")

    if not transcript.strip():
        print("❌ Transcript file is empty.")
        sys.exit(1)

    txt_path, transcript = maybe_apply_terminal_suggestions(txt_path, transcript)

    print(f"✅ Transcript loaded — {len(transcript)} characters, "
          f"{len(transcript.splitlines())} lines")

    output_dir = txt_path.parent
    result = run_requirements_pipeline(transcript, txt_path, output_dir, project_name_override)
    overview = result["overview"]
    requirements = result["requirements"]
    risks = result["risks"]
    assumptions = result["assumptions"]
    constraints = result["constraints"]
    deps = result["dependencies"]
    docx_path = Path(result["docx_path"])
    txt_out_path = Path(result["txt_path"])
    a2a_card_path = Path(result["card_path"])
    project_dir = Path(result["project_dir"])
    cards_dir = Path(result["cards_dir"])
    brd_dir = Path(result["brd_dir"])

    print(f"\n[agent] Parsed:")
    print(f"  Overview fields  : {list(overview.keys())}")
    print(f"  Requirements     : {len(requirements)}")
    print(f"  Risks            : {len(risks)}")
    print(f"  Assumptions      : {len(assumptions)}")
    print(f"  Constraints      : {len(constraints)}")
    print(f"  Dependencies     : {len(deps)}")
    print(f"\n✅ DOCX saved : {docx_path}")

    print("\n[agent] BRD generated with approval block.")
    approval_input = input(
        '[agent] Review the BRD, then type "approved" to trigger estimation now, "not approved" to rerun requirements, or press Enter to leave it pending: '
    ).strip().lower()
    if approval_input == "approved":
        approver_name = input("[agent] Enter approver name: ").strip() or "Terminal Approval"
        approver_position = input("[agent] Enter approver position: ").strip() or "Interactive Review"
        approve_card(
            a2a_card_path,
            approved_by=approver_name,
            approver_role=approver_position,
            approval_comments="Approved in terminal after BRD generation.",
        )
    elif approval_input == "not approved":
        rerun_txt_path = get_updated_transcript_path(txt_path)
        print("[agent] BRD not approved. Re-triggering Requirements Agent...")
        rerun_cmd = [
            sys.executable,
            str(Path(__file__).resolve()),
            str(rerun_txt_path),
        ]
        if selected_project_slug:
            rerun_cmd.extend(["--project", selected_project_slug])
        elif project_name_override:
            rerun_cmd.extend(["--project-name", project_name_override])
        subprocess.run(
            rerun_cmd,
            check=True,
        )
        raise SystemExit(0)
    else:
        print("[agent] Approval not received. Card remains pending.")
        print(
            "[agent] You can approve later with: "
            f"python Requirements.py --approve-card \"{a2a_card_path}\" "
            "--approved-by \"Your Name\" --approver-role \"Business Owner\""
        )

    print(f"\n🎉 Done! All files saved under: {project_dir}")
    print(f"   📁 Cards Dir : {cards_dir}")
    print(f"   📁 BRD Dir   : {brd_dir}")
    print(f"   📄 BRD DOCX  : {docx_path.name}")
    print(f"   📄 BRD TXT   : {txt_out_path.name}")
    print(f"   🔗 A2A Card  : {a2a_card_path.name}")

