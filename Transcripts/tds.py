from dotenv import load_dotenv
import base64
import io
import json
import os
import re
import ssl
import time
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from rag_agent_context import ground_prompt_with_rag, get_rag_related_artifacts

from test_case_generator import generate_requirement_section


SCRIPT_DIR = Path(__file__).parent.resolve()
ENV_PATH = SCRIPT_DIR / ".env"
PEM_PATH = SCRIPT_DIR / "cacert 1 (1).pem"

load_dotenv(dotenv_path=ENV_PATH, override=True)

CLIENT_ID = os.getenv("CLIENT_ID")
USER_NAME = os.getenv("USER_NAME") or "TDS Agent"
AUTH_TOKEN = os.getenv("AUTHENTICATION_TOKEN")
API_URL = "https://api.chathpe.it.hpe.com/v2.8/"
MAX_RETRIES = 5

DEFAULT_TDS_SECTION_PLAN = [
    {
        "title": "System Architecture Diagram",
        "diagram_type": "flowchart TD",
        "focus": "End-to-end system components and integration flow from FDS.",
    },
    {
        "title": "Component Design and Sequences",
        "diagram_type": "sequenceDiagram",
        "focus": "Main request/approval lifecycle across key components from FDS.",
    },
    {
        "title": "Data Model and ERD",
        "diagram_type": "erDiagram",
        "focus": "Core entities and relationships required by FDS features.",
    },
    {
        "title": "API Contracts and NFR Specifications",
        "diagram_type": "flowchart TD",
        "focus": "API endpoints and linked NFR controls from FDS.",
    },
]

MANDATORY_TDS_TEXT_SECTIONS = [
    {
        "title": "Requirement Pseudocode",
        "diagram_type": "text",
        "focus": "Structured pseudocode for the implemented requirement flow based on FDS.",
    },
]


def _ensure_tds_text_sections(section_plan: list) -> list:
    plan = list(section_plan or [])
    existing_titles = {
        sanitize_line(str(section.get("title", "")).lower())
        for section in plan
        if isinstance(section, dict)
    }
    for section in MANDATORY_TDS_TEXT_SECTIONS:
        title_key = sanitize_line(section["title"].lower())
        if title_key not in existing_titles:
            plan.append(dict(section))
            existing_titles.add(title_key)
    return plan[:10]


def build_ssl_context():
    if not PEM_PATH.exists():
        print(f"[ssl] ❌ PEM not found at {PEM_PATH}")
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.load_verify_locations(cafile=str(PEM_PATH))
        print("[ssl] ✅ PEM loaded")
        return ctx
    except Exception as exc:
        print(f"[ssl] ❌ Failed to load PEM: {exc}")
        return None


SSL_CONTEXT = build_ssl_context()


def clean(val):
    return val.strip().strip('"').strip("'") if val else val


def auth_header():
    if not AUTH_TOKEN:
        raise RuntimeError("AUTHENTICATION_TOKEN missing from .env")
    token = clean(AUTH_TOKEN)
    if token.lower().startswith("bearer "):
        return token
    return f"Bearer {token}"


def post(endpoint, payload):
    url = API_URL + endpoint
    headers = {
        "Client-ID": clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Content-Type": "application/json",
        "Cache-Control": "no-cache",
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, headers=headers, data=data, method="POST")
    return _execute_request(req, decode_json=True)


def get(endpoint):
    url = API_URL + endpoint
    headers = {
        "Client-ID": clean(CLIENT_ID),
        "Authorization": auth_header(),
        "Cache-Control": "no-cache",
    }
    req = urllib.request.Request(url, headers=headers, method="GET")
    return _execute_request(req, decode_json=False)


def _extract_retry_seconds(body: str, default_wait: int) -> int:
    match = re.search(r"try again in\s+(\d+)\s+seconds?", body, flags=re.I)
    if match:
        try:
            return max(1, int(match.group(1)))
        except ValueError:
            return default_wait
    return default_wait


def _execute_request(req, decode_json: bool):
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            with urllib.request.urlopen(req, context=SSL_CONTEXT) as resp:
                payload = resp.read()
                if decode_json:
                    return json.loads(payload)
                return payload
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="replace")
            if exc.code == 429 and attempt < MAX_RETRIES:
                wait_seconds = _extract_retry_seconds(body, default_wait=min(2 * attempt, 10))
                print(f"[retry] HTTP 429 received; waiting {wait_seconds}s before retry {attempt + 1}/{MAX_RETRIES} ...")
                time.sleep(wait_seconds)
                continue

            print(f"[error] HTTP {exc.code} — {body}")
            raise


def setup_session():
    print("[session] Logging in ...")
    resp = post("login", {"appId": "1"})
    bot_data = resp["chatHPE_bot_data"]
    user_id = bot_data["userId"]
    username = bot_data["username"]
    print(f"[session] Logged in as {username}")

    post(
        "preferences",
        {
            "agreement": True,
            "dark_mode": True,
            "stream": False,
            "chatHPE_bot_data": {
                "appId": "1",
                "sessionId": "-1",
                "userId": user_id,
                "username": username,
            },
            "webScraping": False,
        },
    )

    raw = get("sessionId_generator")
    session_id = str(raw.split()[2])[2:-2]
    print(f"[session] session_id={session_id}")
    return user_id, username, session_id


def ask(prompt, user_id, username, session_id):
    grounded_prompt = ground_prompt_with_rag(
        prompt,
        task_label="TDS generation from FDS",
    )
    endpoint = (
        "call/chatlite"
        "?force_async=false"
        "&session_management_support=true"
        "&internal_call=false"
        "&proxy=false"
    )
    resp = post(
        endpoint,
        {
            "chatHPE_bot_data": {
                "appId": "1",
                "sessionId": session_id,
                "userId": user_id,
                "username": username,
            },
            "model_name": "gpt-4o-mini",
            "stream": False,
            "webScraping": False,
            "user_query": grounded_prompt,
        },
    )
    return resp.get("Response", "")


PROMPT_ARCHITECTURE = """You are the TDS Agent.
Generate ONLY one Mermaid flowchart for the section "System Architecture Diagram".

Must represent the target application architecture described by the FDS output.
Do NOT describe any document-generation workflow, and do NOT include agent names.

Rules:
- Output ONLY Mermaid code, no prose, no markdown fences.
- First line must be exactly: flowchart TD
- Use short business-readable node labels.
- Use FDS Functional Requirements, Integration Points, Data Requirements, and NFR details to decide nodes/flows.
- Include client/channel, API/gateway, core service(s), data store(s), and external integration(s) when present in FDS.
- Include at least one decision node for a business rule or approval path from FDS.
- Include at least one control node (auth/validation/audit/monitoring) inferred from FDS.

Project: {project_name}

FDS Content:
{fds_context}
"""

PROMPT_COMPONENTS = """You are the TDS Agent.
Generate ONLY one Mermaid sequence diagram for the section "Component Design and Sequences".

Rules:
- Output ONLY Mermaid code, no prose, no markdown fences.
- First line must be exactly: sequenceDiagram
- Include actors/components for User, API, Service, Database, Notification.
- Cover submit request, validation, approval/rejection, persistence, and notification.

Project: {project_name}

FDS Content:
{fds_context}
"""

PROMPT_DATA_MODEL = """You are the TDS Agent.
Generate ONLY one Mermaid ER diagram for the section "Data Model and ERD".

Rules:
- Output ONLY Mermaid code, no prose, no markdown fences.
- First line must be exactly: erDiagram
- Include core entities and relationships inferred from FDS.
- Include primary/foreign key hints in attributes where possible.

Project: {project_name}

FDS Content:
{fds_context}
"""

PROMPT_API_NFR = """You are the TDS Agent.
Generate ONLY one Mermaid flowchart for the section "API Contracts and NFR Specifications".

Rules:
- Output ONLY Mermaid code, no prose, no markdown fences.
- First line must be exactly: flowchart TD
- Include API endpoints as nodes (submit/approve/reject/get status).
- Include NFR nodes: Security, Performance, Scalability, Observability, Reliability.
- Show traceable links from APIs to NFR controls.

Project: {project_name}

FDS Content:
{fds_context}
"""


PROMPT_DYNAMIC_SECTION_PLAN = """You are the TDS Agent.
Create a dynamic TDS section plan from BRD/FDS context, based on actual requirement scope.

Rules:
- Output ONLY valid JSON. No markdown fences and no extra text.
- JSON schema:
    {{
        "sections": [
            {{
                "title": "<section title>",
                "diagram_type": "flowchart TD | sequenceDiagram | erDiagram",
                "focus": "<what this section must cover based on requirements>"
            }}
        ]
    }}
- Do NOT include agent/document-generation workflow sections.
- Section titles must reflect requirement domains in FDS (modules, integrations, data, operations, security, etc.).
- Choose section count dynamically from requirements complexity: minimum 4, maximum 10.
- Keep titles concise and business readable.

Project: {project_name}

FDS Content:
{fds_context}
"""


PROMPT_DYNAMIC_SECTION_DIAGRAM = """You are the TDS Agent.
Generate ONLY one Mermaid diagram for this TDS section.

Section Title: {section_title}
Section Focus: {section_focus}
Required Diagram Type: {diagram_type}

Rules:
- Output ONLY Mermaid code, no prose, no markdown fences.
- First line must be exactly: {diagram_type}
- Use only FDS/BRD-aligned system content.
- Do NOT include any agent/tool/pipeline nodes.
- Keep labels concise and business-readable.

Project: {project_name}

FDS Content:
{fds_context}
"""


def _strip_mermaid_fences(text: str) -> str:
    fenced = re.search(r"```(?:mermaid)?\s*(.*?)```", text, flags=re.S | re.I)
    if fenced:
        return fenced.group(1).strip()
    return text.strip()


def _extract_mermaid(text: str, expected_start: str, fallback: str) -> str:
    cleaned = _strip_mermaid_fences(text)
    lines = [line.rstrip() for line in cleaned.splitlines() if line.strip()]
    if not lines:
        return fallback

    start_index = None
    for idx, line in enumerate(lines):
        if line.strip().lower().startswith(expected_start.lower()):
            start_index = idx
            break

    if start_index is None:
        return fallback

    diagram_lines = lines[start_index:]
    if len(diagram_lines) < 2:
        return fallback

    return "\n".join(diagram_lines)


def _extract_json_object(text: str):
    cleaned = _strip_mermaid_fences(text)
    try:
        return json.loads(cleaned)
    except Exception:
        pass

    match = re.search(r"\{[\s\S]*\}", cleaned)
    if not match:
        return None

    try:
        return json.loads(match.group(0))
    except Exception:
        return None


ARCH_FALLBACK = """flowchart TD
    U[User Channel] --> API[API Gateway]
    API --> V[Validate Request]
    V --> D{Business Rule Check}
    D -->|Pass| S[Core Service]
    D -->|Fail| R[Reject Response]
    S --> DB[(Primary Database)]
    S --> INT[External Integration]
    S --> AUD[Audit and Monitoring]
    DB --> Q[Query and Status API]
"""


COMPONENT_SEQ_FALLBACK = """sequenceDiagram
    actor U as User
    participant API as API Gateway
    participant S as Leave Service
    participant DB as Leave DB
    participant N as Notification Service
    U->>API: Submit leave request
    API->>S: Validate payload
    S->>DB: Persist pending request
    S->>N: Notify approver
    API-->>U: Request accepted
    U->>API: Approve/Reject decision
    API->>S: Process decision
    S->>DB: Update request status
    S->>N: Notify requester
    API-->>U: Final status
"""


ERD_FALLBACK = """erDiagram
    EMPLOYEE {
        int employee_id PK
        string name
        string email
    }
    LEAVE_REQUEST {
        int request_id PK
        int employee_id FK
        date start_date
        date end_date
        string status
    }
    APPROVAL {
        int approval_id PK
        int request_id FK
        string decision
        datetime decided_at
    }
    AUDIT_LOG {
        int audit_id PK
        int request_id FK
        string action
        datetime created_at
    }
    EMPLOYEE ||--o{ LEAVE_REQUEST : submits
    LEAVE_REQUEST ||--o{ APPROVAL : has
    LEAVE_REQUEST ||--o{ AUDIT_LOG : tracks
"""


API_NFR_FALLBACK = """flowchart TD
    A1[POST /leave-requests]
    A2[PUT /leave-requests/{id}/approve]
    A3[PUT /leave-requests/{id}/reject]
    A4[GET /leave-requests/{id}]
    N1[Security: JWT + RBAC + TLS]
    N2[Performance: p95 latency + throughput]
    N3[Scalability: horizontal scaling]
    N4[Observability: logs metrics traces]
    N5[Reliability: retries + idempotency]
    A1 --> N1
    A1 --> N2
    A2 --> N1
    A2 --> N5
    A3 --> N1
    A3 --> N5
    A4 --> N2
    A4 --> N4
    A1 --> N3
"""


def _fallback_for_diagram_type(diagram_type: str) -> str:
    normalized = (diagram_type or "").strip().lower()
    if normalized == "text":
        return ""
    if normalized.startswith("sequencediagram"):
        return COMPONENT_SEQ_FALLBACK
    if normalized.startswith("erdiagram"):
        return ERD_FALLBACK
    return ARCH_FALLBACK


def _normalize_diagram_type(diagram_type: str) -> str:
    normalized = (diagram_type or "").strip().lower()
    if normalized == "text":
        return "text"
    if normalized.startswith("sequencediagram"):
        return "sequenceDiagram"
    if normalized.startswith("erdiagram"):
        return "erDiagram"
    return "flowchart TD"


def derive_tds_section_plan(project_name: str, fds_context: str, user_id, username, session_id):
    raw = ask(
        PROMPT_DYNAMIC_SECTION_PLAN.format(
            project_name=project_name,
            fds_context=fds_context,
        ),
        user_id,
        username,
        session_id,
    )

    parsed = _extract_json_object(raw)
    if not isinstance(parsed, dict):
        return _ensure_tds_text_sections(DEFAULT_TDS_SECTION_PLAN)

    sections = parsed.get("sections")
    if not isinstance(sections, list) or not sections:
        return _ensure_tds_text_sections(DEFAULT_TDS_SECTION_PLAN)

    plan = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = sanitize_line(str(section.get("title", "")))
        if not title:
            continue
        diagram_type = _normalize_diagram_type(str(section.get("diagram_type", "flowchart TD")))
        focus = sanitize_line(str(section.get("focus", "")))
        if not focus:
            focus = f"Generate {title} from FDS requirements."
        plan.append(
            {
                "title": title,
                "diagram_type": diagram_type,
                "focus": focus,
            }
        )

    if not plan:
        return _ensure_tds_text_sections(DEFAULT_TDS_SECTION_PLAN)

    return _ensure_tds_text_sections(plan)


def sanitize_line(text: str) -> str:
    if not text:
        return ""
    return " ".join(text.strip().strip("`").split())


def sanitize_multiline_text(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = []
    blank_pending = False

    for line in lines:
        if not line.strip():
            if cleaned and not blank_pending:
                cleaned.append("")
            blank_pending = True
            continue
        cleaned.append(line.rstrip())
        blank_pending = False

    return "\n".join(cleaned).strip()


def load_tds_a2a_card(card_path: Path) -> dict:
    if card_path.suffix.lower() != ".json":
        raise ValueError("❌ TDS agent only accepts .json a2a cards")

    try:
        with open(card_path, "r", encoding="utf-8") as handle:
            card = json.load(handle)
    except json.JSONDecodeError as exc:
        raise ValueError(f"❌ Invalid JSON file: {exc}") from exc

    if not isinstance(card, dict):
        raise ValueError("❌ Invalid a2a card: root must be a JSON object")

    allowed_statuses = {"ready_for_tds", "approved_for_tds", "tds_approved"}
    if card.get("status") not in allowed_statuses:
        raise ValueError("❌ Invalid a2a card: status must be one of 'ready_for_tds', 'approved_for_tds', or 'tds_approved'")

    payload = card.get("payload")
    if not isinstance(payload, dict):
        raise ValueError("❌ Invalid a2a card: payload must be an object")

    fds_sections = payload.get("fds_sections")
    if not isinstance(fds_sections, dict) or not fds_sections:
        raise ValueError("❌ Invalid a2a card: payload.fds_sections must be a non-empty object")

    return card


def _load_json_file(path: Path):
    try:
        with open(path, "r", encoding="utf-8") as handle:
            return json.load(handle)
    except Exception:
        return None


def extract_brd_requirements(card: dict, card_path: Path) -> list:
    payload = card.get("payload", {}) if isinstance(card, dict) else {}
    requirements = payload.get("requirements", []) if isinstance(payload, dict) else []

    if isinstance(requirements, list) and requirements:
        return [req for req in requirements if isinstance(req, dict)]

    artifacts = card.get("artifacts", {}) if isinstance(card, dict) else {}
    linked_cards = []
    estimation_card = artifacts.get("estimation_card")
    if estimation_card:
        linked_cards.append(Path(estimation_card))

    input_card = artifacts.get("input_card")
    if input_card:
        linked_cards.append(Path(input_card))

    for linked in linked_cards:
        resolved = linked if linked.is_absolute() else (card_path.parent / linked).resolve()
        source = _load_json_file(resolved)
        if not isinstance(source, dict):
            continue
        source_payload = source.get("payload", {})
        source_requirements = source_payload.get("requirements", []) if isinstance(source_payload, dict) else []
        if isinstance(source_requirements, list) and source_requirements:
            return [req for req in source_requirements if isinstance(req, dict)]

    return []


def build_fds_context(card: dict) -> str:
    payload = card.get("payload", {})
    sections = payload.get("fds_sections", {})
    summary = payload.get("summary", {})

    ordered_lines = []
    for key, value in sections.items():
        ordered_lines.append(f"## {key}\n{value}")

    if summary:
        ordered_lines.append("## FDS Summary")
        for key, value in summary.items():
            ordered_lines.append(f"{key}: {value}")

    raw_fds = payload.get("raw_fds", "")
    if raw_fds:
        ordered_lines.append("## Raw FDS")
        ordered_lines.append(raw_fds)

    return "\n\n".join(ordered_lines).strip()


def get_project_name(card: dict) -> str:
    project = card.get("project", {})
    return project.get("name") or "Project"


def resolve_output_dir(card_path: Path, card: dict) -> Path:
    artifacts = card.get("artifacts", {})
    output_dir = (artifacts.get("output_dir") or "").strip()

    # Prefer project-scoped storage: project/tds.
    # If the stored output_dir points to cards, move one level up to project root.
    if output_dir:
        candidate = Path(output_dir)
        if not candidate.is_absolute():
            candidate = (card_path.parent / candidate).resolve()
        if candidate.name.lower() == "cards":
            project_dir = candidate.parent
        else:
            project_dir = candidate
    else:
        project_dir = card_path.parent.parent if card_path.parent.name.lower() == "cards" else card_path.parent

    return (project_dir / "tds").resolve()


def run_tds_from_card(card: dict, card_path: Path = None) -> dict:
    project_name = get_project_name(card)
    fds_context = build_fds_context(card)
    requirements = extract_brd_requirements(card, card_path or Path.cwd())

    rag_seed = f"Project: {project_name}\n\n{fds_context}"
    rag_artifacts = get_rag_related_artifacts(rag_seed, "TDS generation from FDS")
    related_file = rag_artifacts.get("related_file", "")
    related_folder = rag_artifacts.get("related_folder", "")

    user_id, username, session_id = setup_session()

    section_plan = derive_tds_section_plan(project_name, fds_context, user_id, username, session_id)
    total_sections = len(section_plan)

    sections = {}
    section_order = []
    for idx, section in enumerate(section_plan, start=1):
        section_name = section["title"]
        diagram_type = section["diagram_type"]
        section_focus = section["focus"]
        section_order.append(section_name)

        print(f"[tds] Step {idx}/{total_sections} — Generating {section_name} ...")

        # Text sections (pseudocode / test cases) are handled by test_case_generator
        if diagram_type == "text":
            ask_fn = lambda prompt: ask(prompt, user_id, username, session_id)
            sections[section_name] = generate_requirement_section(
                section_name, project_name, fds_context, ask_fn
            )
            continue

        prompt = PROMPT_DYNAMIC_SECTION_DIAGRAM.format(
            section_title=section_name,
            section_focus=section_focus,
            diagram_type=diagram_type,
            project_name=project_name,
            fds_context=fds_context,
        )
        raw_text = ask(prompt, user_id, username, session_id)
        fallback_diagram = _fallback_for_diagram_type(diagram_type)
        sections[section_name] = _extract_mermaid(raw_text, diagram_type, fallback_diagram)

    return {
        "project_name": project_name,
        "requirements": requirements,
        "sections": sections,
        "section_order": section_order,
        "section_plan": section_plan,
        "related_file": related_file,
        "related_folder": related_folder,
    }


HPE_BLUE = RGBColor(0x1F, 0x49, 0x7D)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = HPE_BLUE
    return para


def add_body_paragraph(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return para


def add_preformatted_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def _strip_markdown_fences(text: str) -> str:
    cleaned = (text or "").strip()
    fenced = re.search(r"```(?:[A-Za-z0-9_+-]+)?\s*([\s\S]*?)```", cleaned)
    if fenced:
        return fenced.group(1).strip()
    return cleaned


def _is_code_focused_section(section_name: str) -> bool:
    name = (section_name or "").strip().lower()
    return (
        "pseudocode" in name
        or "junit" in name
        or "test skeleton" in name
        or "sample code" in name
    )


def _normalize_code_text(text: str) -> str:
    lines = _strip_markdown_fences(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    normalized = []
    blank_pending = False
    for line in lines:
        cleaned = line.rstrip()
        if not cleaned.strip():
            if normalized and not blank_pending:
                normalized.append("")
            blank_pending = True
            continue
        normalized.append(cleaned)
        blank_pending = False
    return "\n".join(normalized).strip()


def add_strict_code_block(doc, text: str):
    """Render generated code/pseudocode as a single strict preformatted block."""
    normalized = _normalize_code_text(text)
    if not normalized:
        add_body_paragraph(doc, "Not generated.")
        return
    add_preformatted_block(doc, normalized)


def is_mermaid_diagram(content: str) -> bool:
    stripped = (content or "").strip()
    return (
        stripped.startswith("flowchart TD")
        or stripped.startswith("graph TD")
        or stripped.startswith("graph LR")
        or stripped.startswith("graph RL")
        or stripped.startswith("graph BT")
        or stripped.startswith("sequenceDiagram")
        or stripped.startswith("erDiagram")
    )


def extract_mermaid_block(content: str) -> str:
    lines = (content or "").splitlines()
    start_index = None
    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.lower() == "mermaid":
            start_index = index + 1
            break
        if stripped.startswith((
            "flowchart TD",
            "graph TD",
            "graph LR",
            "graph RL",
            "graph BT",
            "sequenceDiagram",
            "erDiagram",
        )):
            start_index = index
            break

    if start_index is None:
        return ""

    diagram_lines = []
    for line in lines[start_index:]:
        stripped = line.strip()
        if not stripped:
            if diagram_lines:
                break
            continue

        looks_like_diagram = (
            stripped.startswith((
                "flowchart TD",
                "graph TD",
                "graph LR",
                "graph RL",
                "graph BT",
                "sequenceDiagram",
                "erDiagram",
                "subgraph",
                "participant ",
                "actor ",
                "classDef ",
                "class ",
                "style ",
                "linkStyle ",
                "%%",
                "end",
            ))
            or "-->" in stripped
            or "->>" in stripped
            or "-->>" in stripped
            or "||--" in stripped
            or "}|" in stripped
            or re.match(r"^[A-Za-z0-9_]+\s*\[.*\]", stripped)
            or re.match(r"^[A-Za-z0-9_]+\s*\{\{.*\}\}", stripped)
        )
        if not looks_like_diagram and diagram_lines:
            break

        diagram_lines.append(line.rstrip())

    candidate = "\n".join(diagram_lines).strip()
    return candidate if is_mermaid_diagram(candidate) else ""


def _resolve_mmdc_command():
    """Return path to mmdc CLI if available (env var, local workspace, then PATH)."""
    configured = os.getenv("MERMAID_MMDC_PATH")
    if configured and Path(configured).exists():
        return configured

    local_candidates = [
        SCRIPT_DIR / "node_modules" / ".bin" / "mmdc.cmd",
        SCRIPT_DIR / "node_modules" / ".bin" / "mmdc",
        SCRIPT_DIR / "mmdc.cmd",
        SCRIPT_DIR / "mmdc",
    ]
    for candidate in local_candidates:
        if candidate.exists():
            return str(candidate)

    return shutil.which("mmdc.cmd") or shutil.which("mmdc")


def _sanitize_mermaid_for_render(diagram_text: str) -> str:
    """Remove lines that commonly trigger Mermaid parse failures while preserving intent."""
    lines = (diagram_text or "").splitlines()
    if not lines:
        return ""

    sanitized = []
    for index, line in enumerate(lines):
        stripped = line.strip()

        if index == 0:
            sanitized.append(stripped)
            continue

        if not stripped:
            sanitized.append("")
            continue

        # Mermaid fails on dangling connectors (for example: "A -->" with no target).
        if re.search(r"(-->|->>|-->>|\|\|--|\}\|)\s*$", stripped):
            continue

        sanitized.append(line.rstrip())

    while sanitized and not sanitized[-1].strip():
        sanitized.pop()

    return "\n".join(sanitized).strip()


def _render_mermaid_png_remote(diagram_text: str) -> bytes:
    import ssl
    encoded = base64.urlsafe_b64encode(diagram_text.encode("utf-8")).decode("ascii")
    # mermaid.ink renders Mermaid syntax directly to PNG for embedding in documents.
    url = f"https://mermaid.ink/img/{encoded}?bgColor=white"
    req = urllib.request.Request(
        url,
        method="GET",
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
            "Accept": "image/png,image/*;q=0.9,*/*;q=0.8",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return resp.read()
    except Exception:
        # Fallback for restrictive enterprise TLS middleboxes.
        with urllib.request.urlopen(req, timeout=30, context=ssl._create_unverified_context()) as resp:
            return resp.read()


def _render_mermaid_png_local(diagram_text: str):
    """Render diagram via local mmdc CLI. Returns bytes or None on failure."""
    mmdc = _resolve_mmdc_command()
    if not mmdc:
        return None

    temp_dir = tempfile.mkdtemp(prefix="tds_mermaid_")
    input_path = Path(temp_dir) / "diagram.mmd"
    output_path = Path(temp_dir) / "diagram.png"

    try:
        input_path.write_text(diagram_text, encoding="utf-8")
        cmd = [mmdc, "-i", str(input_path), "-o", str(output_path), "-b", "white", "-w", "1600"]
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        return output_path.read_bytes()
    except Exception as exc:
        print(f"[warn] Local Mermaid render failed: {exc}")
        return None
    finally:
        try:
            if input_path.exists():
                input_path.unlink()
            if output_path.exists():
                output_path.unlink()
            Path(temp_dir).rmdir()
        except Exception:
            pass


def render_mermaid_png(diagram_text: str) -> bytes:
    # Try local mmdc first (faster, no network dependency)
    local_png = _render_mermaid_png_local(diagram_text)
    if local_png is not None:
        return local_png

    try:
        return _render_mermaid_png_remote(diagram_text)
    except urllib.error.HTTPError as exc:
        # Retry once after light sanitization for malformed Mermaid output.
        if exc.code == 400:
            sanitized = _sanitize_mermaid_for_render(diagram_text)
            if sanitized and sanitized != (diagram_text or "").strip():
                local_retry = _render_mermaid_png_local(sanitized)
                if local_retry is not None:
                    return local_retry
                return _render_mermaid_png_remote(sanitized)
        raise


def add_mermaid_image(doc, diagram_text: str):
    png_bytes = render_mermaid_png(diagram_text)
    image_stream = io.BytesIO(png_bytes)
    doc.add_picture(image_stream, width=Inches(6.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_content(doc, section_name: str, content: str):
    text = (content or "").strip()

    if _is_code_focused_section(section_name):
        add_strict_code_block(doc, text)
        return

    mermaid_block = text if is_mermaid_diagram(text) else extract_mermaid_block(text)
    if mermaid_block:
        try:
            add_mermaid_image(doc, mermaid_block)
            prose_only = text.replace(mermaid_block, "").strip()
            if prose_only:
                text = prose_only
            else:
                return
        except Exception as exc:
            add_body_paragraph(
                doc,
                "Diagram render failed, keeping Mermaid source. "
                "Allow access to mermaid.ink for remote image rendering. "
                f"Error: {exc}",
            )

    lines = text.splitlines()
    buffer = []
    in_code_block = False

    def flush_buffer():
        nonlocal buffer
        if buffer:
            add_body_paragraph(doc, "\n".join(buffer).strip())
            buffer = []

    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            flush_buffer()
            continue

        if stripped.startswith("flowchart TD") or stripped.startswith("sequenceDiagram") or stripped.startswith("erDiagram"):
            flush_buffer()
            in_code_block = True
            add_preformatted_block(doc, stripped)
            continue

        if in_code_block:
            if stripped.startswith(("1.", "2.", "3.", "4.")):
                in_code_block = False
                buffer.append(stripped)
            elif ":" in stripped or stripped.startswith(("participant ", "actor ", "%%", "classDef ", "class ", "A", "B", "C", "D", "E")) or "-->" in stripped or "||--" in stripped or "{" in stripped:
                add_preformatted_block(doc, stripped)
            else:
                in_code_block = False
                buffer.append(stripped)
            continue

        if stripped.startswith(("1.", "2.", "3.", "4.")):
            flush_buffer()
            add_heading(doc, stripped, level=2)
            continue

        if stripped.startswith(("- ", "* ")):
            flush_buffer()
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        buffer.append(stripped)

    flush_buffer()


def add_requirements_table(doc, requirements: list):
    if not requirements:
        add_body_paragraph(doc, "No BRD requirements were found in the source card payload.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Requirement ID", "Category", "Priority", "Requirement (Verbatim from BRD)"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        header_run = table.rows[0].cells[index].paragraphs[0].runs[0]
        header_run.bold = True

    for req in requirements:
        req_id = str(req.get("id", "")).strip()
        category = str(req.get("category", "")).strip()
        priority = str(req.get("priority", "")).strip()
        description = str(req.get("description", "")).strip()

        if not (req_id or category or priority or description):
            continue

        row = table.add_row().cells
        row[0].text = req_id
        row[1].text = category
        row[2].text = priority
        row[3].text = description


def build_tds_docx(
    project_name: str,
    requirements: list,
    sections: dict,
    section_order: list,
    source_card_name: str,
    source_fds_docx: str,
    related_folder: str = "",
    related_file: str = "",
):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading("Technical Design Specification (TDS)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(project_name).font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Date: {datetime.now().strftime('%Y-%m-%d')}\n"
        f"Prepared By: {USER_NAME}\n"
        f"Source A2A Card: {source_card_name}\n"
        f"Source FDS: {source_fds_docx}\n"
        f"Related File: {related_file or 'N/A'}\n"
        f"Related Folder: {related_folder or 'N/A'}"
    )

    doc.add_page_break()

    add_heading(doc, "1. Document Control", level=1)
    add_body_paragraph(
        doc,
        "Version: 1.0\n"
        "Status: Draft\n"
        f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Generated From: {source_card_name}\n"
        f"Related File: {related_file or 'N/A'}\n"
        f"Related Folder: {related_folder or 'N/A'}",
    )

    add_heading(doc, "2. BRD Requirements (Verbatim)", level=1)
    add_requirements_table(doc, requirements)

    for index, section_name in enumerate(section_order, start=3):
        add_heading(doc, f"{index}. {section_name}", level=1)
        add_section_content(doc, section_name, sections.get(section_name, "Not generated."))

    return doc


def main():
    import sys

    if not CLIENT_ID or not AUTH_TOKEN:
        print("❌ CLIENT_ID or AUTHENTICATION_TOKEN missing from .env")
        sys.exit(1)

    if len(sys.argv) < 2:
        script_name = Path(sys.argv[0]).name
        print(f"❌ Usage: python {script_name} <tds_a2a_card.json>")
        sys.exit(1)

    card_path = Path(sys.argv[1]).expanduser()
    if not card_path.is_absolute():
        card_path = (Path.cwd() / card_path).resolve()

    if not card_path.exists():
        print(f"❌ File not found: {card_path}")
        sys.exit(1)

    try:
        card = load_tds_a2a_card(card_path)
    except ValueError as exc:
        print(str(exc))
        sys.exit(1)

    result = run_tds_from_card(card, card_path)
    output_dir = resolve_output_dir(card_path, card)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"TDS_{timestamp}.docx"
    json_path = output_dir / f"TDS_{timestamp}.json"

    source_fds_docx = card.get("artifacts", {}).get("fds_docx", "Not specified")
    display_folder = result.get("related_folder") or str(output_dir)
    display_file = result.get("related_file") or ""
    doc = build_tds_docx(
        result["project_name"],
        result.get("requirements", []),
        result["sections"],
        result.get("section_order", list(result["sections"].keys())),
        card_path.name,
        source_fds_docx,
        related_folder=display_folder,
        related_file=display_file,
    )
    doc.save(str(docx_path))

    with open(json_path, "w", encoding="utf-8") as handle:
        json.dump(
            {
                "project": result["project_name"],
                "generated_at": datetime.now().isoformat(),
                "source_card": str(card_path),
                "related_file": display_file,
                "related_folder": display_folder,
                "requirements": result.get("requirements", []),
                "section_plan": result.get("section_plan", []),
                "sections": result["sections"],
            },
            handle,
            indent=2,
        )

    print(f"\n✅ TDS DOCX saved: {docx_path}")
    print(f"✅ TDS JSON saved: {json_path}")


if __name__ == "__main__":
    main()