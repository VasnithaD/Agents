"""
Test Case Generator
-------------------
Standalone module responsible for generating requirement pseudocode and
test cases (Low / Medium / Extreme) from FDS context.

Public surface
~~~~~~~~~~~~~~
- enforce_mandatory_requirement_sections(plan) → list
- generate_requirement_section(section_name, project_name, fds_context, ask_fn) → str
- run_test_cases_from_card(card) → dict          # full standalone pipeline
- build_test_cases_docx(...) → docx.Document     # DOCX builder

Designed to be imported by tds.py (for section enforcement) and by app.py
(for the standalone Test Cases endpoint). Has NO dependency on tds.py so
there is no circular import.
"""

import io
import json
import os
import re
import time
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).parent.resolve()
ENV_PATH = SCRIPT_DIR / ".env"
PEM_PATH = SCRIPT_DIR / "cacert 1 (1).pem"

load_dotenv(dotenv_path=ENV_PATH, override=True)

_CLIENT_ID = os.getenv("CLIENT_ID")
_USER_NAME = os.getenv("USER_NAME") or "TDS Agent"
_AUTH_TOKEN = os.getenv("AUTHENTICATION_TOKEN")
_API_URL = "https://api.chathpe.it.hpe.com/v2.8/"
_MAX_RETRIES = 5


# ── Internal LLM helpers (duplicate of minimal session code from tds.py) ─────

def _build_ssl_context():
    import ssl
    if not PEM_PATH.exists():
        return None
    try:
        ctx = ssl.create_default_context()
        ctx.load_verify_locations(cafile=str(PEM_PATH))
        return ctx
    except Exception:
        return None


_SSL_CONTEXT = _build_ssl_context()


def _clean(val):
    return val.strip().strip('"').strip("'") if val else val


def _auth_header():
    if not _AUTH_TOKEN:
        raise RuntimeError("AUTHENTICATION_TOKEN missing from .env")
    token = _clean(_AUTH_TOKEN)
    return token if token.lower().startswith("bearer ") else f"Bearer {token}"


def _execute_request(req, decode_json: bool):
    for attempt in range(1, _MAX_RETRIES + 1):
        try:
            with urllib.request.urlopen(req, context=_SSL_CONTEXT) as resp:
                payload = resp.read()
                return json.loads(payload) if decode_json else payload
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="replace")
            if exc.code == 429 and attempt < _MAX_RETRIES:
                match = re.search(r"try again in\s+(\d+)\s+seconds?", body, re.I)
                wait = max(1, int(match.group(1))) if match else min(2 * attempt, 10)
                print(f"[retry] 429 — waiting {wait}s before retry {attempt + 1}…")
                time.sleep(wait)
                continue
            print(f"[error] HTTP {exc.code} — {body}")
            raise


def _post(endpoint, payload):
    url = _API_URL + endpoint
    headers = {
        "Client-ID": _clean(_CLIENT_ID),
        "Authorization": _auth_header(),
        "Content-Type": "application/json",
        "Cache-Control": "no-cache",
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, headers=headers, data=data, method="POST")
    return _execute_request(req, decode_json=True)


def _get(endpoint):
    url = _API_URL + endpoint
    headers = {
        "Client-ID": _clean(_CLIENT_ID),
        "Authorization": _auth_header(),
        "Cache-Control": "no-cache",
    }
    req = urllib.request.Request(url, headers=headers, method="GET")
    return _execute_request(req, decode_json=False)


def _setup_session():
    resp = _post("login", {"appId": "1"})
    bot = resp["chatHPE_bot_data"]
    user_id, username = bot["userId"], bot["username"]
    _post("preferences", {
        "agreement": True, "dark_mode": True, "stream": False,
        "chatHPE_bot_data": {"appId": "1", "sessionId": "-1",
                             "userId": user_id, "username": username},
        "webScraping": False,
    })
    raw = _get("sessionId_generator")
    session_id = str(raw.split()[2])[2:-2]
    print(f"[tc-session] session_id={session_id}")
    return user_id, username, session_id


def _ask(prompt, user_id, username, session_id):
    endpoint = (
        "call/chatlite"
        "?force_async=false&session_management_support=true"
        "&internal_call=false&proxy=false"
    )
    resp = _post(endpoint, {
        "chatHPE_bot_data": {"appId": "1", "sessionId": session_id,
                             "userId": user_id, "username": username},
        "model_name": "gpt-4o-mini",
        "stream": False, "webScraping": False,
        "user_query": prompt,
    })
    return resp.get("Response", "")


# ─────────────────────────────────────────────────────────────────────────────


# ── Mandatory sections that must always appear in every TDS ───────────────────

MANDATORY_REQUIREMENT_SECTIONS = [
    {
        "title": "Requirement Pseudocode",
        "diagram_type": "text",
        "focus": "Structured pseudocode for the implemented requirement flow based on FDS.",
    },
    {
        "title": "Requirement Test Cases (Low, Medium, Extreme)",
        "diagram_type": "text",
        "focus": "Executable-style test cases that validate low, medium, and extreme scenarios.",
    },
]

TEST_CASES_ONLY_SECTION = {
    "title": "Requirement Test Cases (Low, Medium, Extreme)",
    "diagram_type": "text",
    "focus": "Executable-style test cases that validate low, medium, and extreme scenarios.",
}

JUNIT_TEST_SECTION = {
    "title": "JUnit 5 Test Skeleton",
    "diagram_type": "text",
    "focus": "Executable JUnit 5 starter tests mapped to requirement test cases.",
}


# ── LLM Prompts ───────────────────────────────────────────────────────────────

PROMPT_REQUIREMENT_PSEUDOCODE = """You are the TDS Agent.
Generate implementation-ready pseudocode for requirement coverage from FDS context.

Rules:
- Output plain text only. No markdown fences.
- Begin with: Requirement Pseudocode
- Include 3 to 6 named flows that reflect core requirements from FDS.
- Each flow must follow this exact block format:
    Flow <number>: <Flow Name>
    Inputs: <comma-separated inputs>
    Preconditions: <preconditions>
    Steps:
    1. FUNCTION <Name>(<params>)
    2. IF <condition> THEN <action>
    3. ELSE <action>
    4. RETURN <result>
    Postconditions: <postconditions>
    Error Paths: <error conditions>
- Use deterministic, concise pseudocode style (IF/ELSE, FOR, WHILE, FUNCTION).
- Keep step numbering numeric and continuous within each flow.
- Keep content implementation-neutral but specific enough to code directly.

Project: {project_name}

FDS Content:
{fds_context}
"""

PROMPT_REQUIREMENT_TEST_CASES = """You are the TDS Agent.
Generate requirement validation test cases for low, medium, and extreme conditions.

Rules:
- Output plain text only. No markdown fences.
- Begin with: Requirement Test Cases (Low, Medium, Extreme)
- Include sections in this exact order: Low Cases, Medium Cases, Extreme Cases.
- Provide at least 3 test cases per section.
- For each test case include: Test ID, Requirement Mapping, Input, Steps, Expected Result, Priority.
- Ensure expected results are measurable and implementation-verifiable.

Project: {project_name}

FDS Content:
{fds_context}
"""


PROMPT_JUNIT_TEST_CODE = """You are the TDS Agent.
Generate Java JUnit 5 starter test code from requirement test cases.

Rules:
- Output plain text Java code only. No markdown fences.
- Use JUnit 5 imports (org.junit.jupiter.api.*).
- Provide one test class named {class_name}.
- Include at least 6 test methods that map to low/medium/extreme requirement scenarios.
- Each test must include: Arrange/Act/Assert comments and at least one assertion.
- Use placeholder service/domain classes when needed (for example: RequestService, RequestResult).
- Keep code compile-friendly as a skeleton (minimal dependencies, TODO markers allowed).

Project: {project_name}

Requirement Test Cases:
{test_cases_text}

FDS Content:
{fds_context}
"""


# ── Static fallbacks (used when the LLM returns nothing usable) ───────────────

REQUIREMENT_PSEUDOCODE_FALLBACK = """Requirement Pseudocode

Flow 1: Submit Request
Inputs: actor_id, payload
Preconditions: actor_id is authenticated; payload schema is valid
Steps:
1. FUNCTION SubmitRequest(actor_id, payload)
2.   IF NotAuthenticated(actor_id) THEN RETURN Unauthorized
3.   IF NotValid(payload) THEN RETURN ValidationError
4.   request_id = SaveRequest(payload, status="PENDING")
5.   EmitAudit("REQUEST_SUBMITTED", request_id)
6.   NotifyApprover(request_id)
7.   RETURN Success(request_id)
Postconditions: request persisted with PENDING status; audit recorded; notification sent
Error Paths: auth failure, schema failure, persistence failure, notification failure

Flow 2: Validate Business Rules
Inputs: request_id
Preconditions: request exists
Steps:
1. FUNCTION ValidateRules(request_id)
2.   data = FetchRequest(request_id)
3.   IF data is null THEN RETURN NotFound
4.   IF ViolatesPolicy(data) THEN RETURN RejectedWithReason
5.   RETURN Valid
Postconditions: decision returned with reason when invalid
Error Paths: missing request, rule engine failure

Flow 3: Approve or Reject
Inputs: approver_id, request_id, decision
Preconditions: approver has role; request status is PENDING
Steps:
1. FUNCTION ProcessDecision(approver_id, request_id, decision)
2.   IF NotAuthorized(approver_id) THEN RETURN Forbidden
3.   IF CurrentStatus(request_id) != "PENDING" THEN RETURN Conflict
4.   IF decision == "APPROVE" THEN UpdateStatus(request_id, "APPROVED")
5.   ELSE UpdateStatus(request_id, "REJECTED")
6.   EmitAudit("DECISION_RECORDED", request_id)
7.   NotifyRequester(request_id)
8.   RETURN Success
Postconditions: status finalized; audit recorded; requester notified
Error Paths: authorization error, concurrent update, write failure
"""

REQUIREMENT_TEST_CASES_FALLBACK = """Requirement Test Cases (Low, Medium, Extreme)

Low Cases
1. Test ID: LOW-001
Requirement Mapping: Request submission validation
Input: Minimal valid payload
Steps: Submit one request with mandatory fields only
Expected Result: Request accepted with PENDING status and request ID returned
Priority: High

2. Test ID: LOW-002
Requirement Mapping: Approval workflow
Input: Valid pending request and valid approver
Steps: Approver submits APPROVE decision
Expected Result: Status changes to APPROVED; audit log and notification are generated
Priority: High

3. Test ID: LOW-003
Requirement Mapping: Rejection workflow
Input: Valid pending request and valid approver
Steps: Approver submits REJECT decision
Expected Result: Status changes to REJECTED with reason recorded
Priority: Medium

Medium Cases
1. Test ID: MED-001
Requirement Mapping: Concurrent processing control
Input: Same request updated by two approvers nearly simultaneously
Steps: Send two decisions within short interval
Expected Result: Exactly one decision succeeds; second receives conflict response
Priority: High

2. Test ID: MED-002
Requirement Mapping: Policy rule enforcement
Input: Payload violating one business rule
Steps: Submit request with rule-violating data
Expected Result: Request rejected with specific policy message; no persistence side effects
Priority: High

3. Test ID: MED-003
Requirement Mapping: Retry/reliability behavior
Input: Temporary downstream notification failure
Steps: Submit valid request while notification endpoint intermittently fails
Expected Result: Core transaction succeeds; retry mechanism logs and eventually sends notification
Priority: Medium

Extreme Cases
1. Test ID: EXT-001
Requirement Mapping: High-volume ingestion
Input: Burst of requests at projected peak load
Steps: Submit requests at peak throughput for sustained interval
Expected Result: Throughput and latency remain within NFR thresholds; no data loss
Priority: Critical

2. Test ID: EXT-002
Requirement Mapping: Large payload robustness
Input: Maximum allowed payload size and edge field lengths
Steps: Submit request with boundary-size content
Expected Result: Request processed correctly or rejected with clear bounded-error response
Priority: High

3. Test ID: EXT-003
Requirement Mapping: Recovery and consistency
Input: Forced service restart during in-flight transaction
Steps: Trigger restart while decision update is processing; replay request check
Expected Result: System recovers to a consistent state with idempotent final status
Priority: Critical
"""


JUNIT_TEST_CODE_FALLBACK = """import org.junit.jupiter.api.DisplayName;
import org.junit.jupiter.api.Test;

import static org.junit.jupiter.api.Assertions.assertNotNull;
import static org.junit.jupiter.api.Assertions.assertTrue;

class RequirementWorkflowTest {

    private final RequestService requestService = new RequestService();

    @Test
    @DisplayName("LOW: submit minimal valid request")
    void shouldSubmitMinimalValidRequest() {
        // Arrange
        RequestPayload payload = RequestPayload.minimalValid();

        // Act
        RequestResult result = requestService.submit(payload);

        // Assert
        assertNotNull(result);
        assertTrue(result.isPending());
    }

    @Test
    @DisplayName("LOW: approve pending request")
    void shouldApprovePendingRequest() {
        // Arrange
        RequestId id = requestService.submit(RequestPayload.minimalValid()).requestId();

        // Act
        RequestResult result = requestService.approve(id, "approver-1");

        // Assert
        assertTrue(result.isApproved());
    }

    @Test
    @DisplayName("MEDIUM: reject policy violating request")
    void shouldRejectPolicyViolatingRequest() {
        // Arrange
        RequestPayload payload = RequestPayload.policyViolating();

        // Act
        RequestResult result = requestService.submit(payload);

        // Assert
        assertTrue(result.isRejected());
    }

    @Test
    @DisplayName("MEDIUM: prevent double decision conflict")
    void shouldPreventConflictingDecisions() {
        // Arrange
        RequestId id = requestService.submit(RequestPayload.minimalValid()).requestId();
        requestService.approve(id, "approver-1");

        // Act
        RequestResult secondDecision = requestService.reject(id, "approver-2", "late decision");

        // Assert
        assertTrue(secondDecision.isConflict());
    }

    @Test
    @DisplayName("EXTREME: sustain peak request bursts")
    void shouldHandlePeakRequestBurst() {
        // Arrange
        int accepted = 0;

        // Act
        for (int i = 0; i < 1000; i++) {
            RequestResult result = requestService.submit(RequestPayload.minimalValid());
            if (result.isPending()) accepted++;
        }

        // Assert
        assertTrue(accepted > 0);
    }

    @Test
    @DisplayName("EXTREME: recover consistent state after restart")
    void shouldRecoverConsistentStateAfterRestart() {
        // Arrange
        RequestId id = requestService.submit(RequestPayload.minimalValid()).requestId();

        // Act
        requestService.simulateRestart();
        RequestResult status = requestService.getStatus(id);

        // Assert
        assertNotNull(status);
        assertTrue(status.isTerminal() || status.isPending());
    }
}
"""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _strip_fences(text: str) -> str:
    """Remove markdown/code fences from LLM output."""
    fenced = re.search(r"```(?:mermaid)?\s*(.*?)```", text, flags=re.S | re.I)
    if fenced:
        return fenced.group(1).strip()
    return text.strip()


def _sanitize_multiline(text: str) -> str:
    """Normalise whitespace in multiline plain text."""
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = []
    blank_pending = False
    for line in lines:
        if not line.strip():
            if cleaned and not blank_pending:
                cleaned.append("")
            blank_pending = True
            continue
        cleaned.append(line.rstrip())
        blank_pending = False
    return "\n".join(cleaned).strip()


def _to_java_class_name(project_name: str) -> str:
    chunks = re.findall(r"[A-Za-z0-9]+", project_name or "")
    if not chunks:
        return "RequirementWorkflowTest"
    base = "".join(part[:1].upper() + part[1:] for part in chunks)
    return f"{base}RequirementTest"


# ── Public API ────────────────────────────────────────────────────────────────

def enforce_mandatory_requirement_sections(section_plan: list) -> list:
    """
    Append any missing mandatory requirement sections to the plan.
    Operates on a copy; never mutates the caller's list.
    Caps total sections at 10.
    """
    plan = list(section_plan or [])
    existing_titles = {
        " ".join(str(s.get("title", "")).strip().lower().split())
        for s in plan
        if isinstance(s, dict)
    }

    for mandatory in MANDATORY_REQUIREMENT_SECTIONS:
        title_key = " ".join(mandatory["title"].strip().lower().split())
        if title_key not in existing_titles:
            plan.append(dict(mandatory))
            existing_titles.add(title_key)

    return plan[:10]


def generate_requirement_section(
    section_name: str,
    project_name: str,
    fds_context: str,
    ask_fn,
) -> str:
    """
    Generate content for a text-type requirement section.

    Parameters
    ----------
    section_name : str
        Exact section title from the section plan.
    project_name : str
        Project name for prompt context.
    fds_context : str
        Full FDS content string.
    ask_fn : callable
        Function with signature ask_fn(prompt: str) -> str.
        The caller is responsible for providing the session-bound version.

    Returns
    -------
    str
        Generated (or fallback) plain-text content for the section.
    """
    name_lower = section_name.strip().lower()

    if "pseudocode" in name_lower:
        prompt = PROMPT_REQUIREMENT_PSEUDOCODE.format(
            project_name=project_name,
            fds_context=fds_context,
        )
        fallback = REQUIREMENT_PSEUDOCODE_FALLBACK
    elif "test cases" in name_lower:
        prompt = PROMPT_REQUIREMENT_TEST_CASES.format(
            project_name=project_name,
            fds_context=fds_context,
        )
        fallback = REQUIREMENT_TEST_CASES_FALLBACK
    else:
        # Unknown text section — return empty string; caller handles it
        return ""

    raw = ask_fn(prompt)
    cleaned = _sanitize_multiline(_strip_fences(raw))
    return cleaned or fallback


# ── Standalone pipeline ───────────────────────────────────────────────────────

def run_test_cases_from_card(card: dict) -> dict:
    """
    Generate test cases from a TDS A2A card that contains
    FDS sections in payload.fds_sections.

    Returns
    -------
    dict with keys:
        project_name, sections (dict title→text), section_order (list)
    """
    project = card.get("project", {})
    project_name = project.get("name") or "Project"

    payload = card.get("payload", {})
    fds_sections = payload.get("fds_sections", {})
    summary = payload.get("summary", {})
    raw_fds = payload.get("raw_fds", "")

    parts = [f"## {k}\n{v}" for k, v in fds_sections.items()]
    if summary:
        parts.append("## FDS Summary\n" + "\n".join(f"{k}: {v}" for k, v in summary.items()))
    if raw_fds:
        parts.append(f"## Raw FDS\n{raw_fds}")
    fds_context = "\n\n".join(parts).strip()

    user_id, username, session_id = _setup_session()
    ask_fn = lambda p: _ask(p, user_id, username, session_id)

    sections = {}
    section_order = []

    title = TEST_CASES_ONLY_SECTION["title"]
    print(f"[test-cases] Generating: {title} …")
    test_cases_text = generate_requirement_section(title, project_name, fds_context, ask_fn)
    sections[title] = test_cases_text
    section_order.append(title)

    junit_title = JUNIT_TEST_SECTION["title"]
    junit_prompt = PROMPT_JUNIT_TEST_CODE.format(
        class_name=_to_java_class_name(project_name),
        project_name=project_name,
        test_cases_text=test_cases_text,
        fds_context=fds_context,
    )
    print(f"[test-cases] Generating: {junit_title} …")
    junit_raw = ask_fn(junit_prompt)
    junit_text = _sanitize_multiline(_strip_fences(junit_raw)) or JUNIT_TEST_CODE_FALLBACK
    sections[junit_title] = junit_text
    section_order.append(junit_title)

    return {
        "project_name": project_name,
        "sections": sections,
        "section_order": section_order,
    }


# ── DOCX builder ──────────────────────────────────────────────────────────────

_HPE_BLUE = RGBColor(0x1F, 0x49, 0x7D)


def _tc_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = _HPE_BLUE
    return para


def _tc_body(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    return para


def _tc_preformatted(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return para


def build_test_cases_docx(
    project_name: str,
    sections: dict,
    section_order: list,
    source_card_name: str,
) -> Document:
    """Build and return a python-docx Document for the test cases report."""
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    title = doc.add_heading("Test Cases & Pseudocode Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(project_name).font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Date: {datetime.now().strftime('%Y-%m-%d')}\n"
        f"Prepared By: {_USER_NAME}\n"
        f"Source Card: {source_card_name}"
    )

    doc.add_page_break()

    _tc_heading(doc, "1. Document Control", level=1)
    _tc_body(
        doc,
        "Version: 1.0\n"
        "Status: Draft\n"
        f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Generated From: {source_card_name}",
    )

    for idx, section_name in enumerate(section_order, start=2):
        _tc_heading(doc, f"{idx}. {section_name}", level=1)
        content = (sections.get(section_name) or "Not generated.").strip()

        # Render line-by-line as either preformatted (code-like) or body text
        block = []
        for line in content.splitlines():
            stripped = line.rstrip()
            if not stripped:
                if block:
                    _tc_body(doc, "\n".join(block).strip())
                    block = []
                continue
            # Lines that look like code/pseudocode go in preformatted blocks
            if (
                stripped.startswith(("FUNCTION", "IF ", "ELSE", "FOR ", "WHILE ",
                                     "RETURN ", "  ", "\t", "1.", "2.", "3."))
                or stripped.startswith(("import ", "class ", "public ", "private ", "protected ", "@"))
                or stripped.startswith("-")
            ):
                if block:
                    _tc_body(doc, "\n".join(block).strip())
                    block = []
                _tc_preformatted(doc, stripped)
            else:
                block.append(stripped)
        if block:
            _tc_body(doc, "\n".join(block).strip())

    return doc

